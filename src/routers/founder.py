"""
src/routers/founder.py — Router de Founder (upload + confirm + finalize).

Rutas incluidas:
  POST /upload                       — subir PDF/Excel/CSV (Gemini o Mapper)
  POST /api/founder/confirm-mapping  — confirmar y persistir en BQ
  POST /api/founder/manual-update    — corregir KPIs manualmente en GCS
  GET  /api/founder/config           — auto-detectar empresa del JWT
  POST /api/founder/finalize         — cerrar expediente + email de recibo

Estrategia de no-circular-imports:
  Este módulo NO importa desde src.api. Solo importa:
    - src.dependencies.auth    (JWT + security helpers)
    - src.core.*               (lógica de negocio)
    - src.adapters.*           (GCP, Document AI)
    - stdlib / third-party

  src.api importa ESTE módulo para include_router() — la dependencia es
  de una sola dirección.

Nota de migración (Phase 3):
  Las funciones privadas _get_storage_client, get_company_id, los helpers
  de Gemini (_build_gemini_kpi_schema, _process_tabular, etc.) y los helpers
  GCS se extrajeron de src/api.py aquí. Cuando se creen los routers de
  analyst y admin, las funciones GCS se promoverán a src/core/gcs.py y
  las de Gemini a src/core/gemini_processor.py para compartirlas.
"""
from __future__ import annotations

import copy
import hashlib
import json
import os
import re
import unicodedata
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

import pandas as pd
from fastapi import APIRouter, Depends, File, Form, Header, HTTPException, Request, UploadFile
from fastapi.responses import JSONResponse
from google.api_core.exceptions import Forbidden, Unauthorized
from google.auth.exceptions import DefaultCredentialsError
from google.cloud import storage
from google.oauth2 import service_account
from pydantic import BaseModel

from src.adapters.document_ai import DocumentAIAdapter
from src.adapters.google_cloud import GeminiAuditor
from src.core.data_contract import (
    KPI_REGISTRY,
    build_checklist_status,
    build_contract,
)
from src.core.db_writer import (
    COMPANY_BUCKET,
    PORTFOLIO_MAP,
    detect_company_from_text,
    insert_upload_log,
    lookup_portfolio,
    write_founder_rescue_notes,
)
from src.core.kpi_dispatcher import (
    IngestionMetadata,
    SubmissionBlockedError,
    build_upload_preview,
    dispatch_to_storage,
    get_prev_cash_from_bq,
    upload_gold_layer,
    upload_raw_layer,
    upload_stage_layer,
)
from src.core.kpi_mapper import map_uploaded_file
from src.core.local_db import (
    build_accumulated_kpi_grid,
    build_enriched_audit_response,
    build_kpi_grid_from_contract_rows,
    build_registry_rows,
    build_registry_rows_from_contract,
    check_gate_for_finalize,
    evaluate_commitment_gate,
    generate_jero_contract,
    save_registry_rows,
)
from src.dependencies.auth import (
    MAX_FILE_BYTES,
    MAX_FILE_MB,
    limiter,
    require_auth,
    sanitize_filename,
    validate_company_header,
    validate_email_header,
    validate_magic_bytes,
)

router = APIRouter(tags=["founder"])

# ── GCP / GCS configuration ───────────────────────────────────────────────────
_PROJECT_ID       = os.getenv("GOOGLE_PROJECT_ID", "cometa-mvp")
_LOCATION_DOC_AI  = os.getenv("DOCUMENT_AI_LOCATION", "us")
_PROCESSOR_ID     = os.getenv("DOCUMENT_AI_PROCESSOR_ID", "c5e1adfde68e63cf")
_VERTEX_LOCATION  = os.getenv("VERTEX_AI_LOCATION", "us-central1")
from src.core.buckets import RAW_BUCKET as _GCS_INPUT_BUCKET, STAGE_BUCKET as _GCS_OUTPUT_BUCKET

_BUCKET_TO_VERTICAL: dict[str, str] = {
    "SAAS":  "SAAS",
    "LEND":  "FINTECH",
    "ECOM":  "MARKETPLACE",
    "INSUR": "INSURTECH",
    "OTH":   "GENERAL",
}

# ── Pydantic request models ───────────────────────────────────────────────────

class ConfirmMappingRequest(BaseModel):
    """Body para POST /api/founder/confirm-mapping."""
    load_id:        str
    filename:       str
    company_slug:   str
    company_name:   str
    period_str:     str   # "YYYY-MM"
    sector:         str
    source_type:    str = "verified"


class ManualUpdateRequest(BaseModel):
    """Body for POST /api/founder/manual-update."""
    file_hash: str
    updates:   dict[str, str]


class FinalizeRequest(BaseModel):
    """Body for POST /api/founder/finalize."""
    file_hashes:    list[str]
    company_domain: str = ""   # ignorado — se deriva del JWT (Zero Trust)
    file_names:     list[str] = []
    manual_kpis:    dict[str, str] | None = None


# ═════════════════════════════════════════════════════════════════════════════
# HELPERS PRIVADOS (extraídos de src/api.py durante Phase 3)
# ═════════════════════════════════════════════════════════════════════════════

# ── GCS credential helpers ────────────────────────────────────────────────────

def _resolve_sa_path() -> str | None:
    env_path = os.getenv("GOOGLE_APPLICATION_CREDENTIALS")
    if env_path:
        return env_path
    fallback = os.path.join(
        os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "..")),
        "cometa_key.json",
    )
    return fallback if os.path.exists(fallback) else None


def _parse_sa_json(raw: str) -> dict:
    raw = raw.strip()
    parsed = json.loads(raw)
    if isinstance(parsed, str):
        parsed = json.loads(parsed)
    required = {"type", "project_id", "private_key", "client_email"}
    missing  = required - parsed.keys()
    if missing:
        raise ValueError(f"GCP_SERVICE_ACCOUNT_JSON le faltan campos: {missing}")
    return parsed


def _load_gcp_credentials():
    sa_json_str = os.getenv("GCP_SERVICE_ACCOUNT_JSON")
    if sa_json_str:
        return service_account.Credentials.from_service_account_info(
            _parse_sa_json(sa_json_str)
        )
    sa_path = _resolve_sa_path()
    if not sa_path or not os.path.exists(sa_path):
        raise DefaultCredentialsError(
            "No se encontró GCP_SERVICE_ACCOUNT_JSON ni GOOGLE_APPLICATION_CREDENTIALS"
        )
    return service_account.Credentials.from_service_account_file(sa_path)


def _get_storage_client() -> storage.Client:
    """Crea un Storage client usando credenciales explícitas cuando es posible."""
    try:
        return storage.Client(project=_PROJECT_ID, credentials=_load_gcp_credentials())
    except Exception as e:
        print(f"[GCP] No se pudieron cargar credenciales explícitas: {e}")
        return storage.Client(project=_PROJECT_ID)


# ── GCS hash helpers ──────────────────────────────────────────────────────────

def get_file_hash(file_content: bytes) -> str:
    return hashlib.sha256(file_content).hexdigest()[:16]


def _ensure_dict(obj: Any) -> dict:
    if isinstance(obj, dict):
        return obj
    if isinstance(obj, str):
        parsed = json.loads(obj)
        if isinstance(parsed, str):
            return json.loads(parsed)
        if isinstance(parsed, dict):
            return parsed
        raise TypeError(f"_ensure_dict: json.loads devolvió {type(parsed).__name__}, no dict")
    raise TypeError(f"_ensure_dict: esperaba dict o str, recibió {type(obj).__name__}")


def check_hash_exists_in_gcs(bucket_name: str, file_hash: str) -> bool:
    try:
        sc     = _get_storage_client()
        bucket = sc.bucket(bucket_name)
        for blob in bucket.list_blobs():
            if blob.metadata and blob.metadata.get("file_hash") == file_hash:
                return True
        return False
    except (Forbidden, Unauthorized) as e:
        sa_path = _resolve_sa_path()
        print(f"[GCS] Error de credenciales: {e}  SA={sa_path}")
        raise RuntimeError("GCS_AUTH") from e
    except Exception as e:
        print(f"[GCS] Error verificando hash: {e}")
        raise RuntimeError("GCS_ERROR") from e


def get_existing_result(bucket_name: str, file_hash: str) -> dict:
    try:
        sc     = _get_storage_client()
        bucket = sc.bucket(bucket_name)
        for blob in bucket.list_blobs(prefix="staging/"):
            if blob.name.endswith(".json") and file_hash in blob.name:
                return _ensure_dict(json.loads(blob.download_as_text()))
        return {}
    except Exception as e:
        print(f"[GCS] Error obteniendo resultado existente: {e}")
        return {}


# ── Company catalog helpers ───────────────────────────────────────────────────
# TODO (Phase 4): promover a src/core/company.py cuando el router de analyst
#                 también necesite get_company_id.

_COMPANY_CATALOG: dict[str, dict] = {}   # lazy init


def _build_company_catalog() -> dict[str, dict]:
    catalog: dict[str, dict] = {}
    for key, info in PORTFOLIO_MAP.items():
        bucket   = COMPANY_BUCKET.get(key, "OTH")
        comp_id  = f"COMP_{key.upper().replace('-','_').replace('.','_')}"
        catalog[key] = {
            "comp_id":   comp_id,
            "fund_id":   info["portfolio_id"],
            "bucket_id": bucket,
        }
    return catalog


def _get_company_catalog() -> dict[str, dict]:
    global _COMPANY_CATALOG
    if not _COMPANY_CATALOG:
        _COMPANY_CATALOG = _build_company_catalog()
    return _COMPANY_CATALOG


def get_company_id(name_str: str) -> tuple[str, str, str, bool]:
    """
    Mapea texto libre al ID canónico COMP_XXX.

    Returns
    -------
    (canonical_id, fund_id, bucket_id, is_known)
    """
    catalog = _get_company_catalog()
    s       = name_str.strip().lower()

    # Strip TLD (.com, .mx, .co …)
    s_no_tld = re.sub(r"\.(com|mx|co|vc|ai|io|app|net|org|fund)$", "", s)

    def _try(key: str) -> tuple[str, str, str] | None:
        entry = catalog.get(key)
        if entry:
            return entry["comp_id"], entry["fund_id"], entry["bucket_id"]
        return None

    # 1. Exact match
    for candidate in (s, s_no_tld):
        r = _try(candidate)
        if r:
            return (*r, True)

    # 2. Strip hyphens/underscores
    collapsed = re.sub(r"[-_]", "", s_no_tld)
    for key in catalog:
        if re.sub(r"[-_]", "", key) == collapsed:
            e = catalog[key]
            return e["comp_id"], e["fund_id"], e["bucket_id"], True

    # 3. Prefix / substring
    for key, entry in catalog.items():
        if s_no_tld.startswith(key) or key.startswith(s_no_tld):
            return entry["comp_id"], entry["fund_id"], entry["bucket_id"], True
    for key, entry in catalog.items():
        if key in s_no_tld or s_no_tld in key:
            return entry["comp_id"], entry["fund_id"], entry["bucket_id"], True

    # Unknown company
    _h      = hashlib.sha256(name_str.encode()).hexdigest()[:6].upper()
    unknown = f"COMP_UNKNOWN_{_h}"
    return unknown, "unknown", "OTH", False


# ── Gemini / Document AI helpers ──────────────────────────────────────────────
# TODO (Phase 4): promover a src/core/gemini_processor.py

_MAX_ROWS_PER_SHEET = 500
_PDF_CHUNK_SIZE     = 90

_REQUIRED_FM_SECTIONS = (
    "revenue_growth", "profit_margins", "cash_flow_indicators",
    "debt_ratios", "base_metrics", "sector_metrics",
)


def _ensure_fm_sections(gemini_json: dict) -> dict:
    fm = gemini_json.setdefault("financial_metrics_2025", {})
    for section in _REQUIRED_FM_SECTIONS:
        fm.setdefault(section, {})
    return gemini_json


def _build_gemini_kpi_schema() -> str:
    brain_path = Path(__file__).parent.parent.parent / "assets" / "loading_brain_v1.json"
    try:
        with open(brain_path, encoding="utf-8") as f:
            brain = json.load(f)

        given = [m for m in brain.get("metrics", []) if m.get("given_or_silver") == "GIVEN"]
        if not given:
            raise ValueError("No GIVEN metrics found in brain")

        cats: dict[str, list] = {}
        for m in sorted(given, key=lambda x: x.get("kpi_ref", "")):
            cats.setdefault(m.get("category", "Other"), []).append(m)

        kpi_lines: list[str] = []
        all_items = list(cats.items())
        for cat_idx, (cat, kpis) in enumerate(all_items):
            kpi_lines.append(f'    /* -- {cat} -- */')
            for k_idx, m in enumerate(kpis):
                mid   = m["metric_id"]
                dname = m["display_name"]
                inneg = " INNEGOCIABLE" if m.get("innegociable") else ""
                dtype = m.get("data_type", "")
                unit  = m.get("unit", "")
                sector = m.get("sector_scope", "ALL")
                sector_note = f" | sector: {sector}" if sector != "ALL" else ""
                if dtype == "percentage":
                    ex = '"36%" o "-4.2%"'
                elif dtype == "currency":
                    ex = '"$4.2M" o "-$320K"'
                elif dtype == "integer":
                    ex = '"42"'
                else:
                    ex = f'"{unit or "valor"}"'
                is_last = (cat_idx == len(all_items) - 1) and (k_idx == len(kpis) - 1)
                comma = "" if is_last else ","
                kpi_lines.append(
                    f'    "{mid}": {{"value": "<{dname}{inneg}{sector_note} — ej: {ex}>",'
                    f' "confidence": <float 0.0-1.0>,'
                    f' "description": "<fuente exacta>"}}{comma}'
                )

        inner = "\n".join(kpi_lines)
        return (
            "{\n"
            '  "_document_context": {\n'
            '    "currency": "<ISO 4217>", "period": "<FY/H1/Q>",\n'
            '    "scale": "<units|thousands|millions|billions>", "scale_notes": null\n'
            "  },\n"
            '  "financial_metrics_2025": {\n'
            + inner + "\n  }\n}"
        )
    except Exception as exc:
        print(f"[prompt] WARN: schema dinámico no disponible: {exc}")
        return (
            "{\n"
            '  "_document_context": {\n'
            '    "currency": "<ISO 4217>", "period": "<FY/H1/Q>",\n'
            '    "scale": "<units|thousands|millions|billions>", "scale_notes": null\n'
            "  },\n"
            '  "financial_metrics_2025": {\n'
            '    "revenue":      {"value": "<Revenue — ej: \'$4.2M\'>",   "confidence": 0.9, "description": "<fuente>"},\n'
            '    "gross_profit": {"value": "<Gross Profit — ej: \'$1.8M\'>","confidence": 0.9, "description": "<fuente>"},\n'
            '    "ebitda":       {"value": "<EBITDA>",                     "confidence": 0.9, "description": "<fuente>"},\n'
            '    "cash":         {"value": "<Caja — ej: \'$9.7M\'>",       "confidence": 0.9, "description": "<fuente>"},\n'
            '    "burn":         {"value": "<Burn mensual — ej: \'-$320K\'>","confidence":0.9,"description": "<fuente>"},\n'
            '    "mrr":          {"value": "<MRR SaaS — ej: \'$350K\'>",   "confidence": 0.9, "description": "<fuente>"},\n'
            '    "employees":    {"value": "<Headcount — ej: \'42\'>",     "confidence": 0.9, "description": "<fuente>"}\n'
            "  }\n}"
        )


def _df_to_markdown(df: pd.DataFrame) -> str:
    if df.empty:
        return "*(tabla vacía)*"
    cols      = [str(c).strip() for c in df.columns]
    header    = "| " + " | ".join(cols) + " |"
    separator = "| " + " | ".join(["---"] * len(cols)) + " |"
    rows = []
    for _, row in df.iterrows():
        cells = [
            str(v).replace("|", "\\|").strip() if pd.notna(v) else ""
            for v in row
        ]
        rows.append("| " + " | ".join(cells) + " |")
    return "\n".join([header, separator] + rows)


def _process_tabular(file_path: str, ext: str, gemini: Any, prompt_schema: str) -> str:
    print(f"[Tabular] Leyendo archivo {ext}...")
    try:
        sheets: dict = {}
        if ext == ".csv":
            sheets["Hoja1"] = pd.read_csv(file_path, nrows=_MAX_ROWS_PER_SHEET)
        elif ext in (".xlsx", ".xls"):
            all_sheets = pd.read_excel(file_path, sheet_name=None, engine="openpyxl")
            for name, df in all_sheets.items():
                sheets[str(name)] = df.head(_MAX_ROWS_PER_SHEET)
        elif ext == ".parquet":
            sheets["Hoja1"] = pd.read_parquet(file_path).head(_MAX_ROWS_PER_SHEET)
        else:
            raise ValueError(f"Extensión tabular no reconocida: {ext}")
    except Exception as e:
        raise RuntimeError(f"No se pudo leer el archivo {ext}: {e}") from e

    md_sections = []
    for sheet_name, df in sheets.items():
        df = df.dropna(axis=1, how="all")
        md_sections.append(f"## Pestaña: {sheet_name}\n\n{_df_to_markdown(df)}")
    full_markdown = "\n\n---\n\n".join(md_sections)

    _n_sheets = len(sheets)
    adapter_header = (
        f"Eres un auditor financiero senior especializado en due diligence de startups.\n"
        f"Recibes el contenido completo de un archivo {ext.upper()} con "
        f"{_n_sheets} pestaña(s), convertido a tablas Markdown.\n"
        f"Analiza TODAS las pestañas para localizar las métricas financieras.\n"
        f"Cuando una métrica aparezca en varias hojas, usa la fuente más reciente.\n\n"
        f"Aplica las siguientes instrucciones de extracción EXACTAMENTE:\n\n"
    )
    return gemini.analizar_texto(adapter_header + prompt_schema, full_markdown)


def _process_docx(file_path: str, gemini: Any, prompt_schema: str) -> str:
    try:
        from docx import Document as DocxDocument  # type: ignore
        doc        = DocxDocument(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        table_rows = []
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    table_rows.append(row_text)
        full_text = "\n".join(paragraphs)
        if table_rows:
            full_text += "\n\nTABLAS DEL DOCUMENTO:\n" + "\n".join(table_rows)
    except ImportError:
        raise RuntimeError("python-docx no instalado. Ejecuta: pip install python-docx")
    except Exception as e:
        raise RuntimeError(f"No se pudo procesar el DOCX: {e}") from e

    instruction = (
        "Eres un auditor financiero senior. Analiza el siguiente documento Word (DOCX) "
        "y extrae las métricas del texto y tablas. Emite el JSON estándar.\n\n"
        "CONTENIDO DEL DOCUMENTO:\n"
    )
    return gemini.analizar_texto(instruction + full_text + "\n\n" + prompt_schema, "")


def split_pdf_to_chunks(file_bytes: bytes, size: int = _PDF_CHUNK_SIZE) -> list[bytes]:
    try:
        import fitz  # type: ignore
    except ImportError:
        raise RuntimeError("PyMuPDF no instalado. Ejecuta: pip install pymupdf")

    src_doc     = fitz.open(stream=file_bytes, filetype="pdf")
    total_pages = len(src_doc)

    if total_pages <= size:
        src_doc.close()
        return [file_bytes]

    n_chunks = (total_pages + size - 1) // size
    chunks: list[bytes] = []
    for idx in range(n_chunks):
        start     = idx * size
        end       = min(start + size - 1, total_pages - 1)
        chunk_doc = fitz.open()
        chunk_doc.insert_pdf(src_doc, from_page=start, to_page=end)
        chunks.append(chunk_doc.tobytes())
        chunk_doc.close()
        print(f"[Chunking] Bloque {idx+1}/{n_chunks}: pags {start+1}-{end+1}")
    src_doc.close()
    return chunks


def merge_consolidated_results(jsons: list[dict]) -> dict:
    """Une resultados de múltiples chunks de Gemini en uno solo."""
    if not jsons:
        raise ValueError("merge_consolidated_results: lista vacía")
    if len(jsons) == 1:
        return _ensure_fm_sections(jsons[0])

    merged = copy.deepcopy(jsons[0])

    def _get(obj: dict, path: list[str]) -> Any:
        cur = obj
        for k in path:
            if not isinstance(cur, dict):
                return None
            cur = cur.get(k)
        return cur

    def _set(obj: dict, path: list[str], value: Any) -> None:
        cur = obj
        for k in path[:-1]:
            cur = cur.setdefault(k, {})
        cur[path[-1]] = value

    for chunk_json in jsons[1:]:
        for kpi_def in KPI_REGISTRY:
            path     = kpi_def["path"]
            existing = _get(merged, path)
            incoming = _get(chunk_json, path)
            if not isinstance(incoming, dict) or incoming.get("value") is None:
                continue
            if not isinstance(existing, dict) or existing.get("value") is None:
                _set(merged, path, incoming)
                continue
            if float(incoming.get("confidence") or 0) > float(existing.get("confidence") or 0):
                _set(merged, path, incoming)

    _ensure_fm_sections(merged)
    return merged


def _chunk_and_process_pdf(temp_path: str, gemini: Any, prompt_config: str) -> str:
    with open(temp_path, "rb") as fh:
        file_bytes = fh.read()

    chunks   = split_pdf_to_chunks(file_bytes, size=_PDF_CHUNK_SIZE)
    n_chunks = len(chunks)

    if n_chunks == 1:
        print(f"[Chunking] PDF <= {_PDF_CHUNK_SIZE} pags — llamada directa a Gemini")
        return gemini.extraer_y_auditar(temp_path, prompt_config)

    print(f"[Chunking] {n_chunks} bloques de hasta {_PDF_CHUNK_SIZE} pags")
    chunk_results: list[dict] = []
    for i, chunk_bytes in enumerate(chunks):
        chunk_path = f"{temp_path}_chunk{i}.pdf"
        try:
            with open(chunk_path, "wb") as cf:
                cf.write(chunk_bytes)
            raw = gemini.extraer_y_auditar(chunk_path, prompt_config)
            if isinstance(raw, str):
                try:
                    parsed = json.loads(raw)
                except json.JSONDecodeError:
                    clean  = re.sub(r'^```json\s*|\s*```$', '', raw.strip())
                    parsed = json.loads(clean)
            else:
                parsed = raw
            chunk_results.append(parsed)
            print(f"[Chunking] Bloque {i+1} OK")
        except Exception as chunk_err:
            print(f"[Chunking] Bloque {i+1} falló ({chunk_err}) — omitido")
        finally:
            if os.path.exists(chunk_path):
                os.remove(chunk_path)

    if not chunk_results:
        raise RuntimeError("[Chunking] Ningún bloque fue procesado por Gemini")

    merged = merge_consolidated_results(chunk_results)
    print(f"[Chunking] {len(chunk_results)}/{n_chunks} bloques mergeados")
    return json.dumps(merged, ensure_ascii=False)


def _is_financial_document(resultado: dict) -> bool:
    fm = resultado.get("financial_metrics_2025")
    if not fm or not isinstance(fm, dict):
        return False
    _SENTINEL = {"", "null", "n/a", "--", "0", "none"}
    core_keys = [
        "revenue", "gross_profit", "ebitda", "cash", "burn", "mrr", "employees"
    ]
    for key in core_keys:
        entry = fm.get(key)
        if isinstance(entry, dict):
            val = str(entry.get("value") or "").strip().lower()
            if val and val not in _SENTINEL:
                return True
    return False


def _extract_kpi_confidence_scores(resultado: dict) -> dict[str, int]:
    fm = resultado.get("financial_metrics_2025", {})
    if not isinstance(fm, dict):
        return {}
    scores: dict[str, int] = {}
    for key, entry in fm.items():
        if isinstance(entry, dict):
            raw_conf = entry.get("confidence")
            if raw_conf is not None:
                try:
                    scores[key] = round(float(raw_conf) * 100)
                except (TypeError, ValueError):
                    pass
    return scores


def _apply_contract_normalization(
    contract: dict, raw_company: str, raw_period: str
) -> dict:
    """
    Normaliza period_id y company_id en el contrato en-lugar.
    Devuelve un resumen del resultado de la normalización.
    """
    from src.api import normalize_period  # noqa: PLC0415 — temporal hasta Phase 4

    errors: list[str] = []

    # Period normalization
    period_id, period_ok = normalize_period(raw_period)
    if period_ok:
        contract["submission"]["period_id"] = period_id
    else:
        errors.append(f"Período no reconocido: {raw_period!r} → fallback {period_id!r}")

    # Company normalization
    comp_id, fund_id, bucket_id, company_ok = get_company_id(raw_company)
    if company_ok:
        contract["submission"]["company_id"] = comp_id
        for row in contract.get("kpi_rows", []):
            row["company_id"] = comp_id

    return {
        "period_id":   period_id,
        "period_ok":   period_ok,
        "comp_id":     comp_id,
        "fund_id":     fund_id,
        "bucket_id":   bucket_id,
        "company_ok":  company_ok,
        "errors":      errors,
    }


# ═════════════════════════════════════════════════════════════════════════════
# ROUTE HANDLERS
# ═════════════════════════════════════════════════════════════════════════════

@router.post("/upload")
@limiter.limit("20/minute")
async def upload_pdf(
    request: Request,
    file: UploadFile = File(...),
    founder_email: str = Header(None, description="Email del founder"),
    company_id: str = Header(None, description="Company ID para multi-tenancy"),
    token: dict = Depends(require_auth),
):
    """
    Sube un archivo financiero y lo procesa via Gemini (PDF) o el Mapper (Excel/CSV).

    Flujo:
      1. Validaciones de seguridad (extensión, tamaño, magic bytes).
      2. Deduplicación por SHA-256 en la bóveda GCS.
      3a. Excel/CSV → motor de mapeo (109 KPIs) + preview + commitment gate.
      3b. PDF/DOCX  → Vertex AI / Gemini → build_contract → GCS vault.
    """
    print("=" * 60)
    print(f"[/upload] HIT — archivo={getattr(file, 'filename', 'N/A')}")
    print(f"          founder_email={founder_email!r}  company_id={company_id!r}")
    print("=" * 60)

    try:
        # ── C5: validar headers ───────────────────────────────────────────────
        founder_email = validate_email_header(founder_email)
        company_id    = validate_company_header(company_id)

        # ── 1. Validar extensión ──────────────────────────────────────────────
        ALLOWED_EXTENSIONS = {".pdf", ".csv", ".xlsx", ".xls", ".parquet", ".docx", ".doc"}
        file_ext = os.path.splitext(file.filename or "")[1].lower()
        if not file.filename or file_ext not in ALLOWED_EXTENSIONS:
            raise HTTPException(
                status_code=400,
                detail=f"Formato no soportado. Permitidos: {', '.join(sorted(ALLOWED_EXTENSIONS))}",
            )

        # ── 2. Leer + validar C2 (tamaño) + C7 (magic bytes) ─────────────────
        file_content = await file.read()
        if len(file_content) > MAX_FILE_BYTES:
            raise HTTPException(
                status_code=413,
                detail=f"Archivo supera el límite de {MAX_FILE_MB} MB",
            )
        if not validate_magic_bytes(file_content, file_ext):
            raise HTTPException(
                status_code=415,
                detail=f"El contenido binario no corresponde a un archivo {file_ext} válido",
            )

        # ── 3. Resolver empresa ───────────────────────────────────────────────
        _fname_key, _ = detect_company_from_text(file.filename or "")
        if _fname_key != "unknown":
            company_domain = _fname_key
        elif company_id:
            company_domain = company_id
        elif founder_email and "@" in founder_email:
            company_domain = founder_email.split("@")[-1]
        else:
            company_domain = "pending_detection"

        if not company_domain or company_domain == "unknown":
            company_domain = "pending_detection"

        # Guard: si no se pudo resolver la empresa, rechazar con 400 antes de
        # intentar cualquier operación de GCS o BQ que causaría un 500 opaco.
        if company_domain == "pending_detection":
            raise HTTPException(
                status_code=400,
                detail=(
                    "No se pudo identificar la empresa. "
                    "Selecciona tu empresa en el formulario antes de subir el archivo."
                ),
            )

        # ── 4. Hash y deduplicación en GCS ────────────────────────────────────
        file_hash = get_file_hash(file_content)
        print(f"[upload] hash={file_hash}  vault=vault/{company_domain}/")

        try:
            vault_prefix = f"vault/{company_domain}/"
            sc     = _get_storage_client()
            bucket = sc.bucket(_GCS_OUTPUT_BUCKET)

            for blob in bucket.list_blobs(prefix=vault_prefix):
                if blob.metadata and blob.metadata.get("file_hash") == file_hash:
                    content = blob.download_as_text()
                    result  = _ensure_dict(json.loads(content))
                    try:
                        _dup_contract = build_contract(
                            gemini_json=result, file_hash=file_hash,
                            company_id=company_domain,
                            founder_email=founder_email or "",
                            original_filename=file.filename,
                        )
                        _dup_bucket   = COMPANY_BUCKET.get(company_domain, "UNKNOWN")
                        dup_checklist = build_checklist_status(_dup_contract["kpi_rows"], _dup_bucket)
                    except Exception as _ce:
                        print(f"[upload] checklist recalc failed for dup ({_ce})")
                        dup_checklist = None
                    return JSONResponse(content={
                        "status": "success",
                        "message": "Documento reconocido en la bóveda de Cometa. Sincronizando métricas...",
                        "duplicate": True,
                        "result": result,
                        "file_hash": file_hash,
                        "company_domain": company_domain,
                        "checklist_status": dup_checklist,
                    })

            if check_hash_exists_in_gcs(_GCS_INPUT_BUCKET, file_hash):
                existing_result = get_existing_result(_GCS_OUTPUT_BUCKET, file_hash)
                if existing_result:
                    vault_blob = bucket.blob(f"{vault_prefix}{file_hash}_result.json")
                    vault_blob.metadata = {
                        "file_hash": file_hash,
                        "founder_email": founder_email,
                        "company_domain": company_domain,
                        "copied_from_general": True,
                        "processed_at": datetime.now(timezone.utc).isoformat(),
                    }
                    vault_blob.upload_from_string(
                        json.dumps(existing_result, indent=2),
                        content_type="application/json",
                    )
                    return JSONResponse(content={
                        "status": "success",
                        "message": "Documento reconocido en la bóveda de Cometa. Sincronizando métricas...",
                        "duplicate": True,
                        "result": existing_result,
                        "file_hash": file_hash,
                        "company_domain": company_domain,
                    })

        except RuntimeError as e:
            if str(e) == "GCS_AUTH":
                raise HTTPException(
                    status_code=500,
                    detail="Error de autenticación con GCS. Verifica GOOGLE_APPLICATION_CREDENTIALS.",
                )
            raise HTTPException(status_code=500, detail="Error de conexión con GCS durante deduplicación")

        # ── 5. Guardar archivo temporalmente ─────────────────────────────────
        safe_filename = sanitize_filename(file.filename)
        temp_path     = os.path.join("/tmp", f"{file_hash}_{safe_filename}")
        os.makedirs("/tmp", exist_ok=True)
        with open(temp_path, "wb") as tf:
            tf.write(file_content)

        # ── RAW GCS: archivar original antes de cualquier procesamiento ───────
        try:
            _raw_blob = sc.bucket(_GCS_INPUT_BUCKET).blob(
                f"{company_domain}/{file_hash}_{safe_filename}"
            )
            _raw_blob.upload_from_filename(temp_path, content_type="application/octet-stream")
        except Exception as _raw_err:
            print(f"[upload][raw] GCS raw upload non-fatal: {_raw_err}")

        # ── 6A. RAMA EXCEL/CSV — Motor de Mapeo (109 KPIs) ───────────────────
        if file_ext in (".xlsx", ".xls", ".csv"):
            try:
                _slug      = detect_company_from_text(company_domain)[0]
                _sector    = COMPANY_BUCKET.get(_slug, "ALL")
                _prev_cash = get_prev_cash_from_bq(_slug, datetime.now(timezone.utc).date())

                mapping_result = map_uploaded_file(
                    file_path=temp_path,
                    sector=_sector,
                    prev_cash=_prev_cash,
                )
                _period     = datetime.now(timezone.utc).date().replace(day=1)
                _period_iso = _period.isoformat()
                _meta       = IngestionMetadata(
                    company_name=company_domain.upper(),
                    company_slug=_slug,
                    period=_period,
                    founder_email=founder_email or "",
                    sector=_sector,
                    loaded_by=token.get("email", "unknown"),
                    source_file_hint=file.filename,
                )
                preview  = build_upload_preview(mapping_result, _meta)
                _load_id = preview.get("load_id", "")

                # Commitment gate
                _gate: dict = {}
                try:
                    _gate = evaluate_commitment_gate(
                        company_slug=_slug,
                        period_iso=_period_iso,
                        sector=_sector,
                        mapping_result=mapping_result,
                    )
                except Exception as _gate_err:
                    print(f"[upload][gate] non-fatal: {_gate_err}")

                # SQLite (acumulación multi-archivo)
                _rows_saved   = 0
                _audit_result = {}
                try:
                    _db_rows    = build_registry_rows(mapping_result, _meta, _load_id)
                    _rows_saved = save_registry_rows(_db_rows)
                    _audit_result = build_enriched_audit_response(
                        company_slug=_slug,
                        mapping_result=mapping_result,
                        rows_saved=_rows_saved,
                        load_id=_load_id,
                        period_iso=_period_iso,
                    )
                except Exception as _db_err:
                    print(f"[upload][local_db] non-fatal: {_db_err}")

                # Jero contract + BQ (solo si gate pasa)
                _jero_records = None
                _jero_path    = None
                if _gate.get("gate_passed"):
                    try:
                        _jero_records, _jero_path = generate_jero_contract(
                            mapping_result=mapping_result,
                            metadata=_meta,
                            load_id=_load_id,
                            gate_passed=True,
                        )
                        print(f"[upload][jero] {len(_jero_records)} registros generados")
                    except Exception as _jero_err:
                        print(f"[upload][jero] non-fatal: {_jero_err}")

                # Medallion layers
                upload_raw_layer(file_bytes=file_content, filename=file.filename or safe_filename, metadata=_meta)
                try:
                    upload_stage_layer(
                        gemini_json={
                            "load_id": _load_id, "company_slug": _slug,
                            "period": _period_iso,
                            "coverage_pct": getattr(mapping_result, "coverage_pct", None),
                        },
                        metadata=_meta, load_id=_load_id,
                    )
                except Exception as _st_err:
                    print(f"[upload][stage] non-fatal: {_st_err}")

                if _gate.get("gate_passed") and _jero_records:
                    try:
                        upload_gold_layer(
                            contract_json={
                                "load_id": _load_id, "company": _slug,
                                "period": _period_iso, "records": _jero_records, "gate": _gate,
                            },
                            metadata=_meta, load_id=_load_id,
                        )
                    except Exception as _gld_err:
                        print(f"[upload][gold] non-fatal: {_gld_err}")

                # Persistir pending para confirm-mapping
                _pending_key = f"pending_mapper/{_slug}/{_load_id}_{safe_filename}"
                try:
                    _pblob = _get_storage_client().bucket(_GCS_OUTPUT_BUCKET).blob(_pending_key)
                    _pblob.upload_from_filename(temp_path, content_type="application/octet-stream")
                    _pblob.metadata = {"company_slug": _slug, "sector": _sector,
                                       "founder_email": founder_email or ""}
                    _pblob.patch()
                except Exception as _gcs_err:
                    print(f"[upload][pending] GCS save non-fatal: {_gcs_err}")

                _status = "committed" if _gate.get("gate_passed") else "pending_kpis"
                return JSONResponse(content={
                    "status":          _status,
                    "flow":            "mapper",
                    "file_ext":        file_ext,
                    "commitment_gate": _gate,
                    "preview":         preview,
                    "audit":           _audit_result,
                    "jero_contract":   {
                        "generated":    _jero_records is not None,
                        "total_kpis":   len(_jero_records) if _jero_records else 0,
                        "output_file":  str(_jero_path.name) if _jero_path else None,
                    } if _jero_records else None,
                })

            except SubmissionBlockedError as sbe:
                raise HTTPException(status_code=422, detail=[
                    {"rule_id": f.rule_id, "severity": f.severity, "msg": f.message}
                    for f in sbe.blocking_flags
                ])
            except Exception as mapper_err:
                import traceback as _tb
                print(f"[upload][500] Error en motor de mapeo — stacktrace completo:\n{_tb.format_exc()}")
                raise HTTPException(status_code=500, detail=f"Error en motor de mapeo: {mapper_err}")

        # ── 6B. RAMA PDF/DOCX — Vertex AI / Gemini ───────────────────────────
        try:
            doc_ai = DocumentAIAdapter(_PROJECT_ID, _LOCATION_DOC_AI, _PROCESSOR_ID)
            gemini = GeminiAuditor(_PROJECT_ID, _VERTEX_LOCATION)

            _bucket_id = COMPANY_BUCKET.get(company_domain, "UNKNOWN")
            _sector_hints: dict[str, str] = {
                "SAAS":  "Prioriza MRR, Churn Rate y CAC.",
                "LEND":  "Prioriza Portfolio Size y NPL Ratio.",
                "ECOM":  "Prioriza GMV.",
                "INSUR": "Prioriza Loss Ratio.",
                "OTH":   "Extrae métricas financieras estándar.",
            }
            _prompt_prefix = (
                f"Eres un auditor financiero senior especializado en due diligence de startups.\n"
                f"Vertical del portfolio: {_bucket_id}.\n"
                f"{_sector_hints.get(_bucket_id, 'Extrae todas las métricas del esquema.')}\n\n"
            )
            _kpi_schema   = _build_gemini_kpi_schema()
            _prompt_body  = (
                "REGLAS OBLIGATORIAS:\n"
                "1. Responde ÚNICAMENTE con el objeto JSON.\n"
                "2. Usa EXACTAMENTE las claves del esquema.\n"
                "3. Si una métrica no aparece: escribe null.\n"
                "4. La clave raíz es SIEMPRE \"financial_metrics_2025\".\n\n"
                "ESQUEMA REQUERIDO:\n" + _kpi_schema + "\n\nAnaliza el documento y responde con el JSON completo."
            )
            prompt_config = _prompt_prefix + _prompt_body

            if file_ext in (".csv", ".xlsx", ".xls", ".parquet"):
                resultado_raw = _process_tabular(temp_path, file_ext, gemini, prompt_config)
            elif file_ext in (".docx", ".doc"):
                resultado_raw = _process_docx(temp_path, gemini, prompt_config)
            else:
                resultado_raw = _chunk_and_process_pdf(temp_path, gemini, prompt_config)

            if isinstance(resultado_raw, str):
                try:
                    resultado = json.loads(resultado_raw)
                except json.JSONDecodeError:
                    clean     = re.sub(r'^```json\s*|\s*```$', '', resultado_raw.strip())
                    resultado = json.loads(clean)
            else:
                resultado = resultado_raw

            # Resolución de empresa desde nombre de archivo y contenido
            _fname_key2, _ = detect_company_from_text(file.filename or "")
            _cont_key,   _ = detect_company_from_text(json.dumps(resultado))
            if _fname_key2 != "unknown":
                company_domain = _fname_key2
            elif _cont_key != "unknown":
                company_domain = _cont_key

            portfolio_id = lookup_portfolio(company_domain)
            contract     = build_contract(
                gemini_json=resultado,
                file_hash=file_hash,
                company_id=company_domain,
                founder_email=founder_email or "",
                original_filename=file.filename,
                portfolio_id=portfolio_id,
            )

            # Normalizar period_id y company_id en el contrato
            _raw_period  = resultado.get("_document_context", {}).get("period", "") or ""
            _norm_result = _apply_contract_normalization(
                contract=contract,
                raw_company=company_domain,
                raw_period=_raw_period,
            )
            if _norm_result["company_ok"]:
                company_domain = _norm_result["comp_id"].replace("COMP_", "").lower()

            company_bucket   = _norm_result["bucket_id"] or COMPANY_BUCKET.get(company_domain, "UNKNOWN")
            checklist_status = build_checklist_status(contract["kpi_rows"], company_bucket)
            _conf_scores     = _extract_kpi_confidence_scores(resultado)
            if _conf_scores:
                checklist_status["confidence_scores"] = _conf_scores

            # SQLite acumulación
            _pdf_slug       = company_domain.lower().replace(".", "_")
            _pdf_period_iso = datetime.now(timezone.utc).date().replace(day=1).isoformat()
            _pdf_load_id    = file_hash[:16]
            try:
                _pdf_db_rows = build_registry_rows_from_contract(
                    kpi_rows=contract["kpi_rows"],
                    company_slug=_pdf_slug,
                    company_name=company_domain,
                    period_iso=_pdf_period_iso,
                    load_id=_pdf_load_id,
                    loaded_by=token.get("email", "api"),
                )
                save_registry_rows(_pdf_db_rows)
            except Exception as _pdf_db_err:
                print(f"[upload][sqlite][pdf] non-fatal: {_pdf_db_err}")

            # Grilla acumulada 109 KPIs
            from src.core.local_db import _CATALOG_BY_METRIC_ID as _CAT_IDX  # noqa: PLC0415
            _current_kpis: dict = {}
            for _cr in contract["kpi_rows"]:
                if _cr.get("numeric_value") is not None and _cr.get("is_valid", False):
                    _cat = _CAT_IDX.get(_cr["kpi_key"])
                    if _cat:
                        _current_kpis[_cat["kpi_id"]] = {
                            "value": _cr["numeric_value"], "unit": _cat.get("unit"),
                            "match_type": "EXACT", "audit_status": "GOLD",
                            "quality_score": 1.0, "source": "gemini",
                        }
            _pdf_kpi_grid    = build_accumulated_kpi_grid(
                company_slug=_pdf_slug,
                period_iso=_pdf_period_iso,
                current_kpis=_current_kpis,
            )
            _pdf_found_count = sum(1 for r in _pdf_kpi_grid if r["status"] == "FOUND")

            needs_manual_mapping = not _is_financial_document(resultado)

            # BQ write: solo si el uploader es ANALISTA (Zero-Trust Gatekeeper)
            db_result: dict = {"inserted": False, "duplicate": False, "deferred": False}
            from src.core.db_writer import insert_contract  # noqa: PLC0415
            if token.get("role") == "ANALISTA":
                try:
                    db_result = insert_contract(contract)
                except Exception as db_err:
                    print(f"[upload][BQ] non-fatal: {db_err}")
            else:
                db_result["deferred"] = True

            # Guardar en GCS vault
            result_filename = f"vault/{company_domain}/{file_hash}_result.json"
            res_blob = _get_storage_client().bucket(_GCS_OUTPUT_BUCKET).blob(result_filename)
            res_blob.metadata = {
                "file_hash": file_hash,
                "original_filename": file.filename,
                "founder_email": founder_email,
                "company_domain": company_domain,
                "portfolio_id": portfolio_id,
                "processed_at": datetime.now(timezone.utc).isoformat(),
            }
            res_blob.upload_from_string(
                json.dumps(resultado, indent=2, ensure_ascii=False),
                content_type="application/json",
            )

            return JSONResponse(content={
                "status":               "success",
                "flow":                 "gemini",
                "file_hash":            file_hash,
                "company_domain":       company_domain,
                "result":               resultado,
                "checklist_status":     checklist_status,
                "db_result":            db_result,
                "needs_manual_mapping": needs_manual_mapping,
                "kpi_grid_summary": {
                    "total":   len(_pdf_kpi_grid),
                    "found":   _pdf_found_count,
                    "missing": len(_pdf_kpi_grid) - _pdf_found_count,
                    "source":  "gemini+sqlite",
                },
            })

        except json.JSONDecodeError as json_err:
            raise HTTPException(
                status_code=422,
                detail=f"Gemini devolvió JSON inválido: {json_err}",
            )
        except Exception as gemini_err:
            print(f"[upload][gemini] Error: {gemini_err}")
            raise HTTPException(status_code=500, detail=f"Error procesando con Gemini: {gemini_err}")

    finally:
        try:
            if "temp_path" in locals() and os.path.exists(temp_path):
                os.unlink(temp_path)
        except Exception:
            pass


@router.post("/api/founder/confirm-mapping")
@limiter.limit("10/minute")
async def founder_confirm_mapping(
    request: Request,
    body: ConfirmMappingRequest,
    token: dict = Depends(require_auth),
):
    """
    Confirma y persiste en BigQuery una carga Excel/CSV previamente analizada.

    1. Recupera el archivo temporal de GCS (pending_mapper/{load_id}_{filename}).
    2. Re-ejecuta el motor de mapeo.
    3. Llama dispatch_to_storage() con el load_id del preview para garantizar
       trazabilidad entre preview y escritura real.
    """
    import re as _re
    from datetime import date as _date

    if not _re.match(r"^\d{4}-\d{2}$", body.period_str):
        raise HTTPException(status_code=400, detail="period_str debe tener formato YYYY-MM")

    year_s, month_s = body.period_str.split("-")
    period = _date(int(year_s), int(month_s), 1)

    # Zero Trust: Founder solo puede confirmar su propia empresa
    if token.get("role") == "FOUNDER":
        _email   = token.get("email", "")
        _domain  = _email.split("@")[-1] if "@" in _email else ""
        _slug    = detect_company_from_text(_domain)[0]
        if _slug != "unknown" and _slug != body.company_slug:
            raise HTTPException(status_code=403,
                detail="Founders solo pueden confirmar cargas de su propia empresa.")

    # Recuperar archivo temporal de GCS
    _safe     = sanitize_filename(body.filename)
    _gcs_key  = f"pending_mapper/{body.company_slug}/{body.load_id}_{_safe}"
    _local_tmp = os.path.join("/tmp", f"confirm_{body.load_id}_{_safe}")
    try:
        _sc   = _get_storage_client()
        _blob = _sc.bucket(_GCS_OUTPUT_BUCKET).blob(_gcs_key)
        _blob.download_to_filename(_local_tmp)
    except Exception as dl_err:
        raise HTTPException(status_code=404,
            detail=f"Archivo de mapeo no encontrado en GCS ({_gcs_key}): {dl_err}")

    try:
        _prev_cash     = get_prev_cash_from_bq(body.company_slug, period)
        mapping_result = map_uploaded_file(
            file_path=_local_tmp,
            sector=body.sector,
            prev_cash=_prev_cash,
        )
        meta = IngestionMetadata(
            company_name=body.company_name.upper(),
            company_slug=body.company_slug.lower(),
            period=period,
            founder_email=token.get("email", ""),
            sector=body.sector,
            loaded_by=token.get("email", "pipeline@cometa.vc"),
        )
        dispatch = dispatch_to_storage(
            result=mapping_result,
            metadata=meta,
            source_type=body.source_type,
            load_id=body.load_id,
        )

        # Limpiar pending en GCS (best-effort)
        try:
            _blob.delete()
        except Exception:
            pass

        return JSONResponse(content={
            "status":          "committed",
            "load_id":         dispatch.load_id,
            "rows_inserted":   dispatch.rows_inserted,
            "rows_updated":    dispatch.rows_skipped_dup,
            "rows_error":      dispatch.rows_error,
            "quality_summary": dispatch.quality_summary,
            "warnings":        dispatch.warnings,
        })

    except SubmissionBlockedError as sbe:
        raise HTTPException(status_code=422, detail=[
            {"rule_id": f.rule_id, "severity": f.severity, "msg": f.message}
            for f in sbe.blocking_flags
        ])
    except Exception as confirm_err:
        raise HTTPException(status_code=500, detail=f"Error al confirmar la carga: {confirm_err}")
    finally:
        try:
            os.unlink(_local_tmp)
        except Exception:
            pass


@router.post("/api/founder/manual-update")
@limiter.limit("30/minute")
async def founder_manual_update(
    request: Request,
    body: ManualUpdateRequest,
    token: dict = Depends(require_auth),
) -> JSONResponse:
    """
    Persiste correcciones manuales del founder en el JSON de GCS.
    Localiza el blob por file_hash en la bóveda de la empresa del JWT.
    """
    company_domain: str = token.get("company_id") or token.get("sub", "")
    if "@" in company_domain:
        company_domain = company_domain.split("@")[-1]
    if not company_domain:
        raise HTTPException(status_code=403, detail="company_id no disponible en el token")

    vault_prefix = f"vault/{company_domain}/"
    try:
        sc         = _get_storage_client()
        bucket_obj = sc.bucket(_GCS_OUTPUT_BUCKET)
        blobs      = list(bucket_obj.list_blobs(prefix=vault_prefix))
    except Exception as gcs_err:
        raise HTTPException(status_code=500, detail=f"Error al conectar con GCS: {gcs_err}")

    target_blob = next(
        (b for b in blobs if b.name.endswith(".json") and body.file_hash in b.name),
        None,
    )
    if target_blob is None:
        raise HTTPException(
            status_code=404,
            detail=f"No se encontró resultado para hash '{body.file_hash}' en bóveda de {company_domain}",
        )

    try:
        existing_data = _ensure_dict(json.loads(target_blob.download_as_text()))
    except Exception as read_err:
        raise HTTPException(status_code=500, detail=f"Error al leer resultado existente: {read_err}")

    corrections: dict = existing_data.get("manual_corrections") or {}
    for k, v in body.updates.items():
        corrections[k] = v
    existing_data["manual_corrections"] = corrections

    try:
        target_blob.upload_from_string(
            json.dumps(existing_data, indent=2, ensure_ascii=False),
            content_type="application/json",
        )
    except Exception as write_err:
        raise HTTPException(status_code=500, detail=f"Error al guardar correcciones en GCS: {write_err}")

    updated_fields = list(body.updates.keys())
    print(f"[manual-update] {len(updated_fields)} campo(s) corregidos — hash={body.file_hash} company={company_domain}")
    return JSONResponse(content={"status": "ok", "updated_fields": updated_fields})


@router.get("/api/founder/config")
@limiter.limit("30/minute")
async def founder_config(
    request: Request,
    token: dict = Depends(require_auth),
) -> JSONResponse:
    """
    Auto-detecta company_id y vertical para el founder autenticado.
    Deriva el contexto de empresa desde el dominio del email en el JWT.
    """
    email: str = (token.get("email") or token.get("sub", "")).strip()
    if not email or "@" not in email:
        raise HTTPException(status_code=422, detail="email no disponible en el token")

    domain    = email.split("@", 1)[1].lower()
    comp_id, _, bucket_id, is_known = get_company_id(domain)
    vertical  = _BUCKET_TO_VERTICAL.get(bucket_id, "GENERAL")

    # ── Intentar resolver el company_id real desde BigQuery ───────────────────
    # get_company_id() genera IDs sintéticos "COMP_XXX" que no existen en BQ.
    # resolve_company_id() convierte "COMP_QUINIO" → real ID (ej. "C010").
    # Si BQ no está disponible, se cae al ID sintético sin romper la respuesta.
    try:
        comp_id = _bq_data.resolve_company_id(comp_id)
        is_known = True
    except Exception as _bq_err:
        _log.debug("[founder_config] BQ resolve falló (%s) — usando ID sintético '%s'", _bq_err, comp_id)

    # Display name
    _test_local_re  = re.compile(r"^(founder_test|test|demo|prueba|sandbox)\b", re.IGNORECASE)
    _test_domain_re = re.compile(r"^(test|demo|sandbox|example|localhost)\b", re.IGNORECASE)
    local_part  = email.split("@")[0]
    domain_base = domain.split(".")[0]
    _demo_slugs = {"demo-startup", "demostartup"}

    if (
        _test_local_re.match(local_part)
        or _test_domain_re.match(domain_base)
        or comp_id.replace("COMP_", "").lower().replace("_", "-") in _demo_slugs
    ):
        company_display_name = "Startup Demo"
    elif is_known:
        _pm_entry = PORTFOLIO_MAP.get(domain_base, PORTFOLIO_MAP.get(
            domain_base.replace("-", "").replace("_", ""), {}
        ))
        company_display_name = (
            _pm_entry.get("display_name")
            or domain_base.replace("-", " ").replace("_", " ").title()
        )
    else:
        company_display_name = domain_base.replace("-", " ").replace("_", " ").title() or "tu empresa"

    return JSONResponse(content={
        "company_id":           comp_id,
        "vertical":             vertical,
        "is_known":             is_known,
        "domain":               domain,
        "company_display_name": company_display_name,
    })


@router.post("/api/founder/finalize")
@limiter.limit("10/minute")
async def founder_finalize(
    request: Request,
    body: FinalizeRequest,
    token: dict = Depends(require_auth),
) -> JSONResponse:
    """
    Finaliza el expediente del founder.

    Regla de hierro: Gate 109/109 KPIs.  El backend verifica que el ADN
    financiero esté completo antes de emitir el Vault Seal y el email.

    1. Evalúa commitment gate (presente/total KPIs).
    2. Persiste rescue notes en BQ Truth Shield.
    3. Genera Vault Seal (SHA-256 de integridad).
    4. Registra en BigQuery upload_log.
    5. Envía correo de confirmación HTML al founder.
    """
    from src.services.email_service import send_receipt_email  # noqa: PLC0415
    from src.services.hash_service  import generate_vault_seal  # noqa: PLC0415

    founder_email: str = (token.get("email") or token.get("sub", "")).strip()
    if not founder_email:
        raise HTTPException(status_code=403, detail="email no disponible en el token")

    # Slug de empresa: siempre del JWT (Zero Trust — body.company_domain ignorado)
    company_domain = (token.get("company_slug") or "").strip()
    if not company_domain:
        company_id_claim = (token.get("company_id") or "").strip()
        if "@" in company_id_claim:
            company_domain = company_id_claim.split("@")[-1]
        elif company_id_claim:
            company_domain = company_id_claim
        else:
            company_domain = founder_email.split("@")[-1] if "@" in founder_email else "cometa"

    file_names   = body.file_names or [f[:16] + "…" for f in body.file_hashes]
    raw_manual   = body.manual_kpis or {}
    processed_at = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M")

    # ── Gate de seguridad final: 109/109 o nada ───────────────────────────────
    _raw_manual_keys = {k for k in raw_manual if not k.endswith("_note")}
    _gate = check_gate_for_finalize(
        company_slug=company_domain.replace(".", "_").lower(),
        manual_kpi_refs=_raw_manual_keys,
    )
    if not _gate["gate_passed"]:
        raise HTTPException(
            status_code=403,
            detail=(
                "Sincronización abortada: El ADN financiero está incompleto. "
                f"({_gate['present']}/{_gate['total']} KPIs presentes — faltan {_gate['missing']})"
            ),
        )
    print(f"[finalize] GATE OK — {_gate['present']}/{_gate['total']} company={company_domain!r}")

    # ── Separar valores de notas de rescate ───────────────────────────────────
    manual_kpis:  dict[str, str] = {}
    rescue_notes: dict[str, str] = {}
    for k, v in raw_manual.items():
        if k.endswith("_note"):
            rescue_notes[k[:-5]] = v
        else:
            manual_kpis[k] = v

    # ── Persist rescue notes (Truth Shield) ──────────────────────────────────
    if rescue_notes:
        for fh in body.file_hashes:
            try:
                result = write_founder_rescue_notes(fh, rescue_notes)
                print(f"[finalize] rescue notes — hash={fh[:16]}… updated={result['updated']}")
            except Exception as _note_err:
                print(f"[finalize] note write non-fatal: {_note_err}")

    period_id = datetime.now(timezone.utc).strftime("%Y")

    # ── Vault Seal ────────────────────────────────────────────────────────────
    vault_seal = generate_vault_seal(
        company_id=company_domain,
        file_hash=body.file_hashes[0] if body.file_hashes else "",
        kpi_rows=[
            {"kpi_key": k, "raw_value": v, "unit": "", "is_valid": True}
            for k, v in manual_kpis.items()
        ],
        processed_at=processed_at,
    )
    print(f"[finalize] company={company_domain!r} files={len(body.file_hashes)} seal={vault_seal[:16]}…")

    # ── BigQuery upload_log ───────────────────────────────────────────────────
    try:
        insert_upload_log(
            company_id=company_domain,
            founder_email=founder_email,
            vault_seal=vault_seal,
            file_hashes=body.file_hashes,
            manual_kpis=manual_kpis if manual_kpis else None,
            period_id=period_id,
            manual_edits_count=len(manual_kpis),
        )
    except Exception as _log_err:
        print(f"[finalize] upload_log non-fatal: {_log_err}")

    # ── Correo de confirmación ────────────────────────────────────────────────
    send_receipt_email(
        to_email=founder_email,
        company_domain=company_domain,
        period=period_id,
        vault_seal=vault_seal,
        file_hash=body.file_hashes[0] if body.file_hashes else "",
        kpi_count=len(manual_kpis),
        processed_at=processed_at,
    )

    return JSONResponse(content={
        "status":     "ok",
        "message":    "Expediente registrado. Se ha enviado tu Recibo Digital al correo.",
        "sent_to":    founder_email,
        "vault_seal": vault_seal,
    })


# ══════════════════════════════════════════════════════════════════════════════
# POST /api/founder/ingest-kpis — Ingesta directa a BD_Cometa_Dev
# ══════════════════════════════════════════════════════════════════════════════
#
# Flujo E2E:
#   1. JWT válido → require_auth (sin restricción de dominio: founders pueden usar esto)
#   2. BQService.get_company_context() → valida empresa en dim_company vía BQ
#   3. Genera submission_id único (uuid)
#   4. BQService.save_submission() → escribe submissions + fact_kpi_values
#   5. Retorna JSON de confirmación con submission_id
#
# Este endpoint NO duplica la lógica del /finalize legacy (que trabaja con
# file_hashes, GCS y email). Es el punto de entrada nativo para ingesta
# directa de KPIs procesados desde el pipeline de datos.
# ══════════════════════════════════════════════════════════════════════════════

import uuid as _uuid
from datetime import date as _date

from src.core.bq_service import BQService as _BQService, CompanyNotFoundError as _CompanyNotFoundError

# Singleton: una sola instancia por proceso (el cliente BQ es thread-safe)
_bq = _BQService()


class KpiItem(BaseModel):
    """Un KPI individual dentro del payload de ingesta."""
    metric_id:   str
    value:       float | None = None
    value_notes: str   | None = None


class IngestKpisRequest(BaseModel):
    """
    Payload de POST /api/founder/ingest-kpis.

    Todos los IDs deben seguir el esquema BD_Cometa_Dev:
      company_id  → C001
      fund_id     → F001
      period_id   → P2026Q1M01
      metric_id   → K001
    """
    company_id:   str
    fund_id:      str
    period_id:    str
    period_start: _date              # ej. "2026-01-01" — clave de partición en BQ
    kpis:         list[KpiItem]

    @property
    def kpis_as_dicts(self) -> list[dict]:
        return [kpi.model_dump() for kpi in self.kpis]


@router.post("/api/founder/ingest-kpis", status_code=201)
async def ingest_kpis(
    body:  IngestKpisRequest,
    token: dict = Depends(require_auth),
) -> JSONResponse:
    """
    Ingesta directa de KPIs procesados a BD_Cometa_Dev.

    - Valida la empresa en dim_company (JOIN con dim_bucket) — cero hardcode.
    - Genera un submission_id único.
    - Escribe 1 fila en submissions + N filas en fact_kpi_values.
    - Retorna el submission_id para trazabilidad end-to-end.

    Errores:
      404 → empresa no encontrada en dim_company
      422 → body inválido (Pydantic)
      500 → error de escritura en BigQuery
    """
    submitted_by: str = token.get("email") or token.get("sub", "unknown")

    # ── Paso 1: validar empresa en BQ (cero diccionarios hardcodeados) ────────
    try:
        context = _bq.get_company_context(body.company_id)
    except _CompanyNotFoundError as exc:
        raise HTTPException(status_code=404, detail=str(exc))
    except RuntimeError as exc:
        raise HTTPException(status_code=500, detail=str(exc))

    # ── Paso 2: generar submission_id único ───────────────────────────────────
    submission_id = f"S-{_uuid.uuid4().hex[:8].upper()}"

    # ── Paso 3: persistir en BigQuery ─────────────────────────────────────────
    try:
        rows_inserted = _bq.save_submission(
            submission_id=submission_id,
            company_id=body.company_id,
            period_id=body.period_id,
            fund_id=body.fund_id,
            period_start=body.period_start,
            submitted_by=submitted_by,
            kpis=body.kpis_as_dicts,
        )
    except RuntimeError as exc:
        raise HTTPException(status_code=500, detail=str(exc))

    return JSONResponse(
        status_code=201,
        content={
            "status":        "created",
            "submission_id": submission_id,
            "company_id":    context["company_id"],
            "company_name":  context["company_name"],
            "bucket_name":   context["bucket_name"],
            "period_id":     body.period_id,
            "kpis_inserted": rows_inserted,
            "submitted_by":  submitted_by,
        },
    )


# ══════════════════════════════════════════════════════════════════════════════
# POST /api/founder/process-document — Patrón Adaptador Unificado
# ══════════════════════════════════════════════════════════════════════════════
#
# Punto de entrada único para PDF y Excel/CSV.
#
# Flujo:
#   1. Validación de tamaño + magic bytes.
#   2. Ruteo por tipo de archivo:
#        PDF   → data_contract.extract_pdf_to_contract()
#        Excel → kpi_mapper.extract_excel_to_contract()
#   3. Ambos devuelven UnifiedKPIContract (validado por Pydantic).
#   4. bq_data_service.insert_submission_and_facts() escribe en BD_Cometa_Dev.
#
# Este endpoint reemplaza el /upload legacy para ingesta nueva.
# El /upload se mantiene activo hasta migración completa del frontend.
# ══════════════════════════════════════════════════════════════════════════════

import logging as _logging
import tempfile as _tempfile
from datetime import date as _date_type

from src.core.bq_data_service import (
    BQDataService as _BQDataService,
    BQInsertError as _BQInsertError,
    CompanyNotFoundError as _CompanyNotFoundError2,
)
from src.core.data_contract import extract_pdf_to_contract as _extract_pdf
from src.core.kpi_mapper import (
    detect_company_and_year_from_df as _detect_company_year,
    detect_company_from_excel as _detect_company_excel,
    extract_excel_to_contract as _extract_excel,
    extract_long_format_to_staging_rows as _extract_long_format,
    extract_master_db_to_staging_rows as _extract_master_db,
)
from src.core.vc_validator import validate_financial_physics as _validate_physics
from src.schemas import ProcessDocumentResponse, UnifiedKPIContract  # noqa: F401

_log = _logging.getLogger(__name__)

_bq_data = _BQDataService()


def _period_id_to_date(period_id: str) -> _date_type:
    """
    Convierte un period_id canónico a la fecha de inicio del período.

    P2026Q1M01 → date(2026, 1, 1)
    FY2025     → date(2025, 1, 1)
    H12025     → date(2025, 1, 1)
    """
    m = re.match(r"P(20\d{2})Q[1-4]M(\d{2})", period_id)
    if m:
        return _date_type(int(m.group(1)), int(m.group(2)), 1)
    m2 = re.match(r"FY(20\d{2})", period_id)
    if m2:
        return _date_type(int(m2.group(1)), 1, 1)
    m3 = re.match(r"H[12](20\d{2})", period_id)
    if m3:
        return _date_type(int(m3.group(1)), 1, 1)
    return _date_type.today()


def _extract_company_id_from_file(content: bytes, suffix: str) -> str:
    """
    Bypass de emergencia: intenta extraer el nombre de empresa directamente
    del contenido del archivo cuando company_id llega vacío desde el frontend.

    Escanea las primeras 10 filas × 5 columnas del Excel buscando celdas de
    texto que no sean números ni fechas. El primer candidato razonable (≥3 chars,
    no numérico) se devuelve como identificador para que resolve_company_id()
    lo resuelva en BQ.

    Returns "" si no se puede extraer nada útil.
    """
    if suffix not in (".xlsx", ".xls", ".csv"):
        return ""
    import io as _io
    import tempfile as _tf
    try:
        import pandas as _pd
        with _tf.NamedTemporaryFile(suffix=suffix, delete=False) as _t:
            _t.write(content)
            _tmp = _t.name
        try:
            if suffix in (".xlsx", ".xls"):
                df = _pd.read_excel(_tmp, header=None, nrows=10, dtype=object)
            else:
                df = _pd.read_csv(_tmp, header=None, nrows=10, dtype=str)
        finally:
            import os as _os
            try:
                _os.unlink(_tmp)
            except OSError:
                pass

        # Scan first 10 rows × first 5 cols for text that looks like a company name
        for r in range(min(10, df.shape[0])):
            for c in range(min(5, df.shape[1])):
                cell = df.iloc[r, c]
                if _pd.isna(cell):
                    continue
                s = str(cell).strip()
                # Skip obviously numeric, date, or too-short values
                if not s or len(s) < 3 or s.replace(".", "").replace(",", "").isnumeric():
                    continue
                # Skip cells that look like period headers (Jan 2025, 2025-01, etc.)
                import re as _re
                if _re.search(r"\b(jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec|20\d{2})\b", s, _re.I):
                    continue
                _log.info("[_extract_company_id_from_file] Candidato empresa: '%s'", s)
                return s
    except Exception as exc:
        _log.warning("[_extract_company_id_from_file] Error: %s", exc)
    return ""


_EXCEL_MIME_TYPES: frozenset[str] = frozenset({
    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",  # xlsx
    "application/vnd.ms-excel",                                           # xls
    "text/csv",
    "application/csv",
    "application/octet-stream",  # algunos navegadores envían xlsx como binario genérico
})


@router.get("/api/founder/staging")
@limiter.limit("30/minute")
async def get_founder_staging(
    request:    Request,
    token:      dict = Depends(require_auth),
) -> JSONResponse:
    """
    Devuelve los KPIs en fact_kpi_staging con status='PENDING' para la empresa
    del founder autenticado. Permite al founder ver sus datos recién cargados
    antes de que el analista los mueva a fact_kpi_values.

    La empresa se obtiene siempre del JWT (Zero Trust — nunca del query param).
    """
    raw_company_id: str = (
        token.get("company_id", "")
        or token.get("company_slug", "")
        or ""
    ).strip()

    if not raw_company_id:
        raise HTTPException(status_code=403, detail="company_id no disponible en el token.")

    try:
        company_id = _bq_data.resolve_company_id(raw_company_id)
    except _CompanyNotFoundError2:
        company_id = raw_company_id   # best-effort: use raw value for the query
    except _BQInsertError as exc:
        raise HTTPException(status_code=502, detail=str(exc)) from exc

    try:
        rows = _bq_data.get_staging_pending(company_id=company_id, limit=500)
    except _BQInsertError as exc:
        raise HTTPException(status_code=502, detail=str(exc)) from exc

    # Group by staging_id for easier frontend rendering
    by_batch: dict[str, dict] = {}
    for r in rows:
        sid = r.get("staging_id", "unknown")
        if sid not in by_batch:
            by_batch[sid] = {
                "staging_id":   sid,
                "company_id":   r.get("company_id", ""),
                "submitted_at": str(r.get("submitted_at", "")),
                "source_file":  r.get("source_file", ""),
                "status":       r.get("status", "PENDING"),
                "rows":         [],
            }
        by_batch[sid]["rows"].append({
            "metric_id": r.get("metric_id", ""),
            "value":     r.get("value"),
            "period_id": r.get("period_id", ""),
        })

    return JSONResponse(content={
        "company_id": company_id,
        "total_rows": len(rows),
        "batches":    list(by_batch.values()),
    })


@router.post("/api/founder/process-document", response_model=ProcessDocumentResponse)
@limiter.limit("10/minute")
async def process_document(
    request:    Request,
    file:       UploadFile    = File(...),
    period_id:  str | None   = Form(None),   # opcional — Excel lo detecta del archivo
    company_id: str | None   = Form(None),   # opcional — Excel lo detecta del archivo
    year:       str | None   = Form(None),   # opcional — fallback si scanner falla
    bucket_id:  str | None   = Form(None),   # opcional — no requerido por el scanner
    token:      dict         = Depends(require_auth),
) -> JSONResponse:
    """
    Endpoint unificado de ingesta (Patrón Adaptador).

    Acepta PDF, XLSX, XLS o CSV.

    Flujo Excel Master Database (métricas en filas, meses en columnas):
      - Detecta automáticamente los headers de período (Jan 2025, etc.)
      - Hace unpivot completo → fact_kpi_staging con status PENDING
      - period_id no es necesario (se deriva de los headers del archivo)

    Flujo Excel período único / PDF:
      - period_id es obligatorio
      - Gemini extrae las métricas del período indicado

    Errores:
      413 → archivo demasiado grande
      415 → formato no soportado o magic bytes inválidos
      404 → empresa no encontrada en dim_company
      422 → extracción fallida, contrato inválido o period_id faltante
      500 → error interno / BQ
    """
    # ── 1. Leer y validar archivo ─────────────────────────────────────────────
    content = await file.read()
    if len(content) > MAX_FILE_BYTES:
        raise HTTPException(
            status_code=413,
            detail=f"El archivo excede el límite de {MAX_FILE_MB} MB.",
        )

    filename = sanitize_filename(file.filename or "document")
    suffix   = Path(filename).suffix.lower()

    if not validate_magic_bytes(content, suffix):
        raise HTTPException(
            status_code=415,
            detail=f"El contenido binario no corresponde a un archivo {suffix} válido.",
        )

    content_type_lower = (file.content_type or "").lower()
    is_pdf   = suffix == ".pdf" or content_type_lower == "application/pdf"
    is_excel = suffix in (".xlsx", ".xls", ".csv") or content_type_lower in _EXCEL_MIME_TYPES

    if not is_pdf and not is_excel:
        raise HTTPException(
            status_code=415,
            detail=(
                "Formato no soportado. Se aceptan PDF, XLSX, XLS o CSV. "
                f"Recibido: sufijo='{suffix}' content-type='{content_type_lower}'"
            ),
        )

    submitted_by   = token.get("sub") or token.get("email", "unknown")
    source_file    = f"upload://{filename}"
    company_id     = (company_id or "").strip()   # normalizar None → ""
    # Audit vars — siempre definidos; se sobreescriben en los bloques siguientes
    _raw_file_path = ""
    _preview_url   = ""
    _company_name  = company_id   # se actualiza en 1C con el nombre canónico de BQ
    _detected_year = (year or "").strip() or "auto"

    # ── 1B. SCANNER DE IDENTIDAD — fuente de verdad para Excel ────────────────
    #
    # Para archivos Excel el scanner lee el contenido del archivo y SOBRESCRIBE
    # cualquier parámetro que haya enviado el frontend (company_id, etc.).
    # Razón: el Founder no debe elegir nada — el archivo manda.
    #
    # Para PDF el flujo es distinto: Gemini extrae el período y company_id sigue
    # siendo necesario desde el frontend (o el JWT).
    _excel_scan: dict = {}   # resultado del scanner; vacío si no es Excel

    if is_excel:
        # Leer bytes directamente en un DataFrame — sin tempfile, sin I/O extra
        try:
            import io as _io
            import pandas as _pd_scan
            _df_scan = _pd_scan.read_excel(
                _io.BytesIO(content), header=None, nrows=20, dtype=object
            )
            # Catálogo live de dim_company (cacheado 5 min) — cubre todas las empresas
            _catalog = _bq_data.get_company_catalog_for_scanner()
            _excel_scan = _detect_company_year(_df_scan, company_catalog=_catalog)
        except Exception as _scan_exc:
            _log.warning("[process-document] Scanner de Excel no-fatal: %s", _scan_exc)
            _excel_scan = {}

        # ── Log inmediato del scanner (aparece SIEMPRE, fallo o no) ──────────
        _scanned_id   = _excel_scan.get("company_id")
        _scanned_name = _excel_scan.get("company_name", _scanned_id or "UNKNOWN")
        _scanned_year = _excel_scan.get("year", "no detectado")
        _scanned_bkt  = _excel_scan.get("bucket_id", "?")
        print(
            f"[AUTO-DETECT] Empresa: {_scanned_name} (ID: {_scanned_id}), "
            f"Año: {_scanned_year}, Bucket: {_scanned_bkt}",
            flush=True,
        )
        _log.info(
            "[AUTO-DETECT] Empresa: %s (ID: %s), Año: %s, Bucket: %s",
            _scanned_name, _scanned_id, _scanned_year, _scanned_bkt,
        )

        # ── BLOQUEO si no se identificó empresa ──────────────────────────────
        if not _scanned_id:
            raise HTTPException(
                status_code=422,
                detail=(
                    "No se pudo identificar la empresa en el archivo. "
                    "No se permite la carga manual para este proceso. "
                    "Asegúrate de que el Excel incluya el nombre de la empresa "
                    "en las primeras filas."
                ),
            )

        # ── SOBRESCRITURA FORZOSA — el archivo manda, siempre ─────────────────
        company_id = str(_scanned_id)   # str() explícito: nunca bytes

    # ── 1D. Upload raw a GCS — capa de auditoría ──────────────────────────────
    #
    # Sube el archivo original ANTES de procesarlo, para que el analista pueda
    # abrirlo al lado de los KPIs extraídos.  No-fatal: si GCS falla, el flujo
    # de ingesta continúa sin interrumpirse.
    _raw_ts   = datetime.now(timezone.utc).strftime("%Y%m%d%H%M%S%f")[:18]
    _raw_id   = f"RAW{_raw_ts}"
    _safe_fn  = re.sub(r"[^\w.\-]", "_", filename)
    _det_year = _excel_scan.get("year") or (year or "").strip() or "unknown"
    _gcs_folder = (
        f"{company_id}/{_det_year}" if company_id and _det_year != "unknown"
        else "unidentified"
    )
    _blob_path = f"{_gcs_folder}/{_raw_id}_{_safe_fn}"
    try:
        from datetime import timedelta as _timedelta
        _gcs = _get_storage_client()
        _bkt = _gcs.bucket(_GCS_INPUT_BUCKET)
        _blb = _bkt.blob(_blob_path)
        _blb.upload_from_string(content, content_type=file.content_type or "application/octet-stream")
        _raw_file_path = f"gs://{_GCS_INPUT_BUCKET}/{_blob_path}"
        try:
            _preview_url = _blb.generate_signed_url(
                version="v4",
                expiration=_timedelta(minutes=15),
                method="GET",
            )
        except Exception as _sign_exc:
            _log.warning("[process-document] Signed URL no-fatal: %s", _sign_exc)
            _preview_url = _raw_file_path
        source_file = _raw_file_path   # actualizar referencia para BQ y audit
        _log.info("[GCS] Raw subido: %s", _raw_file_path)
    except Exception as _gcs_exc:
        _log.warning("[process-document] GCS upload no-fatal: %s", _gcs_exc)

    # ── 1C. Resolver company_id canónico en BQ ────────────────────────────────
    # Si el scanner ya devolvió un ID canónico (ej. "C010"), resolve_company_id
    # lo confirmará en BQ.  Si el ID del frontend era basura, ya fue descartado.
    try:
        company_id = _bq_data.resolve_company_id(company_id)
    except _CompanyNotFoundError2 as exc:
        _log.warning("[process-document] resolve_company_id falló: %s", exc)
        raise HTTPException(status_code=404, detail=str(exc)) from exc
    except _BQInsertError as exc:
        _log.error("[process-document] BQ error en resolve_company_id: %s", exc)
        raise HTTPException(status_code=502, detail=str(exc)) from exc

    # Obtener nombre de empresa y sector directamente de BQ (fuente de verdad)
    try:
        _company_meta  = _bq_data.get_company_metadata(company_id)
        _company_name  = _company_meta.get("company_name", company_id)
        _bucket_name   = _company_meta.get("bucket_name") or _excel_scan.get("bucket_id", "UNKNOWN")
    except Exception:
        _company_name = company_id
        _bucket_name  = _excel_scan.get("bucket_id", "UNKNOWN")

    _detected_year = _excel_scan.get("year") or _detected_year   # scanner gana; fallback al Form
    _automation_msg = (
        f"[AUTOMATION] Usando Empresa: {_company_name}, "
        f"Año: {_detected_year}, Sector: {_bucket_name}"
    )
    print(_automation_msg, flush=True)
    _log.info(_automation_msg)

    # ── 2A. RAMA EXCEL MASTER DATABASE (métricas × meses → multi-período) ─────
    if is_excel:
        with _tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as _tmp:
            _tmp.write(content)
            _tmp_path = _tmp.name

        try:
            # Intento 1: formato long/tidy (metric_name, period, value como columnas)
            staging_rows = _extract_long_format(file_path=_tmp_path)
            # Intento 2: formato wide/pivot (métricas en filas, meses en columnas)
            if not staging_rows:
                staging_rows = _extract_master_db(file_path=_tmp_path)
        except Exception as _exc:
            _log.warning("[process-document] Extracción multi-período falló: %s", _exc)
            staging_rows = []
        finally:
            try:
                os.unlink(_tmp_path)
            except OSError:
                pass

        if staging_rows:
            # ── Multi-período detectado: va directo a fact_kpi_staging ─────
            # ID simple basado en timestamp — sin uuid, sin bytes, sin pyarrow
            _ts = datetime.now(timezone.utc).strftime("%Y%m%d%H%M%S%f")[:18]
            staging_id = f"STG{_ts}"
            _staging_ok = False
            try:
                _bq_data.get_company_metadata(company_id)   # valida que la empresa existe
                stg_result = _bq_data.insert_to_staging_multiperiod(
                    staging_id=staging_id,
                    company_id=company_id,
                    submitted_by=submitted_by,
                    source_file=source_file,
                    staging_rows=staging_rows,
                )
                _staging_ok = True
            except _CompanyNotFoundError2 as exc:
                raise HTTPException(status_code=404, detail=str(exc)) from exc
            except Exception as exc:
                import traceback as _tb
                _log.error(
                    "[process-document] Multi-período falló (%s) — fallback a Gemini:\n%s",
                    exc, _tb.format_exc(),
                )
                staging_rows = []   # vaciar para que caiga al flujo Gemini

            if _staging_ok:
                periods = stg_result.get("periods", [])
                _log.info(
                    "[process-document] MasterDB OK — staging=%s company=%s "
                    "metrics=%d rows=%d períodos=%s",
                    staging_id, company_id, len(staging_rows),
                    stg_result["rows_inserted"], periods,
                )
                return JSONResponse(
                    status_code=200,
                    content=ProcessDocumentResponse(
                        submission_id=staging_id,
                        rows_inserted=stg_result["rows_inserted"],
                        timestamp=stg_result["timestamp"],
                        period_id=periods[0] if periods else "MULTI",
                        company_id=company_id,
                        metrics_count=len(staging_rows),
                        cerebro={},
                        periods=periods,
                        audit={
                            "company_name":  _company_name,
                            "year":          _detected_year,
                            "raw_file_path": _raw_file_path,
                            "preview_url":   _preview_url,
                        },
                    ).model_dump(),
                )

        # ── Sin headers de mes: fallback a extractor estándar (período único) ─
        if not period_id:
            raise HTTPException(
                status_code=422,
                detail=(
                    "No se detectaron columnas de período en el Excel. "
                    "Para archivos de período único, incluye 'period_id' en el form "
                    "(ej. 'P2025Q1M01')."
                ),
            )

        with _tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as _tmp2:
            _tmp2.write(content)
            _tmp2_path = _tmp2.name

        try:
            contract: UnifiedKPIContract = _extract_excel(
                file_path=_tmp2_path, period_id=period_id
            )
        except ValueError as exc:
            raise HTTPException(status_code=422, detail=str(exc)) from exc
        except Exception as exc:
            _log.exception("[process-document] Excel extracción fallida: %s", exc)
            raise HTTPException(status_code=500, detail=f"Error en extracción: {exc}") from exc
        finally:
            try:
                os.unlink(_tmp2_path)
            except OSError:
                pass

    # ── 2B. RAMA PDF ──────────────────────────────────────────────────────────
    else:
        if not period_id:
            raise HTTPException(
                status_code=422,
                detail="period_id es obligatorio para archivos PDF.",
            )
        try:
            contract = _extract_pdf(pdf_bytes=content, period_id=period_id)
        except ValueError as exc:
            raise HTTPException(status_code=422, detail=str(exc)) from exc
        except Exception as exc:
            _log.exception("[process-document] PDF extracción fallida: %s", exc)
            raise HTTPException(status_code=500, detail=f"Error en extracción: {exc}") from exc

    # ── 3. Resolver metadatos de empresa desde BQ ─────────────────────────────
    try:
        company_meta = _bq_data.get_company_metadata(company_id)
    except _CompanyNotFoundError2 as exc:
        raise HTTPException(status_code=404, detail=str(exc)) from exc
    except _BQInsertError as exc:
        raise HTTPException(status_code=502, detail=str(exc)) from exc

    fund_id      = company_meta.get("fund_id", "unknown")
    period_start = _period_id_to_date(period_id)  # type: ignore[arg-type]  # period_id validated above
    kpi_rows = [
        {
            "metric_id":   m.metric_id,
            "value":       m.value,
            "value_notes": f"source={m.source} | period={m.period_id}",
        }
        for m in contract.metrics
    ]

    # ── 4. Cerebro — física financiera antes de escribir en BQ ───────────────
    try:
        cerebro_result = _validate_physics(contract)
        _log.info(
            "[process-document] Cerebro OK — violations=%d missing=%d",
            len(cerebro_result.get("violations", [])),
            len(cerebro_result.get("missing_required", [])),
        )
    except Exception as _cerebro_exc:
        _log.warning("[process-document] Cerebro non-fatal: %s", _cerebro_exc)
        cerebro_result = {}

    # ── 5. Persistir en BD_Cometa_Dev ─────────────────────────────────────────
    try:
        bq_result = _bq_data.insert_submission_and_facts(
            company_id=company_id,
            fund_id=fund_id,
            period_id=period_id,
            period_start=period_start,
            submitted_by=submitted_by,
            source_file=source_file,
            kpi_rows=kpi_rows,
            raw_file_path=_raw_file_path or None,
        )
    except _BQInsertError as exc:
        _log.error("[process-document] BQ insert fallido: %s", exc)
        raise HTTPException(status_code=500, detail=str(exc)) from exc

    _log.info(
        "[process-document] OK — submission=%s company=%s metrics=%d bq_rows=%d gcs=%s",
        bq_result["submission_id"], company_id, len(contract.metrics),
        bq_result["rows_inserted"], _raw_file_path or "—",
    )

    return JSONResponse(
        status_code=200,
        content=ProcessDocumentResponse(
            submission_id=bq_result["submission_id"],
            rows_inserted=bq_result["rows_inserted"],
            timestamp=bq_result["timestamp"],
            period_id=period_id,
            company_id=company_id,
            metrics_count=len(contract.metrics),
            cerebro=cerebro_result,
            periods=[period_id],
            audit={
                "company_name":  _company_name,
                "year":          _detected_year,
                "raw_file_path": _raw_file_path,
                "preview_url":   _preview_url,
            },
        ).model_dump(),
    )
