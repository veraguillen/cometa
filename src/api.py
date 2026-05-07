import sys, io
if sys.stdout.encoding != "utf-8":
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")
if sys.stderr.encoding != "utf-8":
    sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding="utf-8", errors="replace")

from dotenv import load_dotenv
load_dotenv()  # Carga .env desde el directorio de trabajo

from fastapi import FastAPI, UploadFile, File, Header, HTTPException, Depends, Request, Query
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse, StreamingResponse
from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials
from fastapi.staticfiles import StaticFiles
from fastapi.exceptions import RequestValidationError
from slowapi import Limiter, _rate_limit_exceeded_handler
from slowapi.util import get_remote_address
from slowapi.errors import RateLimitExceeded
from pydantic import BaseModel, ValidationError, model_validator
from src.schemas import UserSchema, UserOut, PortfolioMetadataResponse  # UserOut = alias de UserSchema
import bcrypt
from jose import jwt, JWTError
import os
import re
import unicodedata
import hashlib
import json
import traceback
from datetime import datetime, timezone, timedelta
import secrets
from pathlib import Path
from src.auth_utils import (
    create_access_token,
    enforce_internal_role,
    generate_hybrid_id,
    is_hybrid_id,
    JWT_SECRET as _AUTH_JWT_SECRET,
    JWT_ALGORITHM as _AUTH_JWT_ALGORITHM,
)
from google.cloud import storage
from google.auth.exceptions import DefaultCredentialsError
from google.api_core.exceptions import Forbidden, Unauthorized
from google.oauth2 import service_account
from src.adapters.google_cloud import GeminiAuditor
from src.adapters.document_ai import DocumentAIAdapter
from src.core.data_contract import build_contract
import logging
log = logging.getLogger(__name__)

# ── Constantes de portfolio migradas desde db_writer (v1) ────────────────────
# Fuente de verdad para sector lookup en el flujo PDF/Excel upload legacy.
# En rutas nuevas, usar _bq_svc.get_portfolio_catalog() (datos vivos de BQ).

_COMPANY_BUCKET: dict[str, str] = {
    # ── Fondo VII ─────────────────────────────────────────────────────────────
    "conekta":     "SAAS",
    "kueski":      "LEND",
    "mpower":      "LEND",
    "bnext":       "SAAS",
    "yotepresto":  "LEND",
    "ivoy":        "ECOM",
    "bewe":        "SAAS",
    "skydropx":    "ECOM",
    "gaia":        "SAAS",
    # ── Fondo CIII ────────────────────────────────────────────────────────────
    "simetrik":    "SAAS",
    "guros":       "INSUR",
    "quinio":      "ECOM",
    "hackmetrix":  "SAAS",
    "hunty":       "SAAS",
    "atani":       "OTH",
    "cluvi":       "SAAS",
    "kuona":       "SAAS",
    "prometeo":    "OTH",
    "territorium": "SAAS",
    "morgana":     "INSUR",
    "duppla":      "LEND",
    "kala":        "OTH",
    "pulsar":      "SAAS",
    "solvento":    "LEND",
    "numia":       "SAAS",
    # ── Demo ──────────────────────────────────────────────────────────────────
    "demo-startup":  "SAAS",
    "demostartup":   "SAAS",
}

_PORTFOLIO_MAP: dict[str, dict] = {
    "fund_vii_overview": {"portfolio_id": "VII",  "portfolio_name": "Fondo VII", "display_name": "Fondo VII — Overview", "entity_type": "FUND_OVERVIEW"},
    "conekta":     {"portfolio_id": "VII",  "portfolio_name": "Fondo VII"},
    "kueski":      {"portfolio_id": "VII",  "portfolio_name": "Fondo VII"},
    "mpower":      {"portfolio_id": "VII",  "portfolio_name": "Fondo VII"},
    "bnext":       {"portfolio_id": "VII",  "portfolio_name": "Fondo VII"},
    "yotepresto":  {"portfolio_id": "VII",  "portfolio_name": "Fondo VII"},
    "ivoy":        {"portfolio_id": "VII",  "portfolio_name": "Fondo VII"},
    "bewe":        {"portfolio_id": "VII",  "portfolio_name": "Fondo VII"},
    "skydropx":    {"portfolio_id": "VII",  "portfolio_name": "Fondo VII"},
    "gaia":        {"portfolio_id": "VII",  "portfolio_name": "Fondo VII"},
    "simetrik":    {"portfolio_id": "CIII", "portfolio_name": "Fondo CIII"},
    "guros":       {"portfolio_id": "CIII", "portfolio_name": "Fondo CIII"},
    "quinio":      {"portfolio_id": "CIII", "portfolio_name": "Fondo CIII"},
    "hackmetrix":  {"portfolio_id": "CIII", "portfolio_name": "Fondo CIII"},
    "hunty":       {"portfolio_id": "CIII", "portfolio_name": "Fondo CIII"},
    "atani":       {"portfolio_id": "CIII", "portfolio_name": "Fondo CIII"},
    "cluvi":       {"portfolio_id": "CIII", "portfolio_name": "Fondo CIII"},
    "kuona":       {"portfolio_id": "CIII", "portfolio_name": "Fondo CIII"},
    "prometeo":    {"portfolio_id": "CIII", "portfolio_name": "Fondo CIII"},
    "territorium": {"portfolio_id": "CIII", "portfolio_name": "Fondo CIII"},
    "morgana":     {"portfolio_id": "CIII", "portfolio_name": "Fondo CIII"},
    "duppla":      {"portfolio_id": "CIII", "portfolio_name": "Fondo CIII"},
    "kala":        {"portfolio_id": "CIII", "portfolio_name": "Fondo CIII"},
    "pulsar":      {"portfolio_id": "CIII", "portfolio_name": "Fondo CIII"},
    "solvento":    {"portfolio_id": "CIII", "portfolio_name": "Fondo CIII"},
    "numia":       {"portfolio_id": "CIII", "portfolio_name": "Fondo CIII"},
    "demo-startup":  {"portfolio_id": "DEMO", "portfolio_name": "Demo Environment", "display_name": "Startup Demo", "entity_type": "DEMO"},
    "demostartup":   {"portfolio_id": "DEMO", "portfolio_name": "Demo Environment", "display_name": "Startup Demo", "entity_type": "DEMO"},
}

# Alias de compatibilidad — el código legacy referencia COMPANY_BUCKET y PORTFOLIO_MAP
# en mayúsculas (nombres originales de db_writer). Apuntan a los dicts locales.
COMPANY_BUCKET = _COMPANY_BUCKET
PORTFOLIO_MAP  = _PORTFOLIO_MAP


def _lookup_portfolio(company_id: str) -> str:
    """
    Infiere portfolio_id desde slug/dominio. Usa BQ catalog primero (datos vivos);
    cae al dict estático si BQ no está disponible. Devuelve "UNKNOWN" como fallback final.
    """
    slug = company_id.lower().split(".")[0].replace("-", "").replace("_", "")
    try:
        catalog = _bq_svc.get_portfolio_catalog()
        for entry in catalog:
            if entry["company_id"].lower() == slug or entry.get("company_name", "").lower() == slug:
                return entry["fund_id"]
    except Exception:
        pass
    # Fallback al dict estático
    for key, info in _PORTFOLIO_MAP.items():
        if slug in key.replace("-", "").replace("_", ""):
            return info["portfolio_id"]
    return "UNKNOWN"
from src.core.master_db_preprocessor import process_all as _process_all_master_db
from src.core.data_contract import (
    KPI_REGISTRY,
    parse_numeric,
    # build_checklist_status  — RETIRADO: solo usado en POST /upload (legacy)
    # validate_founder_submission — RETIRADO: solo usado en POST /upload (legacy)
)
# kpi_mapper: map_uploaded_file — RETIRADO: solo usado en endpoints 410 y confirm-gold (reescrito)
# kpi_dispatcher — ELIMINADO completamente en v2.0 (Star Schema).
# Toda escritura va por _bq_svc (BQDataService → BD_Cometa_Dev).
from src.core.local_db import (
    build_registry_rows,
    save_registry_rows,
    build_enriched_audit_response,
    build_kpi_status_grid,
    evaluate_commitment_gate,
    generate_jero_contract,
    check_gate_for_finalize,
    build_accumulated_kpi_grid,
    build_kpi_grid_from_contract_rows,
    build_registry_rows_from_contract,
    company_exists as _db_company_exists,
)
import pandas as pd
from src.core.bq_data_service import (
    BQDataService        as _BQDataServiceCatalog,
    BQInsertError        as _BQCatalogError,
    CompanyNotFoundError as _CompanyNotFoundError,
)

# Singleton para los endpoints que reemplazaron PORTFOLIO_MAP / COMPANY_BUCKET.
# La caché interna (TTL 5 min) garantiza que los dashboards carguen sin
# golpear BigQuery en cada render.
_bq_svc = _BQDataServiceCatalog()

app = FastAPI(title="Cometa Pipeline API", version="1.0.0")

# ── Routers ────────────────────────────────────────────────────────────────────
# founder_router expone /api/founder/ingest-kpis (ingesta directa a BD_Cometa_Dev).
# Las rutas legacy de founder (confirm-mapping, manual-update, finalize)
# siguen viviendo directamente en este módulo — include_router agrega solo
# las rutas del router sin eliminar las ya registradas en app.
from src.routers.founder import router as _founder_router
app.include_router(_founder_router)

from src.routers.analyst import router as _analyst_router
app.include_router(_analyst_router)

# ── A2: Rate limiting ──────────────────────────────────────────────────────────
limiter = Limiter(key_func=get_remote_address)
app.state.limiter = limiter
app.add_exception_handler(RateLimitExceeded, _rate_limit_exceeded_handler)


# ── E1: Manejo global de errores de validación Pydantic → 422 ─────────────────
#
# Por qué dos handlers:
#   - ValidationError      → errores internos (UserSchema.model_validate, etc.)
#   - RequestValidationError → errores de FastAPI al parsear el body/query params
#
# Ambos devuelven la misma estructura {detail: [{loc, msg, type}]} para que el
# frontend pueda parsearlos de forma uniforme sin lógica de ramificación.
#
# GARANTÍA DE ORDEN: los handlers se registran aquí, antes de cualquier ruta.
# Python evalúa el decorador en el momento de la definición, así que estos
# handlers están activos desde el primer request — incluido cualquier intento
# de escritura que falle en UserSchema.model_validate().

def _format_validation_errors(errors: list[dict]) -> list[dict]:
    """
    Normaliza la lista de errores de Pydantic v2 al subset {loc, msg, type}.
    Excluye 'url', 'input' y 'ctx' que son ruido para el cliente.
    """
    return [
        {
            "loc":  list(e.get("loc", [])),
            "msg":  e.get("msg", ""),
            "type": e.get("type", ""),
        }
        for e in errors
    ]


@app.exception_handler(ValidationError)
async def pydantic_validation_handler(
    request: Request,
    exc: ValidationError,
) -> JSONResponse:
    """
    Captura pydantic.ValidationError lanzado dentro de cualquier route handler.
    Ejemplo: UserSchema.model_validate() falla → este handler retorna 422
    ANTES de que se abra ningún archivo para escritura.
    """
    return JSONResponse(
        status_code=422,
        content={"detail": _format_validation_errors(exc.errors(include_url=False))},
    )


@app.exception_handler(RequestValidationError)
async def request_validation_handler(
    request: Request,
    exc: RequestValidationError,
) -> JSONResponse:
    """
    Sobreescribe el handler por defecto de FastAPI para request body / query params.
    Misma estructura {detail} que pydantic_validation_handler → frontend unificado.
    """
    print(f"❌ ERROR DE VALIDACIÓN DETECTADO en {request.method} {request.url.path}: {exc.errors()}")
    return JSONResponse(
        status_code=422,
        content={"detail": _format_validation_errors(exc.errors())},
    )


@app.on_event("startup")
async def _startup():
    """Bootstrap BigQuery analytics schema once at server start.

    Si no se puede conectar a cometa-mvp.BD_Cometa el servidor se detiene
    inmediatamente con un error claro — nunca arranca contra un dataset incorrecto.
    """
    import asyncio
    import sys

    from src.core.bq_data_service import _DATASET as _BQ_DATASET
    _bq_project = os.getenv("GOOGLE_PROJECT_ID", "cometa-mvp")
    print(f"🚀 Conectado a Producción: Proyecto {_bq_project} - Dataset {_BQ_DATASET}")

    try:
        # Startup health-check: verificar conectividad BQ con get_portfolio_catalog()
        # (TTL 5 min, non-blocking). En v2.0 ya no hay DDL propio — el schema lo
        # gestiona el equipo de data; nosotros solo insertamos filas.
        await asyncio.wait_for(
            asyncio.get_event_loop().run_in_executor(None, _bq_svc.get_portfolio_catalog),
            timeout=30.0,
        )
        print("✅ [Startup] Servidor listo para recibir archivos en :8000")
    except asyncio.TimeoutError:
        msg = (
            f"❌ [Startup] TIMEOUT (30 s) al conectar con {_bq_project}.{_bq_dataset}.\n"
            f"   Verifica conectividad y permisos IAM. El servidor se detiene."
        )
        print(msg)
        sys.exit(1)
    except Exception as exc:
        print(
            f"❌ [Startup] Fallo crítico de BigQuery — el servidor se detiene.\n{exc}"
        )
        sys.exit(1)

# ── CORS — orígenes permitidos desde variable de entorno ─────────────────────
# ALLOWED_ORIGINS acepta una lista separada por comas.
# Fallback seguro: solo localhost:3000 para desarrollo local.
# En producción (Cloud Run) inyectar vía Secret Manager o panel de env vars.
#
# Ejemplo .env:
#   ALLOWED_ORIGINS=https://cometa-vault-frontend-xxx.run.app,https://cometa.vc
#
# NOTA: allow_credentials puede cambiarse a True cuando los orígenes sean
# explícitos y la app use cookies de sesión. Con Bearer tokens, False es
# la opción más segura.

def _parse_allowed_origins() -> list[str]:
    raw = os.getenv("ALLOWED_ORIGINS", "").strip()
    if not raw:
        return ["http://localhost:3000"]
    return [o.strip() for o in raw.split(",") if o.strip()]

_ALLOWED_ORIGINS: list[str] = _parse_allowed_origins()

app.add_middleware(
    CORSMiddleware,
    allow_origins=_ALLOWED_ORIGINS,
    allow_credentials=False,
    allow_methods=["*"],
    allow_headers=["*"],
)

# ═══════════════════════════════════════════════════════════════════════════════
# BLOQUE DE SEGURIDAD
# ═══════════════════════════════════════════════════════════════════════════════

# ── C3: JWT Authentication ────────────────────────────────────────────────────
_bearer_scheme = HTTPBearer(auto_error=False)
_JWT_SECRET    = _AUTH_JWT_SECRET   # from auth_utils → JWT_SECRET env var
_JWT_ALGORITHM = _AUTH_JWT_ALGORITHM

# ── Auth: ruta al fichero de usuarios ─────────────────────────────────────────
_USERS_FILE = Path(__file__).parent / "users.json"

async def _require_auth(
    credentials: HTTPAuthorizationCredentials = Depends(_bearer_scheme),
) -> dict:
    """
    Valida el JWT HS256 emitido por /api/auth/token (Next.js).
    Lanza 401 si el token es inválido, expirado o ausente.
    """
    if not credentials:
        raise HTTPException(status_code=401, detail="Token de autenticación requerido")
    if not _JWT_SECRET:
        raise HTTPException(status_code=500, detail="NEXTAUTH_SECRET no configurado en el servidor")
    try:
        payload = jwt.decode(
            credentials.credentials,
            _JWT_SECRET,
            algorithms=[_JWT_ALGORITHM],
            options={"verify_aud": False},
        )
        return payload
    except JWTError as exc:
        raise HTTPException(status_code=401, detail=f"Token inválido: {exc}")


async def _require_analyst_auth(
    token: dict = Depends(_require_auth),
) -> dict:
    """
    Dependencia compuesta para todos los endpoints /api/analyst/*.

    Encadena tres capas de control de acceso:
      1. JWT válido y no expirado  (_require_auth ya lo garantiza).
      2. Dominio @cometa.vc        (defensa en profundidad más allá del login).
      3. Rol ANALISTA               (control de autorización estándar).

    Usar como: token: dict = Depends(_require_analyst_auth)
    """
    _check_cometa_domain(token)
    role = token.get("role")
    if role != "ANALISTA":
        raise HTTPException(
            status_code=403,
            detail="Acceso denegado. Se requiere rol ANALISTA con dominio @cometa.vc.",
        )
    return token


def _check_cometa_domain(token: dict) -> None:
    """
    Valida que el JWT pertenezca al dominio @cometa.vc.

    Se aplica a todos los endpoints /api/analyst/* y a /api/me cuando el rol
    es ANALISTA. Lanza 403 si el dominio no coincide — incluso si el JWT es
    técnicamente válido — para añadir defensa en profundidad más allá de la
    restricción de login.
    """
    email  = (token.get("email") or token.get("sub") or "").strip().lower()
    domain = email.split("@")[-1] if "@" in email else ""
    if domain != "cometa.vc":
        raise HTTPException(
            status_code=403,
            detail=(
                "Acceso denegado. Este recurso está restringido al dominio "
                "@cometa.vc. Token emitido para un dominio no autorizado."
            ),
        )


# ── C7: Magic bytes validation ────────────────────────────────────────────────
_MAGIC_BYTES: dict[str, list[bytes]] = {
    ".pdf":     [b"%PDF"],
    ".xlsx":    [b"PK\x03\x04"],
    ".xls":     [b"\xd0\xcf\x11\xe0"],
    ".docx":    [b"PK\x03\x04"],
    ".doc":     [b"\xd0\xcf\x11\xe0"],
    ".parquet": [b"PAR1"],
    ".csv":     [],  # Texto plano — no tiene magic bytes fijos
}

def _validate_magic_bytes(file_content: bytes, ext: str) -> bool:
    """
    Verifica que los primeros bytes del contenido coincidan con la extensión declarada.
    Protege contra archivos renombrados (p.ej. malware.exe → informe.pdf).
    """
    signatures = _MAGIC_BYTES.get(ext, [])
    if not signatures:
        return True
    return any(file_content[:8].startswith(sig) for sig in signatures)

# ── C2: Límite de tamaño de archivo ──────────────────────────────────────────
_MAX_FILE_MB    = int(os.getenv("MAX_FILE_SIZE_MB", "50"))
_MAX_FILE_BYTES = _MAX_FILE_MB * 1024 * 1024

# ── C6: Sanitización de nombre de archivo ────────────────────────────────────
_SAFE_FILENAME_RE = re.compile(r"[^\w\-.]")

def _sanitize_filename(filename: str) -> str:
    """
    Protege contra path traversal y caracteres peligrosos en nombres de archivo.
    Pasos: normalizar unicode → extraer basename → eliminar chars no seguros
           → eliminar puntos dobles → limitar a 200 chars.
    """
    filename = unicodedata.normalize("NFKD", filename)
    filename = os.path.basename(filename)                     # Bloquea ../../
    filename = _SAFE_FILENAME_RE.sub("_", filename)           # Solo alfanum + -_.
    filename = re.sub(r"\.{2,}", ".", filename)               # Elimina ..
    stem, ext = os.path.splitext(filename)
    return f"{stem[:196]}{ext}" if len(filename) > 200 else filename

# ── C5: Validación de headers de entrada ─────────────────────────────────────
_COMPANY_ID_RE = re.compile(r"^[a-zA-Z0-9_\-\.]{1,64}$")

def _validate_company_header(company_id: str | None) -> str | None:
    """Valida que company_id sea alfanumérico + guiones/puntos (sin path traversal)."""
    if not company_id:
        return None
    if not _COMPANY_ID_RE.match(company_id):
        raise HTTPException(
            status_code=400,
            detail=f"company_id contiene caracteres no permitidos: {company_id!r}"
        )
    return company_id

def _validate_email_header(email: str | None) -> str | None:
    """Valida formato básico de email. No verifica entregabilidad."""
    if not email:
        return None
    _EMAIL_RE = re.compile(r"^[a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,}$")
    if not _EMAIL_RE.match(email):
        raise HTTPException(status_code=400, detail=f"founder-email inválido: {email!r}")
    return email.lower()

# ── C4: Verificación de origen / preparación Cloud IAP ───────────────────────
_SKIP_ORIGIN_CHECK = os.getenv("SKIP_ORIGIN_CHECK", "false").lower() == "true"
_INTERNAL_SOURCE_HEADER = "x-cometa-source"
_IAP_USER_HEADER = "x-goog-authenticated-user-email"
_VALID_COMETA_SOURCES = {"dashboard", "analyst-portal", "internal-tool"}

async def _verify_origin(request: Request) -> None:
    """
    C4: Verifica que la petición provenga de una fuente autorizada.
    En producción (Cloud IAP) valida el header X-Goog-Authenticated-User-Email.
    En entornos sin IAP, acepta el header X-Cometa-Source con valor válido.
    SKIP_ORIGIN_CHECK=true lo deshabilita para desarrollo local.
    """
    if _SKIP_ORIGIN_CHECK:
        return

    # Cloud IAP en producción inyecta este header automáticamente
    iap_user = request.headers.get(_IAP_USER_HEADER)
    if iap_user:
        return  # IAP verificó la identidad; request autorizada

    # Fallback para entornos sin IAP (staging interno, tests de integración)
    source = request.headers.get(_INTERNAL_SOURCE_HEADER, "").strip().lower()
    if source in _VALID_COMETA_SOURCES:
        return

    raise HTTPException(
        status_code=403,
        detail="Acceso denegado: origen no autorizado. Se requiere X-Goog-Authenticated-User-Email o X-Cometa-Source válido.",
    )

# ── A1: Derivación de tenant desde JWT ───────────────────────────────────────
_INTERNAL_DOMAINS = {"cometa.vc", "cometa.fund", "cometavc.com"}

def _derive_tenant_from_token(token: dict) -> str | None:
    """
    A1: Extrae company_id del dominio del email en el JWT.
    - Analistas internos (@cometa.vc, etc.) → None (pueden consultar cualquier empresa)
    - Founders externos → company_id derivado de su dominio de email
    El llamador NO puede sobreescribir esto con company_id del body.
    """
    email: str = token.get("email", "")
    if not email or "@" not in email:
        return None
    domain = email.split("@", 1)[1].lower()
    if domain in _INTERNAL_DOMAINS:
        return None  # Analista interno — sin restricción de tenant
    # Founder externo: derivar company_id canónico desde su dominio
    comp_id, _, _, _ = get_company_id(domain)
    return comp_id


# ═══════════════════════════════════════════════════════════════════════════════
# BLOQUE DE NORMALIZACIÓN DE CONTRATO DE DATOS
#   R1 — normalize_period()  →  fecha libre → PYYYYQxMyy
#   R2 — get_company_id()    →  nombre libre → COMP_XXX + fund_id + bucket_id
# ═══════════════════════════════════════════════════════════════════════════════

# ── R1: Tablas de traducción para normalización de períodos ───────────────────

_MONTH_TO_NUM: dict[str, str] = {
    # Inglés
    "january":"01","february":"02","march":"03","april":"04",
    "may":"05","june":"06","july":"07","august":"08",
    "september":"09","october":"10","november":"11","december":"12",
    # Español
    "enero":"01","febrero":"02","marzo":"03","abril":"04",
    "mayo":"05","junio":"06","julio":"07","agosto":"08",
    "septiembre":"09","octubre":"10","noviembre":"11","diciembre":"12",
    # Abreviaturas EN
    "jan":"01","feb":"02","mar":"03","apr":"04","jun":"06",
    "jul":"07","aug":"08","sep":"09","oct":"10","nov":"11","dec":"12",
}

_MONTH_TO_QUARTER: dict[str, str] = {
    "01":"Q1","02":"Q1","03":"Q1",
    "04":"Q2","05":"Q2","06":"Q2",
    "07":"Q3","08":"Q3","09":"Q3",
    "10":"Q4","11":"Q4","12":"Q4",
}

# Primer mes canónico de cada quarter (para cuando solo tenemos Q sin mes exacto)
_QUARTER_FIRST_MONTH: dict[str, str] = {
    "1":"01","2":"04","3":"07","4":"10",
}

# Período de cierre de cada semestre (H1 → jun, H2 → dic)
_HALF_CLOSE_MONTH: dict[str, str] = {"1":"06","2":"12"}
_HALF_QUARTER: dict[str, str]     = {"1":"Q2","2":"Q4"}

# Período ya en formato canónico — pasa sin transformación
_CANONICAL_RE = re.compile(r"^P(20\d{2})Q([1-4])M(\d{2})$")


def normalize_period(date_str: str) -> tuple[str, bool]:
    """
    R1 — Normaliza cualquier representación de período al formato PYYYYQxMyy.

    Entradas reconocidas (insensibles a mayúsculas/espacios extra):
      "March 2025"    → "P2025Q1M03"
      "Q1 2025"       → "P2025Q1M01"
      "Q4 2024"       → "P2024Q4M10"
      "H1 2025"       → "P2025Q2M06"
      "H2 2024"       → "P2024Q4M12"
      "FY2025"        → "P2025Q4M12"
      "2025"          → "P2025Q4M12"
      "2025M03"       → "P2025Q1M03"
      "P2025Q1M03"    → "P2025Q1M03"  (passthrough)

    Returns
    -------
    (canonical_period_id, is_valid)
      is_valid=False  →  el input no pudo mapearse; se devuelve un fallback
                        con el año actual para no romper el pipeline.
    """
    if not date_str or not isinstance(date_str, str):
        fallback = f"P{datetime.now(timezone.utc).year}Q4M12"
        return fallback, False

    s = date_str.strip()

    # 0. Ya canónico — passthrough sin coste
    if _CANONICAL_RE.match(s):
        return s, True

    sl = s.lower()

    # 1. "March 2025" / "marzo 2025" — nombre de mes + año
    for month_name, month_num in _MONTH_TO_NUM.items():
        m = re.search(rf"\b{re.escape(month_name)}\b\s*(20\d{{2}})", sl)
        if not m:
            # También acepta "2025 March"
            m = re.search(rf"(20\d{{2}})\s*\b{re.escape(month_name)}\b", sl)
        if m:
            year = m.group(1) if m.lastindex == 1 else (m.group(1) if m.group(1).startswith("20") else m.group(2))
            # Reparar: si el grupo de año no es el que tiene 20xx, tomar el otro
            year = next((g for g in m.groups() if g and g.startswith("20")), None)
            if not year:
                continue
            quarter = _MONTH_TO_QUARTER[month_num]
            return f"P{year}{quarter}M{month_num}", True

    # 2. "Q1 2025" / "Q4 2024"
    m = re.search(r"q([1-4])\s*[/\-]?\s*(20\d{2})", sl)
    if not m:
        m = re.search(r"(20\d{2})\s*[/\-]?\s*q([1-4])", sl)
    if m:
        groups = m.groups()
        if groups[0].startswith("20"):
            year, qnum = groups[0], groups[1]
        else:
            qnum, year = groups[0], groups[1]
        return f"P{year}Q{qnum}M{_QUARTER_FIRST_MONTH[qnum]}", True

    # 3. "H1 2025" / "H2 2024"
    m = re.search(r"h([12])\s*(20\d{2})", sl)
    if not m:
        m = re.search(r"(20\d{2})\s*h([12])", sl)
    if m:
        groups = m.groups()
        if groups[0].startswith("20"):
            year, half = groups[0], groups[1]
        else:
            half, year = groups[0], groups[1]
        return f"P{year}{_HALF_QUARTER[half]}M{_HALF_CLOSE_MONTH[half]}", True

    # 4. "FY2025" / "FY 2025" / "fiscal year 2025"
    m = re.search(r"fy\s*(20\d{2})", sl)
    if not m:
        m = re.search(r"fiscal\s+year\s*(20\d{2})", sl)
    if m:
        year = m.group(1)
        return f"P{year}Q4M12", True

    # 5. "2025M03" (formato partial canónico sin prefijo P)
    m = re.match(r"^(20\d{2})m(\d{2})$", sl)
    if m:
        year, month_num = m.group(1), m.group(2).zfill(2)
        quarter = _MONTH_TO_QUARTER.get(month_num, "Q4")
        return f"P{year}{quarter}M{month_num}", True

    # 6. Año suelto: "2025"
    m = re.fullmatch(r"20\d{2}", s.strip())
    if m:
        return f"P{s.strip()}Q4M12", True

    # 6b. "YYYYMMDD" o "YYYYMM" sin separador — extrae año y mes
    m = re.match(r"^(20\d{2})(\d{2})(?:\d{2})?$", s.strip())
    if m:
        year, month_num = m.group(1), m.group(2).zfill(2)
        quarter = _MONTH_TO_QUARTER.get(month_num, "Q4")
        return f"P{year}{quarter}M{month_num}", True

    # 7. Cualquier año 20xx + número de mes detectados en el string
    year_m  = re.search(r"(20\d{2})", s)
    month_m = re.search(r"\b(0?[1-9]|1[0-2])\b", s)
    if year_m and month_m:
        year      = year_m.group(1)
        month_num = month_m.group(1).zfill(2)
        quarter   = _MONTH_TO_QUARTER.get(month_num, "Q4")
        return f"P{year}{quarter}M{month_num}", True

    # 8. Solo año detectado — fallback degradado
    if year_m:
        return f"P{year_m.group(1)}Q4M12", False

    # 9. Sin información → año corriente, marcar como inválido
    fallback = f"P{datetime.now(timezone.utc).year}Q4M12"
    return fallback, False


# ── R2: Resolución de empresa desde BigQuery (BD_Cometa_Dev.dim_company) ──────
# get_company_id() consulta _bq_svc.get_portfolio_catalog() — resultado cacheado
# con TTL de 5 min en BQDataService. Sin estado local en este módulo.

def get_company_id(name_str: str) -> tuple[str, str, str, bool]:
    """
    R2 — Mapea texto libre al company_id canónico de BQ y hereda fund_id /
    bucket_name desde dim_company + dim_bucket (via BQDataService).

    Estrategia de resolución (para en el primer match):
      1. Match exacto contra company_id (ej. "C001")
      2. Match exacto contra company_name normalizado
      3. Strip guiones/guiones bajos y match exacto de company_name
      4. Prefijo: input starts with company_name token
      5. Substring: company_name en input normalizado o viceversa

    Returns
    -------
    (company_id, fund_id, bucket_name, is_known)
      is_known=False → empresa no registrada; company_id es COMP_UNKNOWN_<hash>
                       para garantizar trazabilidad sin perder la submission.
    """
    import hashlib as _hashlib

    if not name_str:
        return "COMP_UNKNOWN", "unknown", "OTH", False

    try:
        catalog = _bq_svc.get_portfolio_catalog()
    except Exception:
        catalog = []

    # ── Normalización base ────────────────────────────────────────────────
    base     = name_str.lower().strip()
    base     = base.split(".")[0]            # strip TLD: "solvento.com" → "solvento"
    base     = re.sub(r"\s+", " ", base)    # colapsar espacios
    stripped = re.sub(r"[-_\s]", "", base)

    def _entry(e: dict) -> tuple[str, str, str, bool]:
        return e["company_id"], e["fund_id"], e["bucket_name"], True

    # 1. Exacto contra company_id (ej. "C001")
    for e in catalog:
        if e["company_id"].lower() == base:
            return _entry(e)

    # 2. Exacto contra company_name normalizado
    for e in catalog:
        if e["company_name"].lower().strip() == base:
            return _entry(e)

    # 3. Strip separadores y match exacto de company_name
    for e in catalog:
        if re.sub(r"[-_\s]", "", e["company_name"].lower()) == stripped:
            return _entry(e)

    # 4. Prefijo: input starts with company_name token
    for e in sorted(catalog, key=lambda x: len(x["company_name"]), reverse=True):
        cname = e["company_name"].lower().strip()
        if base.startswith(cname + "-") or base.startswith(cname + "_") or base.startswith(cname + " "):
            return _entry(e)

    # 5. Substring: company_name en input o input en company_name
    for e in sorted(catalog, key=lambda x: len(x["company_name"]), reverse=True):
        cstripped = re.sub(r"[-_\s]", "", e["company_name"].lower())
        pattern   = r"(?<![a-z])" + re.escape(cstripped) + r"(?![a-z])"
        if re.search(pattern, stripped) or (stripped and stripped in cstripped):
            return _entry(e)

    # Sin match → COMP_UNKNOWN_<fingerprint> (determinístico, trazable)
    fingerprint = _hashlib.sha1(name_str.lower().encode()).hexdigest()[:8].upper()
    unknown_id  = f"COMP_UNKNOWN_{fingerprint}"
    print(f"[R2] '{name_str}' no encontrado en BQ → {unknown_id}")
    return unknown_id, "unknown", "OTH", False


def _apply_contract_normalization(
    contract: dict,
    raw_company: str,
    raw_period: str,
) -> dict:
    """
    Aplica R1 y R2 al contrato ya construido por build_contract().

    Muta el contrato in-place y devuelve un dict con los resultados
    de normalización para logging y trazabilidad.

    Regla de oro:
      Si period o company no se pueden mapear → submission.status = "error"
      El contrato SE MUTA en todo caso para garantizar que el pipeline
      continúa y no pierde datos: is_known/period_valid se registran en
      submission para que el analista pueda corregir manualmente.
    """
    # ── R1: Normalizar período ────────────────────────────────────────────
    # Prioridad: raw_period del _document_context de Gemini (más específico)
    # → contiene "March 2025", "Q1 2025", etc. que infer_period_id() ignora.
    # Fallback: el period_id ya inferido por build_contract() (solo tiene año).
    current_period = raw_period or contract["submission"].get("period_id", "")
    norm_period, period_ok = normalize_period(current_period)

    # ── R2: Normalizar company_id ─────────────────────────────────────────
    comp_id, fund_id, bucket_id, company_ok = get_company_id(raw_company)

    # ── Determinar status final ───────────────────────────────────────────
    normalization_errors: list[str] = []
    if not period_ok:
        normalization_errors.append(
            f"period '{current_period}' no reconocido — "
            f"se usó fallback '{norm_period}'"
        )
    if not company_ok:
        normalization_errors.append(
            f"company '{raw_company}' no está en el catálogo maestro — "
            f"se asignó '{comp_id}'"
        )

    # Solo marca error si ambos fallan simultáneamente (un error de compañía
    # con período conocido es recuperable vía edición manual del analista)
    if not period_ok and not company_ok:
        contract["submission"]["status"] = "error"
    elif not company_ok:
        contract["submission"]["status"] = "pending_review"

    # ── Mutar submission ──────────────────────────────────────────────────
    contract["submission"]["period_id"]            = norm_period
    contract["submission"]["company_id"]           = comp_id
    contract["submission"]["fund_id"]              = fund_id
    contract["submission"]["bucket_id"]            = bucket_id
    contract["submission"]["period_normalized"]    = period_ok
    contract["submission"]["company_known"]        = company_ok
    if normalization_errors:
        contract["submission"]["normalization_errors"] = normalization_errors

    # ── Mutar kpi_rows — propagar IDs canónicos a cada hecho ─────────────
    for row in contract["kpi_rows"]:
        row["period_id"]  = norm_period   # sobrescribe el inferido por build_contract
        row["company_id"] = comp_id       # añade company_id a la fila (requerido por contrato)
        row["fund_id"]    = fund_id

    return {
        "period_id":    norm_period,
        "period_ok":    period_ok,
        "comp_id":      comp_id,
        "fund_id":      fund_id,
        "bucket_id":    bucket_id,
        "company_ok":   company_ok,
        "errors":       normalization_errors,
    }


# ── Rutas de archivos estáticos ───────────────────────────────────────────────
# BASE_DIR apunta a la raíz del proyecto (/app en Cloud Run) sin importar
# desde qué directorio se invoque uvicorn. Evita problemas con rutas relativas.
BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
_assets_dir = os.path.join(BASE_DIR, "assets")
if os.path.isdir(_assets_dir):
    app.mount("/assets", StaticFiles(directory=_assets_dir), name="assets")

@app.get("/api/v2/health", include_in_schema=False)
async def root():
    return {"status": "backend_online", "version": "2.0"}


# ── Portfolio metadata — diccionarios que antes vivían en el frontend ──────────
# Centralizados aquí para que un cambio se propague a todos los clientes
# sin redeploy de Next.js.
#
# TODO (v2): sustituir _COVERAGE_MAP y _LAST_MONTH_MAP por queries directas a
# BD_Cometa_Dev.fact_kpi_values via _bq_svc.

_VERTICAL_LABEL: dict[str, str] = {
    "SAAS":  "SaaS",
    "LEND":  "Lending",
    "ECOM":  "eCommerce",
    "INSUR": "InsurTech",
    "OTH":   "Other",
}

def _build_vertical_map_bq(catalog: list[dict]) -> dict[str, str]:
    """
    Construye company_id → vertical label desde el catálogo de BQ.
    bucket_name ("SAAS", "FINTECH") se mapea a la etiqueta de display.
    """
    return {
        c["company_id"]: _VERTICAL_LABEL.get(c["bucket_name"], c["bucket_name"])
        for c in catalog
    }


def _build_vertical_map_legacy() -> dict[str, str]:
    """Fallback estático usando COMPANY_BUCKET. Solo se usa si BQ no está disponible."""
    return {
        slug: _VERTICAL_LABEL.get(bucket, "Other")
        for slug, bucket in COMPANY_BUCKET.items()
        if slug not in ("demo-startup", "demostartup")
    }

_COVERAGE_MAP_STATIC: dict[str, int] = {
    "pulsar": 84, "rintin": 71, "cluvi": 70, "m1": 68, "numia": 68,
    "simetrik": 65, "prometeo": 63, "kuona": 57, "solvento": 47,
    "territorium": 34, "hunty": 22, "quinio": 18, "kala": 13,
    "duppla": 1, "guros": 0, "hackmetrix": 0,
}

_LAST_MONTH_MAP_STATIC: dict[str, str] = {
    "simetrik": "Mar 2026",
    "cluvi": "Dic 2025", "hunty": "Dic 2025", "kuona": "Dic 2025",
    "numia": "Dic 2025", "prometeo": "Dic 2025", "pulsar": "Dic 2025",
    "duppla": "Dic 2025", "guros": "Dic 2025", "hackmetrix": "Dic 2025",
    "kala": "Dic 2025", "m1": "Dic 2025", "quinio": "Dic 2025",
    "rintin": "Dic 2025", "solvento": "Dic 2025", "territorium": "Dic 2025",
}


@app.get("/api/metadata", response_model=PortfolioMetadataResponse)
async def get_portfolio_metadata():
    """
    Devuelve los metadatos del portfolio que el frontend necesita para
    renderizar el dashboard: vertical de cada empresa, cobertura de KPIs y
    último período reportado.

    No requiere autenticación — los datos son informativos y no contienen
    información financiera sensible.

    vertical_map se construye leyendo dim_company + dim_bucket desde BQ
    (caché de 5 min).  Si BQ no está disponible cae al mapa estático legacy.
    """
    try:
        catalog      = _bq_svc.get_portfolio_catalog()
        vertical_map = _build_vertical_map_bq(catalog)
    except Exception as _exc:
        # BQ no disponible (local sin credenciales, etc.) → fallback estático
        import logging as _log
        _log.getLogger(__name__).warning(
            "[metadata] BQ no disponible, usando mapa estático legacy: %s", _exc
        )
        vertical_map = _build_vertical_map_legacy()

    return PortfolioMetadataResponse(
        vertical_map=vertical_map,
        coverage_map=_COVERAGE_MAP_STATIC,
        last_month_map=_LAST_MONTH_MAP_STATIC,
    )

# Configuración
PROJECT_ID = os.getenv("GOOGLE_PROJECT_ID", "cometa-mvp")
BQ_DATASET = os.getenv("GOOGLE_BIGQUERY_DATASET") or os.getenv("BIGQUERY_DATASET", "BD_Cometa")
LOCATION_DOC_AI = os.getenv("DOCUMENT_AI_LOCATION", "us")
PROCESSOR_ID = os.getenv("DOCUMENT_AI_PROCESSOR_ID", "c5e1adfde68e63cf")
VERTEX_LOCATION = os.getenv("VERTEX_AI_LOCATION", "us-central1")
from src.core.buckets import RAW_BUCKET, STAGE_BUCKET, GOLD_BUCKET, HIST_BUCKET
GCS_INPUT_BUCKET  = RAW_BUCKET    # alias legacy — usar RAW_BUCKET en código nuevo
GCS_OUTPUT_BUCKET = STAGE_BUCKET  # alias legacy — usar STAGE_BUCKET en código nuevo

def _resolve_service_account_path() -> str | None:
    env_path = os.getenv("GOOGLE_APPLICATION_CREDENTIALS")
    if env_path:
        return env_path
    # Fallback al JSON en raíz del repo
    repo_root = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))
    fallback = os.path.join(repo_root, "cometa_key.json")
    return fallback if os.path.exists(fallback) else None

def _parse_sa_json(raw: str, tag: str = "") -> dict:
    """
    Parsea el JSON de una service account desde una variable de entorno,
    tolerando: espacios/newlines extra y doble-serialización (string dentro de string).
    Lanza ValueError con diagnóstico si el resultado no tiene los campos requeridos.
    """
    raw = raw.strip()
    parsed = json.loads(raw)
    # Doble-serialización: el valor del secreto era una cadena JSON (string de string)
    if isinstance(parsed, str):
        print(f"⚠️  {tag} GCP_SERVICE_ACCOUNT_JSON estaba doblemente serializado — decodificando de nuevo")
        parsed = json.loads(parsed)
    if not isinstance(parsed, dict):
        raise ValueError(f"{tag} GCP_SERVICE_ACCOUNT_JSON no es un objeto JSON válido (tipo: {type(parsed).__name__})")
    required = {"type", "project_id", "private_key", "client_email"}
    missing = required - parsed.keys()
    if missing:
        raise ValueError(
            f"{tag} GCP_SERVICE_ACCOUNT_JSON le faltan campos: {missing}. "
            f"Claves presentes: {list(parsed.keys())}"
        )
    print(f"✅  {tag} Service account JSON OK — client_email: {parsed.get('client_email')}")
    return parsed


def _load_gcp_credentials():
    # ── Prioridad 1: GCP_SERVICE_ACCOUNT_JSON (Cloud Run + Secret Manager) ──────
    # El JSON completo se inyecta como variable de entorno desde Secret Manager.
    # No requiere archivo físico — compatible con contenedores inmutables.
    sa_json_str = os.getenv("GCP_SERVICE_ACCOUNT_JSON")
    if sa_json_str:
        print("🔐 [GCP] Usando GCP_SERVICE_ACCOUNT_JSON (Secret Manager)")
        sa_info = _parse_sa_json(sa_json_str, "[GCP]")
        creds = service_account.Credentials.from_service_account_info(sa_info)
        creds_project = getattr(creds, "project_id", None)
        if creds_project and creds_project != PROJECT_ID:
            print(f"⚠️  [GCP] project_id del JSON ({creds_project}) != PROJECT_ID ({PROJECT_ID})")
        return creds

    # ── Prioridad 2: GOOGLE_APPLICATION_CREDENTIALS (archivo, dev local) ────────
    sa_path = _resolve_service_account_path()
    if sa_path:
        if not os.path.isabs(sa_path):
            sa_path = os.path.abspath(sa_path)
        if os.path.exists(sa_path):
            print(f"🔐 [GCP] Usando Service Account JSON: {sa_path}")
            creds = service_account.Credentials.from_service_account_file(sa_path)
            creds_project = getattr(creds, "project_id", None)
            if creds_project and creds_project != PROJECT_ID:
                print(
                    f"⚠️  [GCP] project_id del JSON ({creds_project}) no coincide con PROJECT_ID ({PROJECT_ID})"
                )
            return creds

    # ── Prioridad 3: Application Default Credentials ──────────────────────────
    # Funciona automáticamente en Cloud Run con SA adjunta al servicio,
    # en GKE con Workload Identity, y localmente tras:
    #   gcloud auth application-default login
    try:
        import google.auth as _google_auth
        creds, adc_project = _google_auth.default()
        print(f"✅ [GCP] Usando Application Default Credentials (ADC) — proyecto: {adc_project}")
        return creds
    except Exception as adc_err:
        raise DefaultCredentialsError(
            "Ninguna credencial GCP disponible. Opciones:\n"
            "  · Cloud Run: adjunta una SA al servicio (no requiere JSON)\n"
            "  · CI/CD: inyecta GCP_SERVICE_ACCOUNT_JSON desde Secret Manager\n"
            "  · Local: ejecuta `gcloud auth application-default login`\n"
            f"  Error ADC: {adc_err}"
        ) from adc_err

def _get_storage_client() -> storage.Client:
    """Crea un Storage client usando credenciales explícitas cuando es posible."""
    try:
        creds = _load_gcp_credentials()
        return storage.Client(project=PROJECT_ID, credentials=creds)
    except Exception as e:
        print(f"❌ [GCP] No se pudieron cargar credenciales explícitas: {e}")
        # Intentar fallback a ADC (podría funcionar en Cloud Run/GCE)
        return storage.Client(project=PROJECT_ID)

def _validate_company_slug(slug: str) -> tuple[bool, list[str]]:
    """
    Valida que *slug* existe en el CSV histórico de historicofund.

    Lee gs://{HISTORICOFUND_BUCKET}/CIII/Batch_final_normalizado_CIII.csv y
    busca el slug (lowercase, normalizado) en las columnas comunes de nombre de
    empresa: 'startup', 'company', 'company_name', 'startup_name', 'nombre'.

    Returns
    -------
    (is_valid, valid_slugs)
        is_valid   — True si el slug fue encontrado en el CSV.
        valid_slugs — lista de slugs canónicos disponibles (siempre retornada
                      para que el analista pueda elegir el correcto).
    """
    import io as _io
    import csv as _csv
    import re as _re

    bucket_name = os.getenv("HISTORICOFUND_BUCKET", "historicofund")
    blob_path   = os.getenv(
        "HISTORICOFUND_CIII_FILE",
        "CIII/Kpi_master.csv",
    )

    def _to_slug(name: str) -> str:
        s = name.strip().lower()
        s = _re.sub(r"[^a-z0-9]+", "_", s)
        return s.strip("_")

    try:
        gcs_client = _get_storage_client()
        content    = gcs_client.bucket(bucket_name).blob(blob_path).download_as_text(encoding="utf-8")
        reader     = _csv.DictReader(_io.StringIO(content))

        # Detectar columna de nombre de empresa de forma flexible
        _NAME_COLS = {"startup", "company", "company_name", "startup_name", "nombre", "empresa"}
        col_name: str | None = None
        if reader.fieldnames:
            for fn in reader.fieldnames:
                if fn.strip().lower() in _NAME_COLS:
                    col_name = fn
                    break

        if col_name is None:
            # Usar primera columna como fallback
            col_name = (reader.fieldnames or ["startup"])[0]

        seen: dict[str, str] = {}  # slug → original name
        for row in reader:
            raw = row.get(col_name, "").strip()
            if raw:
                seen[_to_slug(raw)] = raw

        normalized = _to_slug(slug)
        valid_slugs = sorted(seen.keys())
        return (normalized in seen), valid_slugs

    except Exception as exc:
        # CSV no accesible — no bloqueamos el flujo, solo logueamos en DEBUG
        import logging as _logging
        _logging.getLogger(__name__).debug("[historicofund] CSV no accesible (non-fatal): %s", exc)
        return True, []


def _resolve_slug(raw_slug: str) -> tuple[str, str]:
    """
    Resuelve un slug crudo al slug canónico y al nombre de empresa para mostrar.

    Convierte IDs de empresa ('c010', 'COMP_C010') a nombres reales ('quinio')
    buscando en el catálogo de BQ.  Non-fatal: si BQ no responde, devuelve el
    slug tal cual con un nombre humanizado.

    Returns
    -------
    (slug_canónico, nombre_display)
      slug_canónico — slug legible y normalizado ('quinio', 'simetrik', …)
      nombre_display — nombre para mensajes al analista ('Quinio', 'Simetrik', …)
    """
    if not raw_slug:
        return "", ""

    slug_clean = raw_slug.lower().replace("comp_", "").strip()

    # 1. Coincidencia directa en _DOMAIN_SLUG_MAP (slug ya es canónico)
    for _, (map_slug, map_name) in _DOMAIN_SLUG_MAP.items():
        if map_slug == slug_clean:
            return map_slug, map_name

    # 2. Patrón de company_id: c010, c001, etc.
    _is_id_pattern = bool(re.match(r"^c\d{3,}$", slug_clean))
    if _is_id_pattern:
        try:
            catalog     = _bq_svc.get_portfolio_catalog()
            comp_upper  = slug_clean.upper()   # "C010"
            match       = next(
                (c for c in catalog if c.get("company_id", "").upper() == comp_upper),
                None,
            )
            if match:
                name      = match.get("company_name", "")
                canonical = re.sub(r"[^a-z0-9]+", "_", name.lower()).strip("_")
                return canonical, name
        except Exception as _exc:
            log.debug("[_resolve_slug] BQ lookup non-fatal: %s", _exc)

    # 3. Fallback: devolver tal cual con nombre humanizado
    display = slug_clean.replace("_", " ").replace("-", " ").title()
    return slug_clean, display


def _load_historicofund_map() -> dict[str, str]:
    """
    Carga el CSV de historicofund y devuelve un dict {slug -> official_name}.

    Usado por GET /api/analyst/buckets para mostrar el nombre oficial de cada
    startup y marcar si la empresa tiene datos históricos certificados.

    Returns {} (vacío) si el CSV no es accesible — el endpoint no se bloquea.
    """
    import io as _io
    import csv as _csv
    import re as _re

    bucket_name = os.getenv("HISTORICOFUND_BUCKET", "historicofund")
    blob_path   = os.getenv(
        "HISTORICOFUND_CIII_FILE",
        "CIII/Kpi_master.csv",
    )

    def _to_slug(name: str) -> str:
        s = name.strip().lower()
        s = _re.sub(r"[^a-z0-9]+", "_", s)
        return s.strip("_")

    try:
        gcs_client = _get_storage_client()
        content    = gcs_client.bucket(bucket_name).blob(blob_path).download_as_text(encoding="utf-8")
        reader     = _csv.DictReader(_io.StringIO(content))

        _NAME_COLS = {"startup", "company", "company_name", "startup_name", "nombre", "empresa"}
        col_name: str | None = None
        if reader.fieldnames:
            for fn in reader.fieldnames:
                if fn.strip().lower() in _NAME_COLS:
                    col_name = fn
                    break
        if col_name is None:
            col_name = (reader.fieldnames or ["startup"])[0]

        result: dict[str, str] = {}
        for row in reader:
            raw = row.get(col_name, "").strip()
            if raw:
                result[_to_slug(raw)] = raw
        return result

    except Exception as exc:
        import logging as _logging
        _logging.getLogger(__name__).debug("[historicofund] mapa no cargado (non-fatal): %s", exc)
        return {}


def get_file_hash(file_content: bytes) -> str:
    """Genera hash SHA-256 del contenido del archivo"""
    return hashlib.sha256(file_content).hexdigest()[:16]


def _build_gemini_kpi_schema() -> str:
    """
    Genera el esquema JSON de extracción para el prompt de Gemini a partir de
    loading_brain_v1.json — cubre los 89 KPIs GIVEN agrupados por categoría.

    El esquema usa rutas planas bajo financial_metrics_2025, que coinciden con
    las paths de KPI_REGISTRY en data_contract.py.

    Returns
    -------
    str  — bloque JSON listo para insertar en el prompt.
         Retorna el esquema estático mínimo si el brain no está disponible.
    """
    import json as _json
    from pathlib import Path as _Path

    brain_path = _Path(__file__).parent.parent / "assets" / "loading_brain_v1.json"
    try:
        with open(brain_path, encoding="utf-8") as f:
            brain = _json.load(f)

        # Solo KPIs GIVEN (los SILVER se calculan en BQ)
        given = [m for m in brain.get("metrics", []) if m.get("given_or_silver") == "GIVEN"]
        if not given:
            raise ValueError("No GIVEN metrics found in brain")

        # Agrupar por categoría preservando orden de KPI-ref
        cats: dict[str, list] = {}
        for m in sorted(given, key=lambda x: x.get("kpi_ref", "")):
            cat = m.get("category", "Other")
            cats.setdefault(cat, []).append(m)

        kpi_lines: list[str] = []
        all_items = list(cats.items())
        for cat_idx, (cat, kpis) in enumerate(all_items):
            kpi_lines.append(f'    /* ── {cat} ───────────────────────────── */')
            for k_idx, m in enumerate(kpis):
                mid   = m["metric_id"]
                dname = m["display_name"]
                inneg = " ★ INNEGOCIABLE" if m.get("innegociable") else ""
                dtype = m.get("data_type", "")
                unit  = m.get("unit", "")
                sector = m.get("sector_scope", "ALL")
                sector_note = f" | sector: {sector}" if sector != "ALL" else ""

                if dtype == "percentage":
                    ex = '"36%" o "-4.2%"'
                elif dtype == "currency":
                    ex = '"$4.2M" o "-$320K"'
                elif dtype == "integer":
                    ex = '"42"'
                else:
                    ex = f'"{unit or "valor"}"'

                # Trailing comma: all entries except the very last
                is_last_cat = (cat_idx == len(all_items) - 1)
                is_last_kpi = (k_idx == len(kpis) - 1)
                comma = "" if (is_last_cat and is_last_kpi) else ","

                kpi_lines.append(
                    f'    "{mid}": {{"value": "<{dname}{inneg}{sector_note} — ej: {ex}>",'
                    f' "confidence": <float 0.0-1.0>,'
                    f' "description": "<fuente exacta en el doc>"}}{comma}'
                )

        inner = "\n".join(kpi_lines)
        return (
            "{\n"
            '  "_document_context": {\n'
            '    "currency":    "<ISO 4217 de la moneda dominante, ej. \'MXN\'>",\n'
            '    "period":      "<período fiscal, ej. \'FY2025\', \'H1 2025\', \'Q4 2025\'>",\n'
            '    "scale":       "<escala: \'units\', \'thousands\', \'millions\', \'billions\'>",\n'
            '    "scale_notes": "<dónde se encontró la indicación de escala, o null>"\n'
            "  },\n"
            '  "financial_metrics_2025": {\n'
            + inner + "\n"
            "  }\n"
            "}"
        )

    except Exception as exc:
        print(f"[prompt] WARN: no se pudo generar schema dinámico desde brain: {exc}")
        # Fallback mínimo con los 7 innegociables
        return (
            "{\n"
            '  "_document_context": {\n'
            '    "currency": "<ISO 4217>", "period": "<FY/H1/Q>",\n'
            '    "scale": "<units|thousands|millions|billions>", "scale_notes": null\n'
            "  },\n"
            '  "financial_metrics_2025": {\n'
            '    "revenue":     {"value": "<Revenue total — ej: \'$4.2M\'>",     "confidence": <float>, "description": "<fuente>"},\n'
            '    "gross_profit":{"value": "<Gross Profit — ej: \'$1.8M\'>",      "confidence": <float>, "description": "<fuente>"},\n'
            '    "ebitda":      {"value": "<EBITDA — puede ser negativo>",        "confidence": <float>, "description": "<fuente>"},\n'
            '    "cash":        {"value": "<Caja disponible — ej: \'$9.7M\'>",   "confidence": <float>, "description": "<fuente>"},\n'
            '    "burn":        {"value": "<Burn mensual — ej: \'-$320K\'>",      "confidence": <float>, "description": "<fuente>"},\n'
            '    "mrr":         {"value": "<MRR — solo SaaS — ej: \'$350K\'>",   "confidence": <float>, "description": "<fuente>"},\n'
            '    "employees":   {"value": "<Headcount total — ej: \'42\'>",       "confidence": <float>, "description": "<fuente>"}\n'
            "  }\n"
            "}"
        )


def _is_financial_document(resultado: dict) -> bool:
    """
    Devuelve True si Gemini extrajo al menos 1 KPI financiero con valor real.
    Cualquiera de estos campos con un valor no-nulo califica el documento.
    Usado como gate antes de persistir en GCS / BigQuery.
    """
    fm = resultado.get("financial_metrics_2025")
    if not fm or not isinstance(fm, dict):
        return False

    # Rutas a los KPIs "core" — basta con que uno sea no-nulo
    _SENTINEL = {"", "null", "n/a", "--", "0", "none"}
    core_paths = [
        ["revenue_growth", "value"],
        ["base_metrics", "revenue", "value"],
        ["base_metrics", "ebitda", "value"],
        ["profit_margins", "gross_profit_margin", "value"],
        ["profit_margins", "ebitda_margin", "value"],
        ["cash_flow_indicators", "cash_in_bank_end_of_year", "value"],
        ["cash_flow_indicators", "annual_cash_flow", "value"],
        ["debt_ratios", "working_capital_debt", "value"],
        ["sector_metrics", "mrr", "value"],
        ["sector_metrics", "gmv", "value"],
        ["sector_metrics", "portfolio_size", "value"],
        ["sector_metrics", "loss_ratio", "value"],
    ]
    for path in core_paths:
        node = fm
        for key in path:
            if not isinstance(node, dict):
                node = None
                break
            node = node.get(key)
        if node is not None and str(node).strip().lower() not in _SENTINEL:
            return True
    return False


def _extract_kpi_confidence_scores(resultado: dict) -> dict[str, int]:
    """
    Extract per-KPI confidence scores from a parsed Gemini JSON result.

    Traverses every KPI path defined in KPI_REGISTRY. If a node has a
    ``confidence`` field (float 0.0–1.0 as instructed in the prompt), it is
    converted to an integer 0–100 and stored under the ``kpi_key``.

    KPIs without a value (null) or without a ``confidence`` field are omitted.

    Parameters
    ----------
    resultado : dict
        Parsed Gemini JSON (the ``resultado`` dict produced by the upload pipeline).

    Returns
    -------
    dict[str, int]
        Mapping of ``kpi_key`` → confidence integer (0–100).
        Empty dict if no confidence scores are available.
    """
    scores: dict[str, int] = {}
    for kpi_def in KPI_REGISTRY:
        path: list[str] = kpi_def["path"]
        kpi_key: str    = kpi_def["kpi_key"]

        node = resultado
        for segment in path:
            if not isinstance(node, dict):
                node = None
                break
            node = node.get(segment)

        if not isinstance(node, dict):
            continue

        raw_conf = node.get("confidence")
        if raw_conf is None:
            continue

        try:
            # Gemini returns float 0.0–1.0; convert to 0–100 integer
            conf_float = float(raw_conf)
            # Accept both 0-1 and 0-100 ranges defensively
            if conf_float <= 1.0:
                conf_int = round(conf_float * 100)
            else:
                conf_int = round(conf_float)
            scores[kpi_key] = max(0, min(100, conf_int))
        except (TypeError, ValueError):
            continue

    return scores


# ── RAG / Chat helpers ────────────────────────────────────────────────────────

def _get_bq_client_for_api():
    """BigQuery client using the same credential chain as the rest of the API."""
    from google.cloud import bigquery as bq
    try:
        creds = _load_gcp_credentials()
        return bq.Client(project=PROJECT_ID, credentials=creds)
    except Exception:
        return bq.Client(project=PROJECT_ID)


def _build_results_from_bq(company_id: str) -> list[dict]:
    """
    BigQuery fallback for /api/results.

    Reads fact_kpi_values for a company and synthesises AnalysisResult objects
    (one per distinct period_id) with a financial_metrics_2025 payload — the
    same structure the frontend's extractKPIs() already knows how to consume.

    Parameters
    ----------
    company_id : Lowercase canonical slug, e.g. ``"quinio"``.

    Returns
    -------
    list[dict]  — Empty list on any BQ error (never raises).
    """
    # kpi_key → (financial_metrics_2025 section, sub-key)
    # All known aliases from fact_kpi_values are listed; cash_at_hand is an
    # alternate column name used in some legacy loads alongside cash_in_bank_end_of_year.
    _KPI_PATH: dict[str, tuple[str, str]] = {
        "revenue":                  ("revenue",              "total_revenue"),
        "ebitda":                   ("income",               "net_income"),
        "net_income":               ("income",               "net_income"),
        "gross_profit_margin":      ("profit_margins",       "gross_profit_margin"),
        "gross_margin":             ("profit_margins",       "gross_profit_margin"),  # alias
        "ebitda_margin":            ("profit_margins",       "ebitda_margin"),
        "revenue_growth":           ("revenue_growth",       "value"),
        "cash_in_bank_end_of_year": ("cash_flow_indicators", "cash_in_bank_end_of_year"),
        "cash_at_hand":             ("cash_flow_indicators", "cash_in_bank_end_of_year"),  # alias
        "annual_cash_flow":         ("cash_flow_indicators", "annual_cash_flow"),
        "cogs":                     ("cost_structure",       "cogs"),
        "working_capital_debt":     ("debt_ratios",          "working_capital_debt"),
        "net_working_capital":      ("debt_ratios",          "net_working_capital"),
        "mrr":                      ("revenue",              "mrr"),
        "gmv":                      ("revenue",              "gmv"),
    }

    def _fmt(value: float, unit: str) -> str:
        """Format a raw numeric value into a display string matching the frontend's toK() expectations."""
        if unit == "%":
            return f"{value:.1f}%"
        if abs(value) >= 1_000_000:
            return f"${value / 1_000_000:.1f}M"
        if abs(value) >= 1_000:
            return f"${value / 1_000:.0f}K"
        return f"{value:,.2f}"

    try:
        from google.cloud import bigquery as _bq
        _client = _get_bq_client_for_api()

        sql = f"""
            SELECT
                submission_id,
                period_id,
                metric_name_std         AS kpi_key,
                value                   AS num_value,
                COALESCE(unit_type, '') AS unit,
                value_status
            FROM `{PROJECT_ID}.{BQ_DATASET}.Vista_valores_H`
            WHERE LOWER(company_id) = LOWER(@company_id)
              AND value IS NOT NULL
            ORDER BY period_id ASC, metric_name_std
        """
        job  = _client.query(
            sql,
            job_config=_bq.QueryJobConfig(query_parameters=[
                _bq.ScalarQueryParameter("company_id", "STRING", company_id)
            ])
        )
        rows = list(job.result())

        if not rows:
            return []

        # Group rows by period_id
        from collections import defaultdict
        by_period: dict[str, list] = defaultdict(list)
        for r in rows:
            by_period[r.period_id].append(r)

        # ── First pass: raw numeric values per period (for derived calculations) ──
        raw_by_period: dict[str, dict[str, float]] = {}
        for pid, period_rows in by_period.items():
            raw_by_period[pid] = {}
            for r in period_rows:
                if r.num_value is not None and r.kpi_key:
                    raw_by_period[pid][r.kpi_key] = float(r.num_value)

        sorted_periods = sorted(by_period.keys())

        results = []
        for i, period_id in enumerate(sorted_periods):
            period_rows = by_period[period_id]
            raw_vals    = raw_by_period[period_id]

            # Build financial_metrics_2025 from flat BQ rows
            fm: dict = {}
            for r in period_rows:
                if r.num_value is None:
                    continue  # skip nulls — leave section absent so frontend shows "—"
                path = _KPI_PATH.get(r.kpi_key)
                if not path:
                    continue
                section, subkey = path
                fm.setdefault(section, {})
                # Don't overwrite a key already populated by an earlier alias
                if subkey != "value" and subkey in fm[section]:
                    continue
                unit_str    = r.unit or ""
                display_val = _fmt(float(r.num_value), unit_str)
                if subkey == "value":
                    fm[section]["value"] = {"value": display_val, "unit": unit_str}
                else:
                    fm[section][subkey] = {"value": display_val, "unit": unit_str}

            # ── Derived: Revenue Growth from consecutive periods ──────────────
            if "value" not in fm.get("revenue_growth", {}):
                if i > 0:
                    rev_now  = raw_vals.get("revenue")
                    rev_prev = raw_by_period[sorted_periods[i - 1]].get("revenue")
                    if rev_now is not None and rev_prev and rev_prev != 0:
                        growth = (rev_now - rev_prev) / abs(rev_prev) * 100
                        fm.setdefault("revenue_growth", {})["value"] = {
                            "value": f"{growth:.1f}%", "unit": "%",
                        }

            # ── Derived: Gross Margin = (revenue − cogs) / revenue ────────────
            if "gross_profit_margin" not in fm.get("profit_margins", {}):
                rev  = raw_vals.get("revenue")
                cogs = raw_vals.get("cogs")
                if rev and rev != 0 and cogs is not None:
                    gm = (rev - cogs) / rev * 100
                    fm.setdefault("profit_margins", {})["gross_profit_margin"] = {
                        "value": f"{gm:.1f}%", "unit": "%",
                    }

            # ── Derived: EBITDA Margin = ebitda / revenue ─────────────────────
            if "ebitda_margin" not in fm.get("profit_margins", {}):
                rev    = raw_vals.get("revenue")
                ebitda = raw_vals.get("ebitda")
                if rev and rev != 0 and ebitda is not None:
                    em = ebitda / rev * 100
                    fm.setdefault("profit_margins", {})["ebitda_margin"] = {
                        "value": f"{em:.1f}%", "unit": "%",
                    }

            # ── Fidelity calculation ───────────────────────────────────────
            # A row is "verified" when the view returns value_status='verified'
            # (driven by is_manually_edited=TRUE in fact_kpi_values).
            non_null_rows   = [r for r in period_rows if r.num_value is not None]
            verified_count  = sum(1 for r in non_null_rows if r.value_status == "verified")
            total_count     = len(non_null_rows)
            fidelity_pct    = int(verified_count / total_count * 100) if total_count else 0
            period_status   = "verified" if fidelity_pct == 100 else "legacy"
            submission_ids  = list({r.submission_id for r in period_rows if r.submission_id})

            result_item = {
                "id":         f"legacy_{period_id}",
                "company_id": company_id,
                "slug":       company_id.lower(),
                "date":       period_id,
                "data": {
                    "financial_metrics_2025": fm,
                    "_source":          "bigquery_legacy",
                    "_period_id":       period_id,
                    "_value_status":    period_status,
                    "_fidelity_pct":    fidelity_pct,
                    "_submission_ids":  submission_ids,
                },
                "metadata": {
                    "original_filename": f"histórico {period_id}",
                    "founder_email":     "",
                    "file_hash":         "",
                    "processed_at":      period_id,
                    "gcs_path":          "",
                    "company_domain":    company_id,
                    "portfolio_id":      "",
                },
            }
            results.append(result_item)

        print(f" [API/BQ] Resultados históricos para '{company_id}': "
              f"{len(results)} periodos, {len(rows)} filas")
        return results

    except Exception as _err:
        print(f" [API/BQ] Fallback BQ failed for '{company_id}' (non-fatal): {_err}")
        return []


# DEPRECATED — solo era llamado desde GET /api/results/all (ahora retirado).
# Mantener hasta confirmar que ningún otro path lo invoca.
def _build_all_results_from_bq(gcs_companies: set[str]) -> list[dict]:
    """
    Returns ONE synthetic AnalysisResult per company (latest available period)
    for companies NOT already covered by gcs_companies.

    KPI values are formatted for the extractTopKpis() path in the portfolio page:
      revenue_growth    → revenue_growth.value.value  (formatted as "X.X%")
      gross_profit_margin → profit_margins.gross_profit_margin.value
      ebitda_margin     → profit_margins.ebitda_margin.value

    Revenue Growth is calculated as:
      If revenue_growth kpi_key present → use directly.
      Otherwise → (revenue_last - revenue_prev) / |revenue_prev| * 100.

    Returns empty list on any error (never raises).
    """
    _KPI_PATH: dict[str, tuple[str, str]] = {
        "revenue":                  ("revenue",              "total_revenue"),
        "ebitda":                   ("income",               "net_income"),
        "net_income":               ("income",               "net_income"),
        "gross_profit_margin":      ("profit_margins",       "gross_profit_margin"),
        "gross_margin":             ("profit_margins",       "gross_profit_margin"),  # alias
        "ebitda_margin":            ("profit_margins",       "ebitda_margin"),
        "revenue_growth":           ("revenue_growth",       "value"),
        "cash_in_bank_end_of_year": ("cash_flow_indicators", "cash_in_bank_end_of_year"),
        "cash_at_hand":             ("cash_flow_indicators", "cash_in_bank_end_of_year"),  # alias
        "annual_cash_flow":         ("cash_flow_indicators", "annual_cash_flow"),
        "cogs":                     ("cost_structure",       "cogs"),
        "working_capital_debt":     ("debt_ratios",          "working_capital_debt"),
        "net_working_capital":      ("debt_ratios",          "net_working_capital"),
        "mrr":                      ("revenue",              "mrr"),
        "gmv":                      ("revenue",              "gmv"),
    }

    def _fmt(value: float, unit: str) -> str:
        if unit == "%":
            return f"{value:.1f}%"
        if abs(value) >= 1_000_000:
            return f"${value / 1_000_000:.1f}M"
        if abs(value) >= 1_000:
            return f"${value / 1_000:.0f}K"
        return f"{value:,.2f}"

    try:
        from collections import defaultdict as _dd
        _client = _get_bq_client_for_api()

        # Latest period per company + one-prior period for revenue growth calc
        sql = f"""
            WITH ranked AS (
                SELECT
                    company_id,
                    period_id,
                    metric_name_std         AS kpi_key,
                    value                   AS num_value,
                    COALESCE(unit_type, '') AS unit,
                    value_status,
                    ROW_NUMBER() OVER (
                        PARTITION BY company_id, metric_name_std
                        ORDER BY period_id DESC
                    ) AS rn
                FROM `{PROJECT_ID}.{BQ_DATASET}.Vista_valores_H`
                WHERE value IS NOT NULL
            )
            SELECT company_id, period_id, kpi_key, num_value, unit, value_status, rn
            FROM ranked
            WHERE rn <= 2
            ORDER BY company_id, kpi_key, rn
        """
        rows = list(_client.query(sql).result())

        # Separate: rn=1 (latest) and rn=2 (prior) per company+kpi
        latest: dict[tuple, float | None] = {}   # (company, kpi) → value
        prior:  dict[tuple, float | None] = {}
        period_map: dict[str, str]         = {}   # company → latest period_id

        for r in rows:
            # missing_legacy rows mark expected-but-absent values — exclude from calcs
            if r.value_status == "missing_legacy" or r.num_value is None:
                continue
            key = (r.company_id, r.kpi_key)
            if r.rn == 1:
                latest[key]              = r.num_value
                period_map[r.company_id] = r.period_id
            else:
                prior[key] = r.num_value

        # Build one result per company
        all_companies = {r.company_id for r in rows}
        results: list[dict] = []

        for company_id in sorted(all_companies):
            if company_id.lower() in gcs_companies:
                continue   # already covered by GCS

            period_id = period_map.get(company_id, "")

            # ── financial_metrics_2025 payload ───────────────────────────────
            fm: dict = {}
            for kpi_key, (section, subkey) in _KPI_PATH.items():
                val = latest.get((company_id, kpi_key))
                if val is None:
                    continue
                unit_key = next(
                    (r.unit for r in rows if r.company_id == company_id and r.kpi_key == kpi_key and r.rn == 1),
                    ""
                ) or ""

                # Revenue Growth: prefer stored value, else compute from revenue
                if kpi_key == "revenue_growth":
                    display = _fmt(val, "%")
                else:
                    display = _fmt(val, unit_key)

                fm.setdefault(section, {})
                if subkey == "value":
                    fm[section]["value"] = {"value": display, "unit": unit_key}
                else:
                    fm[section][subkey] = {"value": display, "unit": unit_key}

            # Derived Revenue Growth if kpi not stored
            if "revenue_growth" not in fm.get("revenue_growth", {}):
                rev_now  = latest.get((company_id, "revenue"))
                rev_prev = prior.get((company_id,  "revenue"))
                if rev_now is not None and rev_prev and rev_prev != 0:
                    growth = (rev_now - rev_prev) / abs(rev_prev) * 100
                    fm.setdefault("revenue_growth", {})["value"] = {
                        "value": f"{growth:.1f}%", "unit": "%"
                    }

            # Derived Gross Margin from revenue + cogs if not stored
            if "gross_profit_margin" not in fm.get("profit_margins", {}):
                rev  = latest.get((company_id, "revenue"))
                cogs = latest.get((company_id, "cogs"))
                if rev and rev != 0 and cogs is not None:
                    gm = (rev - cogs) / rev * 100
                    fm.setdefault("profit_margins", {})["gross_profit_margin"] = {
                        "value": f"{gm:.1f}%", "unit": "%"
                    }

            # Derived EBITDA Margin from ebitda + revenue if not stored
            if "ebitda_margin" not in fm.get("profit_margins", {}):
                rev    = latest.get((company_id, "revenue"))
                ebitda = latest.get((company_id, "ebitda"))
                if rev and rev != 0 and ebitda is not None:
                    em = ebitda / rev * 100
                    fm.setdefault("profit_margins", {})["ebitda_margin"] = {
                        "value": f"{em:.1f}%", "unit": "%"
                    }

            results.append({
                "id":   f"legacy_{company_id}_{period_id}",
                "date": period_id,
                "data": {
                    "financial_metrics_2025": fm,
                    "_source":    "bigquery_legacy",
                    "_period_id": period_id,
                },
                "metadata": {
                    "original_filename": f"histórico {period_id}",
                    "founder_email":     "",
                    "file_hash":         "",
                    "processed_at":      period_id,
                    "gcs_path":          "",
                    "company_domain":    company_id,
                    "portfolio_id":      _lookup_portfolio(company_id),
                },
            })

        print(f"[API/all/BQ] {len(results)} empresas históricas desde BigQuery")
        return results

    except Exception as _err:
        print(f"[API/all/BQ] Fallback failed (non-fatal): {_err}")
        return []


def _resolve_company_name(company_id: str | None) -> str | None:
    """Resuelve el nombre legible de una empresa a partir de su company_id.

    Consulta el catálogo en caché de BQDataService (TTL 5 min).
    Devuelve None si company_id es None o no se encuentra en el catálogo.
    No lanza excepciones — degradación graceful.
    """
    if not company_id:
        return None
    try:
        catalog = _bq_svc.get_portfolio_catalog()
        for entry in catalog:
            if entry.get("company_id", "").upper() == company_id.upper():
                return entry.get("company_name") or None
    except Exception:
        pass
    return None


def _query_rag_context(portfolio_id: str | None, company_id: str | None) -> list[dict]:
    """
    Fetches KPI rows from BD_Cometa_Dev.fact_kpi_values via BQDataService.

    Fuente única: BD_Cometa_Dev (Star Schema v2.0).
    La fusión con fact_portfolio_kpis (legacy MasterDB) fue eliminada —
    todo el data histórico debe migrarse a fact_kpi_values.
    """
    try:
        return _bq_svc.get_rag_context(
            company_id=company_id,
            fund_id=portfolio_id,
            limit=400,
        )
    except Exception as exc:
        print(f"[RAG] get_rag_context failed (non-fatal): {exc}")
        return []


# ── KPI Dictionary for RAG ────────────────────────────────────────────────────

def _fetch_kpi_dict_for_rag() -> dict[str, dict]:
    """
    Fetch the full KPI dictionary from dim_metric keyed by kpi_key (metric_id).

    Returns a dict mapping kpi_key → {display_name, description, unit,
    min_historical_year, vertical} for use in the Gemini prompt.

    Non-fatal: on BQ error returns an empty dict so the RAG prompt is
    built without metadata (graceful degradation, no 500 thrown).
    """
    try:
        items = _bq_svc.get_kpi_metadata()
        return {
            item["kpi_key"]: {
                "display_name":        item.get("display_name", ""),
                "description":         item.get("description", ""),
                "unit":                item.get("unit", ""),
                "min_historical_year": item.get("min_historical_year"),
                "vertical":            item.get("vertical", "GENERAL"),
            }
            for item in items
        }
    except Exception as exc:
        print(f"[RAG/dict] KPI metadata fetch failed (non-fatal): {exc}")
        return {}


# ── A3: RAG Leak Protection ───────────────────────────────────────────────────

def _verify_rag_integrity(
    rows: list[dict],
    expected_company_id: str,
) -> list[dict]:
    """
    Post-fetch verification that every BQ row belongs to the requested company.

    Controle A3: La query de BigQuery ya filtra por company_id, pero esta función
    actúa como segunda línea de defensa en caso de que el filtro LIKE sea demasiado
    permisivo o sea bypasseado por una condición de carrera.

    Lógica:
    - Si expected_company_id está vacío, no hay restricción → devuelve todo.
    - Para cada row, verifica que su company_id CONTENGA la cadena esperada
      (case-insensitive). Rows que no coincidan son "contaminados".
    - Rows contaminados se eliminan del contexto y se emite una SECURITY ALERT.
    - Si TODOS los rows son contaminados (bypass total del filtro BQ),
      se lanza HTTPException 500 — generación abortada.

    Returns
    -------
    list[dict]  — Solo rows validados para el company_id solicitado.

    Raises
    ------
    HTTPException 500  — Si la contaminación es total (fuga de datos detectada).
    """
    if not expected_company_id or not rows:
        return rows

    needle = expected_company_id.lower().strip()
    clean: list[dict] = []
    contaminated: list[dict] = []

    for row in rows:
        row_company = str(row.get("company_id") or "").lower().strip()
        # Bidirectional containment: handles "solvento" ↔ "solvento.com"
        if needle in row_company or row_company in needle:
            clean.append(row)
        else:
            contaminated.append(row)

    if contaminated:
        leaked_companies = list({r.get("company_id", "?") for r in contaminated})
        print(
            f"🚨 [RAG/A3] SECURITY ALERT — {len(contaminated)} row(s) contaminado(s) "
            f"de {len(rows)} para company='{expected_company_id}'. "
            f"Companies ajenas detectadas: {leaked_companies}"
        )
        if not clean:
            # Contaminación total — el filtro BQ puede haber sido bypasseado.
            raise HTTPException(
                status_code=500,
                detail=(
                    "[A3] Integridad del contexto comprometida: ningún dato "
                    "pertenece a la empresa solicitada. Consulta abortada por seguridad."
                ),
            )

    return clean


# ── Multi-format processing helpers ──────────────────────────────────────────

# Maximum rows rendered per sheet to stay within Gemini's token budget.
# 500 rows ≈ 40–80 KB of Markdown — well within the 1 M-token context window.
_MAX_ROWS_PER_SHEET = 500


def _df_to_markdown(df: pd.DataFrame) -> str:
    """
    Convert a DataFrame to a GitHub-flavored Markdown table.
    No external dependencies (no tabulate required).
    """
    if df.empty:
        return "*(tabla vacía)*"

    # Stringify column names
    cols = [str(c).strip() for c in df.columns]
    header    = "| " + " | ".join(cols) + " |"
    separator = "| " + " | ".join(["---"] * len(cols)) + " |"

    rows = []
    for _, row in df.iterrows():
        cells = [
            str(v).replace("|", "\\|").strip() if pd.notna(v) else ""
            for v in row
        ]
        rows.append("| " + " | ".join(cells) + " |")

    return "\n".join([header, separator] + rows)


def _process_tabular(file_path: str, ext: str, gemini, prompt_schema: str) -> str:
    """
    Read ALL sheets from CSV / XLSX / PARQUET, convert each to Markdown,
    then call Gemini with the identical financial-audit prompt used for PDFs.

    Changes vs. previous implementation:
    - XLSX: reads every sheet (sheet_name=None), not just the first one.
    - Row cap raised to _MAX_ROWS_PER_SHEET (500) per sheet.
    - Output format: GitHub-flavored Markdown tables (better structure preservation).
    - Prompt: uses the full FASE 1 + FASE 2 schema (same as PDF pipeline).
    """
    print(f"📊 [Tabular] Leyendo archivo {ext} (todas las hojas)...")

    try:
        sheets: dict = {}

        if ext == ".csv":
            df = pd.read_csv(file_path, nrows=_MAX_ROWS_PER_SHEET)
            sheets["Hoja1"] = df

        elif ext in (".xlsx", ".xls"):
            # sheet_name=None → devuelve un dict {nombre: DataFrame}
            all_sheets = pd.read_excel(
                file_path, sheet_name=None, engine="openpyxl"
            )
            for name, df in all_sheets.items():
                sheets[str(name)] = df.head(_MAX_ROWS_PER_SHEET)

        elif ext == ".parquet":
            df = pd.read_parquet(file_path).head(_MAX_ROWS_PER_SHEET)
            sheets["Hoja1"] = df

        else:
            raise ValueError(f"Extensión tabular no reconocida: {ext}")

    except Exception as e:
        print(f"❌ [Tabular] Error leyendo archivo: {e}")
        raise RuntimeError(f"No se pudo leer el archivo {ext}: {e}") from e

    _sheet_keys = list(sheets.keys())
    print(f"   Pestañas encontradas: {_sheet_keys}")

    # ── Build Markdown content (one section per sheet) ────────────────────
    md_sections = []
    for sheet_name, df in sheets.items():
        # Drop columns that are 100 % null — they add noise, no signal
        df = df.dropna(axis=1, how="all")
        _n_rows = df.shape[0]
        _n_cols = df.shape[1]
        print(f"   Hoja '{sheet_name}': {_n_rows} filas × {_n_cols} columnas")
        md = _df_to_markdown(df)
        md_sections.append(f"## Pestaña: {sheet_name}\n\n{md}")

    full_markdown = "\n\n---\n\n".join(md_sections)
    _md_len      = len(full_markdown)
    _sheet_count = len(sheets)
    print(f"   Markdown generado: {_md_len:,} caracteres, {_sheet_count} pestaña(s)")

    # ── Prompt: same FASE 1 + FASE 2 schema as the PDF pipeline ─────────
    # Prepend a short adapter note so Gemini knows the source is a spreadsheet,
    # then inject the full audit schema unchanged.
    _ext_upper   = ext.upper()
    _n_sheets    = len(sheets)
    adapter_header = (
        f"Eres un auditor financiero senior especializado en due diligence de startups.\n"
        f"Recibes el contenido completo de un archivo {_ext_upper} con "
        f"{_n_sheets} pestaña(s), convertido a tablas Markdown.\n"
        f"Analiza TODAS las pestañas para localizar las métricas financieras.\n"
        f"Cuando una métrica aparezca en varias hojas, usa la fuente más reciente o la más detallada.\n"
        f"Si los valores están en términos absolutos y tienes los ingresos, calcula los márgenes (%).\n"
        f"Para Revenue Growth YoY: (valor_último - valor_anterior) / |valor_anterior| × 100.\n\n"
        f"Aplica las siguientes instrucciones de extracción EXACTAMENTE:\n\n"
    )
    full_prompt = adapter_header + prompt_schema

    # Pass the prompt as first arg and the Markdown content as second —
    # analizar_texto calls generate_content([prompt, contenido_texto]).
    return gemini.analizar_texto(full_prompt, full_markdown)


def _process_docx(file_path: str, gemini, prompt_schema: str) -> str:
    """Extract text from DOCX and send to Gemini with the financial audit prompt."""
    print(f"📝 [DOCX] Extrayendo texto...")
    try:
        from docx import Document as DocxDocument
        doc = DocxDocument(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also extract tables
        table_texts = []
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    table_texts.append(row_text)
        full_text = "\n".join(paragraphs)
        if table_texts:
            full_text += "\n\nTABLAS DEL DOCUMENTO:\n" + "\n".join(table_texts)
    except ImportError:
        raise RuntimeError("python-docx no está instalado. Ejecuta: pip install python-docx")
    except Exception as e:
        print(f"❌ [DOCX] Error extrayendo texto: {e}")
        raise RuntimeError(f"No se pudo procesar el DOCX: {e}") from e

    print(f"   Texto extraído: {len(full_text)} caracteres, {len(paragraphs)} párrafos")

    docx_instruction = (
        "Eres un auditor financiero senior. Analiza el siguiente documento Word (DOCX) "
        "que contiene información financiera de una startup. Extrae las métricas del texto "
        "y tablas, y emite el JSON estándar.\n\nCONTENIDO DEL DOCUMENTO:\n"
    )
    return gemini.analizar_texto(docx_instruction + full_text + "\n\n" + prompt_schema, "")


# ── PDF helpers ───────────────────────────────────────────────────────────────

_PDF_CHUNK_SIZE = 90  # páginas máximas por llamada a Gemini

# Sections that MUST exist in the merged Gemini JSON for build_contract() to
# find all 16 KPIs.  If a section is absent (e.g. all chunks had null values)
# we guarantee an empty dict so downstream code never KeyErrors.
_REQUIRED_FM_SECTIONS = (
    "revenue_growth",
    "profit_margins",
    "cash_flow_indicators",
    "debt_ratios",
    "base_metrics",
    "sector_metrics",
)


def _ensure_dict(obj) -> dict:
    """
    Garantiza que el resultado recuperado de GCS sea un dict de Python.

    GCS puede almacenar JSONs con doble serialización (el string ya fue
    serializado una vez antes de ser guardado como string). Este helper
    maneja ambos casos sin excepciones silenciosas.

    Raises TypeError si el objeto no puede convertirse a dict.
    """
    if isinstance(obj, dict):
        return obj
    if isinstance(obj, str):
        parsed = json.loads(obj)
        # Doble serialización: json.loads devolvió otro string
        if isinstance(parsed, str):
            return json.loads(parsed)
        if isinstance(parsed, dict):
            return parsed
        raise TypeError(f"_ensure_dict: json.loads devolvió {type(parsed).__name__}, no dict")
    raise TypeError(f"_ensure_dict: esperaba dict o str, recibió {type(obj).__name__}")


def _ensure_fm_sections(gemini_json: dict) -> dict:
    """
    Garantiza que financial_metrics_2025 tenga todas las sub-secciones
    requeridas para que build_contract() pueda iterar KPI_REGISTRY completo.

    Modifica el dict en-lugar y también lo devuelve (para encadenamiento).
    """
    fm = gemini_json.setdefault("financial_metrics_2025", {})
    for section in _REQUIRED_FM_SECTIONS:
        fm.setdefault(section, {})
    return gemini_json


def split_pdf_to_chunks(file_bytes: bytes, size: int = _PDF_CHUNK_SIZE) -> list[bytes]:
    """
    Divide los bytes de un PDF en fragmentos de máximo `size` páginas.

    Devuelve una lista de bytes — cada elemento es un PDF autónomo válido
    que puede guardarse en disco o enviarse directamente a Gemini.

    Parameters
    ----------
    file_bytes : bytes — contenido completo del PDF original.
    size       : int   — máximo de páginas por fragmento (default 50).

    Returns
    -------
    list[bytes] con 1..N fragmentos. Si el PDF tiene ≤ size páginas,
    la lista tiene exactamente un elemento (los bytes originales sin modificar).
    """
    try:
        import fitz
    except ImportError:
        raise RuntimeError("PyMuPDF no instalado. Ejecuta: pip install pymupdf")

    src_doc     = fitz.open(stream=file_bytes, filetype="pdf")
    total_pages = len(src_doc)

    if total_pages <= size:
        src_doc.close()
        return [file_bytes]

    n_chunks = (total_pages + size - 1) // size
    chunks: list[bytes] = []

    for idx in range(n_chunks):
        start = idx * size
        end   = min(start + size - 1, total_pages - 1)

        chunk_doc = fitz.open()
        chunk_doc.insert_pdf(src_doc, from_page=start, to_page=end)
        chunks.append(chunk_doc.tobytes())
        chunk_doc.close()

        _blk_num    = idx + 1
        _page_start = start + 1
        _page_end   = end + 1
        print(f"[Chunking] Bloque {_blk_num}/{n_chunks}: páginas {_page_start}–{_page_end}")

    src_doc.close()
    return chunks


def merge_consolidated_results(jsons: list[dict]) -> dict:
    """
    Une los resultados de múltiples llamadas a Gemini (chunks de PDF)
    en un único diccionario consolidado.

    Regla de consolidación por KPI:
    - Si un KPI aparece en varios chunks con value != null, se queda el de
      mayor confidence score.
    - Empate de confidence → gana el chunk anterior (índice más bajo),
      preservando el orden documental del PDF.
    - _document_context se toma siempre del primer chunk que lo contenga.
    - base_metrics y sector_metrics se garantizan presentes en el resultado
      (secciones vacías si ningún chunk extrajo datos para ellas).

    Parameters
    ----------
    jsons : list de dicts Gemini ya parseados — mínimo uno.

    Returns
    -------
    Un único dict Gemini con la misma estructura del schema Cometa-Vault.
    """
    import copy

    if not jsons:
        raise ValueError("merge_consolidated_results: lista vacía")
    if len(jsons) == 1:
        return _ensure_fm_sections(jsons[0])

    merged = copy.deepcopy(jsons[0])

    def _get(obj: dict, path: list[str]):
        cur = obj
        for k in path:
            if not isinstance(cur, dict):
                return None
            cur = cur.get(k)
        return cur

    def _set(obj: dict, path: list[str], value) -> None:
        cur = obj
        for k in path[:-1]:
            cur = cur.setdefault(k, {})
        cur[path[-1]] = value

    for chunk_json in jsons[1:]:
        for kpi_def in KPI_REGISTRY:
            path     = kpi_def["path"]
            existing = _get(merged, path)
            incoming = _get(chunk_json, path)

            if not isinstance(incoming, dict) or incoming.get("value") is None:
                continue  # este chunk no tiene dato para este KPI

            if not isinstance(existing, dict) or existing.get("value") is None:
                _set(merged, path, incoming)
                continue

            # Ambos tienen valor — queda el de mayor confidence
            existing_conf = float(existing.get("confidence") or 0.0)
            incoming_conf = float(incoming.get("confidence") or 0.0)
            if incoming_conf > existing_conf:
                _set(merged, path, incoming)

    # Garantizar secciones requeridas aunque todos los chunks tengan null
    _ensure_fm_sections(merged)
    return merged


# Alias de compatibilidad (el upload path lo llamaba merge_kpi_results en
# versiones anteriores — conservamos el nombre por si algún test lo importa)
merge_kpi_results = merge_consolidated_results


def _chunk_and_process_pdf(temp_path: str, gemini, prompt_config: str) -> str:
    """
    Motor de Chunking — lee el PDF desde disco, lo divide en bloques de 50
    páginas con split_pdf_to_chunks(), llama a Gemini por cada bloque y
    consolida con merge_consolidated_results().

    - PDFs ≤ 50 páginas: una sola llamada directa a Gemini (sin overhead).
    - PDFs > 50 páginas: N bloques secuenciales; si un bloque falla se omite
      y se continúa. Si TODOS fallan → RuntimeError.

    Returns
    -------
    String JSON listo para json.loads() — mismo contrato que
    gemini.extraer_y_auditar().
    """
    import re as _re

    with open(temp_path, "rb") as fh:
        file_bytes = fh.read()

    chunks = split_pdf_to_chunks(file_bytes, size=_PDF_CHUNK_SIZE)
    n_chunks = len(chunks)

    if n_chunks == 1:
        print(f"[Chunking] PDF ≤ {_PDF_CHUNK_SIZE} páginas — llamada directa a Gemini")
        return gemini.extraer_y_auditar(temp_path, prompt_config)

    print(f"[Chunking] {n_chunks} bloques de hasta {_PDF_CHUNK_SIZE} páginas")

    chunk_results: list[dict] = []
    for i, chunk_bytes in enumerate(chunks):
        chunk_path = f"{temp_path}_chunk{i}.pdf"
        try:
            with open(chunk_path, "wb") as cf:
                cf.write(chunk_bytes)

            _blk = i + 1
            print(f"[Chunking] Enviando bloque {_blk}/{n_chunks} a Gemini...")
            raw = gemini.extraer_y_auditar(chunk_path, prompt_config)

            if isinstance(raw, str):
                try:
                    parsed = json.loads(raw)
                except json.JSONDecodeError:
                    clean  = _re.sub(r'^```json\s*|\s*```$', '', raw.strip())
                    parsed = json.loads(clean)
            else:
                parsed = raw

            chunk_results.append(parsed)
            print(f"[Chunking] Bloque {_blk} OK")

        except Exception as chunk_err:
            _blk = i + 1
            print(f"[Chunking] Bloque {_blk} falló ({chunk_err}) — omitido")
        finally:
            if os.path.exists(chunk_path):
                os.remove(chunk_path)

    if not chunk_results:
        raise RuntimeError("[Chunking] Ningún bloque fue procesado por Gemini")

    merged = merge_consolidated_results(chunk_results)
    print(f"[Chunking] {len(chunk_results)}/{n_chunks} bloques mergeados exitosamente")
    return json.dumps(merged, ensure_ascii=False)


def check_hash_exists_in_gcs(bucket_name: str, file_hash: str) -> bool:
    """Verifica si existe un archivo con el mismo hash en custom_metadata"""
    try:
        storage_client = _get_storage_client()
        bucket = storage_client.bucket(bucket_name)
        
        # Listar blobs y buscar por metadata
        blobs = bucket.list_blobs()
        
        for blob in blobs:
            if blob.metadata and blob.metadata.get('file_hash') == file_hash:
                print(f"📋 [API] Hash duplicado encontrado en GCS: {file_hash}")
                return True
        
        return False
    except (DefaultCredentialsError, Forbidden, Unauthorized) as e:
        sa_path = _resolve_service_account_path()
        print(f"❌ [API] Error de credenciales/permisos en GCS: {e}")
        print(f"   GOOGLE_APPLICATION_CREDENTIALS={os.getenv('GOOGLE_APPLICATION_CREDENTIALS')}")
        print(f"   Service Account resuelto={sa_path}")
        raise RuntimeError("GCS_AUTH") from e
    except Exception as e:
        print(f"❌ [API] Error verificando hash en GCS: {e}")
        raise RuntimeError("GCS_ERROR") from e

def get_existing_result(bucket_name: str, file_hash: str) -> dict:
    """Obtiene el resultado JSON existente para un hash específico"""
    try:
        storage_client = _get_storage_client()
        bucket = storage_client.bucket(bucket_name)
        
        # Buscar en staging por hash
        blobs = bucket.list_blobs(prefix="staging/")
        
        for blob in blobs:
            if blob.metadata and blob.metadata.get('file_hash') == file_hash:
                # Descargar y retornar el resultado existente
                content = blob.download_as_text()
                result = json.loads(content)
                print(f"📋 [API] Resultado existente encontrado para hash: {file_hash}")
                return result
        
        return None
    except (DefaultCredentialsError, Forbidden, Unauthorized) as e:
        sa_path = _resolve_service_account_path()
        print(f"❌ [API] Error de credenciales/permisos obteniendo resultado: {e}")
        print(f"   GOOGLE_APPLICATION_CREDENTIALS={os.getenv('GOOGLE_APPLICATION_CREDENTIALS')}")
        print(f"   Service Account resuelto={sa_path}")
        raise RuntimeError("GCS_AUTH") from e
    except Exception as e:
        print(f"❌ [API] Error obteniendo resultado existente: {e}")
        raise RuntimeError("GCS_ERROR") from e

# ── LEGACY RETIRADO ──────────────────────────────────────────────────────────
# POST /upload fue reemplazado por POST /api/founder/process-document
# (Patrón Adaptador Unificado → BD_Cometa_Dev via BQDataService).
# Retirado en la desconexión del legacy. No eliminar la función: conserva el
# histórico de implementación y permite rollback comentando el raise.
@app.post("/upload", deprecated=True, include_in_schema=False)
@limiter.limit("20/minute")
async def upload_pdf(
    request: Request,
    file: UploadFile = File(...),
    founder_email: str = Header(None, description="Email del founder para identificación"),
    company_id: str = Header(None, description="Company ID para multi-tenancy"),
    token: dict = Depends(_require_auth),
):
    """[RETIRADO] Reemplazado por POST /api/founder/process-document."""
    raise HTTPException(
        status_code=410,
        detail=(
            "Este endpoint ha sido retirado. "
            "Usa POST /api/founder/process-document para subir documentos."
        ),
    )
    # ── DEBUG: log inmediato al primer byte de la request ─────────────────
    print("=" * 60)
    print(f"📥 [DEBUG] /upload HIT — conexión recibida")
    _filename     = getattr(file, "filename", "N/A")
    _content_type = getattr(file, "content_type", "N/A")
    print(f"   filename     : {_filename}")
    print(f"   content_type : {_content_type}")
    print(f"   founder_email: {founder_email!r}")
    print(f"   company_id   : {company_id!r}")
    print("=" * 60)
    # ─────────────────────────────────────────────────────────────────────

    try:
        # ── C5: Validar y sanitizar headers de entrada ──────────────────────
        founder_email = _validate_email_header(founder_email)
        company_id    = _validate_company_header(company_id)

        # 1. Validar extensión
        ALLOWED_EXTENSIONS = {'.pdf', '.csv', '.xlsx', '.xls', '.parquet', '.docx', '.doc'}
        file_ext = os.path.splitext(file.filename or "")[1].lower()
        if not file.filename or file_ext not in ALLOWED_EXTENSIONS:
            _allowed_str = ", ".join(sorted(ALLOWED_EXTENSIONS))
            raise HTTPException(
                status_code=400,
                detail=f"Formato no soportado. Permitidos: {_allowed_str}"
            )
        print(f"📁 [DEBUG] Extensión detectada: {file_ext}")

        # 4. Leer contenido (adelantado para C2 y C7 antes de procesar)
        file_content = await file.read()

        # ── C2: Límite de tamaño ────────────────────────────────────────────
        if len(file_content) > _MAX_FILE_BYTES:
            raise HTTPException(
                status_code=413,
                detail=f"Archivo supera el límite de {_MAX_FILE_MB} MB"
            )

        # ── C7: Validar magic bytes ─────────────────────────────────────────
        if not _validate_magic_bytes(file_content, file_ext):
            raise HTTPException(
                status_code=415,
                detail=f"El contenido binario no corresponde a un archivo {file_ext} válido"
            )

        # 2. Resolución de identidad — prioridad: filename > header/email
        # El nombre del archivo es la señal más fiable: el Founder sabe qué sube.
        # Solo usamos el dominio del email como último recurso.
        _fname_key, _ = detect_company_from_text(file.filename or "")
        if _fname_key != "unknown":
            # Filename contiene una empresa conocida → prioridad máxima
            company_domain = _fname_key
            print(f"🏢 [API] Empresa detectada del filename '{file.filename}': {company_domain}")
        elif company_id:
            company_domain = company_id
            print(f"🏢 [API] Empresa desde header company_id: {company_domain}")
        elif founder_email and "@" in founder_email:
            company_domain = founder_email.split("@")[-1]
            print(f"🏢 [API] Empresa inferida del email: {company_domain}")
        else:
            company_domain = "pending_detection"
            print(f"[API] No company_id en header — se identificará desde el PDF")

        # 3. Fallback si llegó vacío
        if not company_domain or company_domain == 'unknown':
            company_domain = 'pending_detection'
        
        # 4. Calcular hash (file_content ya fue leído antes de las validaciones C2/C7)
        file_hash = get_file_hash(file_content)
        
        print(f"📤 [API] Archivo recibido: {file.filename}")
        print(f"🔍 [API] Hash calculado: {file_hash}")
        print(f"👤 [API] Founder: {founder_email}")
        print(f"🏢 [API] Vault path: vault/{company_domain}/")
        
        # 5. Verificar si ya existe por hash en la bóveda específica
        try:
            # Primero verificar en la bóveda de la empresa
            vault_prefix = f"vault/{company_domain}/"
            storage_client = _get_storage_client()
            bucket = storage_client.bucket(GCS_OUTPUT_BUCKET)
            
            # Buscar en la bóveda específica de la empresa
            blobs = bucket.list_blobs(prefix=vault_prefix)
            
            for blob in blobs:
                if blob.metadata and blob.metadata.get('file_hash') == file_hash:
                    # Encontrar resultado existente en la bóveda de la empresa
                    content = blob.download_as_text()
                    # Fix-1: garantizar dict (GCS puede devolver doble-serialización)
                    result = _ensure_dict(json.loads(content))
                    print(f"📋 [API] Resultado duplicado encontrado en bóveda de {company_domain}: {file_hash}")
                    # OBS-04: recalcular checklist_status desde el JSON cacheado
                    try:
                        _dup_contract = build_contract(
                            gemini_json=result,
                            file_hash=file_hash,
                            company_id=company_domain,
                            founder_email=founder_email or "",
                            original_filename=file.filename,
                        )
                        _dup_bucket = COMPANY_BUCKET.get(company_domain, "UNKNOWN")
                        dup_checklist = build_checklist_status(_dup_contract["kpi_rows"], _dup_bucket)
                    except Exception as _ce:
                        print(f"[API] checklist recalc failed for duplicate ({_ce}) — omitting")
                        dup_checklist = None
                    return JSONResponse(
                        content={
                            "status": "success",
                            "message": "Documento reconocido en la bóveda de Cometa. Sincronizando métricas...",
                            "duplicate": True,
                            "result": result,
                            "file_hash": file_hash,
                            "company_domain": company_domain,
                            "checklist_status": dup_checklist,
                        },
                        status_code=200
                    )
            
            # Si no existe en la bóveda de la empresa, buscar en el bucket general
            if check_hash_exists_in_gcs(GCS_INPUT_BUCKET, file_hash):
                # Copiar resultado a la bóveda de la empresa
                existing_result_raw = get_existing_result(GCS_OUTPUT_BUCKET, file_hash)
                if existing_result_raw:
                    # Fix-1: garantizar dict antes de operar con él
                    existing_result = _ensure_dict(existing_result_raw)
                    # Copiar a la bóveda específica
                    vault_result_filename = f"{vault_prefix}{file_hash}_result.json"
                    vault_blob = bucket.blob(vault_result_filename)

                    vault_blob.metadata = {
                        'file_hash': file_hash,
                        'original_filename': existing_result.get('original_filename', 'unknown'),
                        'founder_email': founder_email,
                        'company_domain': company_domain,
                        'vault_path': vault_prefix,
                        'processed_at': datetime.now(timezone.utc).isoformat(),
                        'copied_from_general': True
                    }
                    
                    vault_blob.upload_from_string(
                        json.dumps(existing_result, indent=2),
                        content_type='application/json'
                    )
                    
                    print(f"📋 [API] Resultado copiado a bóveda de {company_domain}: {vault_result_filename}")
                    # OBS-04: recalcular checklist_status desde el JSON cacheado
                    try:
                        _dup_contract2 = build_contract(
                            gemini_json=existing_result,
                            file_hash=file_hash,
                            company_id=company_domain,
                            founder_email=founder_email or "",
                            original_filename=file.filename,
                        )
                        _dup_bucket2 = COMPANY_BUCKET.get(company_domain, "UNKNOWN")
                        dup_checklist2 = build_checklist_status(_dup_contract2["kpi_rows"], _dup_bucket2)
                    except Exception as _ce2:
                        print(f"[API] checklist recalc failed for duplicate-copy ({_ce2}) — omitting")
                        dup_checklist2 = None
                    return JSONResponse(
                        content={
                            "status": "success",
                            "message": "Documento reconocido en la bóveda de Cometa. Sincronizando métricas...",
                            "duplicate": True,
                            "result": existing_result,
                            "file_hash": file_hash,
                            "company_domain": company_domain,
                            "checklist_status": dup_checklist2,
                        },
                        status_code=200
                    )
        except RuntimeError as e:
            if str(e) == "GCS_AUTH":
                raise HTTPException(
                    status_code=500,
                    detail=(
                        "Error de autenticación/permisos con GCS. "
                        "Verifica GOOGLE_APPLICATION_CREDENTIALS y permisos del bucket."
                    ),
                )
            raise HTTPException(
                status_code=500,
                detail="Error de conexión/lectura con GCS durante deduplicación",
            )
        
        # 6. Si es nuevo, iniciar procesamiento asíncrono
        print(f"[API] Archivo nuevo ({file_ext}), iniciando procesamiento...")

        # Guardar temporalmente para procesamiento
        safe_filename = _sanitize_filename(file.filename)  # C6: sanitización segura
        temp_path = os.path.join('/tmp', f"{file_hash}_{safe_filename}")
        os.makedirs('/tmp', exist_ok=True)

        with open(temp_path, "wb") as temp_file:
            temp_file.write(file_content)

        # ── RAW GCS: persistir archivo original SIEMPRE, antes de cualquier procesamiento ──
        # Garantiza que el Analista pueda recuperar el PDF desde el KpiReviewPanel
        # aunque Gemini falle o no reconozca los KPIs.
        try:
            _raw_sc   = _get_storage_client()
            _raw_blob = _raw_sc.bucket(GCS_INPUT_BUCKET).blob(
                f"{company_domain}/{file_hash}_{safe_filename}"
            )
            _raw_blob.upload_from_filename(temp_path, content_type="application/octet-stream")
            print(
                f"[API][raw] Archivo guardado en gs://{GCS_INPUT_BUCKET}/"
                f"{company_domain}/{file_hash}_{safe_filename}"
            )
        except Exception as _raw_err:
            print(f"[API][raw] GCS raw upload non-fatal: {_raw_err}")

        # ── 7A. RAMA EXCEL/CSV — Motor de Mapeo (109 KPIs) ───────────────────────
        # Para archivos estructurados (Excel / CSV), omitimos Vertex AI y
        # ejecutamos el nuevo motor de mapeo directamente.  El resultado es un
        # preview JSON que el frontend usa para mostrar el modal de confirmación.
        # La escritura real en BQ ocurre en POST /api/founder/confirm-mapping.
        if file_ext in ('.xlsx', '.xls', '.csv'):
            try:
                _slug       = detect_company_from_text(company_domain)[0]
                _sector     = COMPANY_BUCKET.get(_slug, "ALL")
                _prev_cash  = get_prev_cash_from_bq(_slug, datetime.now(timezone.utc).date())

                mapping_result = map_uploaded_file(
                    file_path=temp_path,
                    sector=_sector,
                    prev_cash=_prev_cash,
                )

                _period = datetime.now(timezone.utc).date().replace(day=1)
                _meta   = IngestionMetadata(
                    company_name  = company_domain.upper(),
                    company_slug  = _slug,
                    period        = _period,
                    founder_email = founder_email or "",
                    sector        = _sector,
                    loaded_by     = token.get("email", "unknown"),
                    source_file_hint = file.filename,
                )
                preview = build_upload_preview(mapping_result, _meta)

                _load_id    = preview.get("load_id", "")
                _period_iso = _period.isoformat()

                # ── COMMITMENT GATE: evaluar umbral de calidad ────────────────
                _gate = {}
                try:
                    _gate = evaluate_commitment_gate(
                        company_slug   = _slug,
                        period_iso     = _period_iso,
                        sector         = _sector,
                        mapping_result = mapping_result,
                    )
                    print(
                        f"[API][gate] {_gate['counter']} KPIs  "
                        f"gate_passed={_gate['gate_passed']}  "
                        f"sector={_gate['sector']}"
                    )
                except Exception as _gate_err:
                    print(f"[API][gate] Gate eval non-fatal error: {_gate_err}")

                # ── DB: guardar SIEMPRE (acumulación multi-archivo) ───────────
                # Los KPIs se persisten aunque el gate no pase, para que un
                # segundo archivo pueda completar el reporte del mismo período.
                _rows_saved   = 0
                _audit_result = {}
                try:
                    _db_rows    = build_registry_rows(mapping_result, _meta, _load_id)
                    _rows_saved = save_registry_rows(_db_rows)
                    print(f"[API][local_db] {_rows_saved} filas guardadas (acumulacion)")

                    _audit_result = build_enriched_audit_response(
                        company_slug   = _slug,
                        mapping_result = mapping_result,
                        rows_saved     = _rows_saved,
                        load_id        = _load_id,
                        period_iso     = _period_iso,
                    )
                    _grid_found = sum(
                        1 for r in _audit_result.get("kpi_grid", [])
                        if r.get("status") == "FOUND"
                    )
                    print(
                        f"DEBUG: Enviando grilla acumulada al UI con "
                        f"{_grid_found}/109 KPIs encontrados "
                        f"(periodo={_period_iso!r} empresa={_slug!r})."
                    )
                except Exception as _db_err:
                    print(f"[API][local_db] DB write non-fatal error: {_db_err}")

                # ── COMMIT 109/109: Contrato Jero + BigQuery ─────────────────
                # Solo ejecuta cuando los 109 KPIs están completos.
                # Orden obligatorio: Jero contract primero (prueba de integridad),
                # luego BQ MERGE (persistencia cloud).
                _bq_dispatch   = None
                _jero_records  = None
                _jero_path     = None

                if _gate.get("gate_passed"):
                    # 1. Generar output_to_jero.json (escritura atómica)
                    try:
                        _jero_records, _jero_path = generate_jero_contract(
                            mapping_result = mapping_result,
                            metadata       = _meta,
                            load_id        = _load_id,
                            gate_passed    = True,
                        )
                        print(
                            f"[API][jero] Contrato generado: {len(_jero_records)} registros "
                            f"→ {_jero_path.name}"
                        )
                    except Exception as _jero_err:
                        print(f"[API][jero] Error generando contrato (non-fatal): {_jero_err}")

                    # 2. BigQuery MERGE diferido — NO se ejecuta aqui.
                    # El MERGE a fact_portfolio_kpis solo ocurre cuando un
                    # analista aprueba la carga via POST /api/analyst/confirm-gold.
                    # El contrato Jero + los archivos GCS son la fuente de verdad
                    # hasta que el analista lo valide.
                    print(
                        f"[API][BQ] MERGE diferido — pendiente aprobacion de analista "
                        f"(load_id={_load_id!r})"
                    )

                # ── MEDALLION RAW: archivar archivo original ──────────────
                _raw_gcs_uri = upload_raw_layer(
                    file_bytes  = file_content,
                    filename    = file.filename or safe_filename,
                    metadata    = _meta,
                )
                if _raw_gcs_uri:
                    print(f"[API][medallion] raw → {_raw_gcs_uri}")

                # ── MEDALLION STAGE: archivar resultado del mapper ────────
                try:
                    _stage_payload = {
                        "load_id":       _load_id,
                        "company_slug":  _slug,
                        "period":        _period_iso,
                        "coverage_pct":  getattr(mapping_result, "coverage_pct", None),
                        "found_count":   len(getattr(mapping_result, "found", [])),
                        "missing_count": len(getattr(mapping_result, "missing_kpis", [])),
                    }
                    _stage_uri = upload_stage_layer(
                        gemini_json = _stage_payload,
                        metadata    = _meta,
                        load_id     = _load_id,
                    )
                    if _stage_uri:
                        print(f"[API][medallion] stage → {_stage_uri}")
                except Exception as _st_err:
                    print(f"[API][medallion] stage upload non-fatal: {_st_err}")

                # ── MEDALLION GOLD: solo cuando gate pasa ─────────────────
                if _gate.get("gate_passed") and _jero_records:
                    try:
                        _gold_uri = upload_gold_layer(
                            contract_json = {
                                "load_id":    _load_id,
                                "company":    _slug,
                                "period":     _period_iso,
                                "records":    _jero_records,
                                "gate":       _gate,
                            },
                            metadata  = _meta,
                            load_id   = _load_id,
                        )
                        if _gold_uri:
                            print(f"[API][medallion] gold → {_gold_uri}")
                    except Exception as _gld_err:
                        print(f"[API][medallion] gold upload non-fatal: {_gld_err}")

                # ── GCS: guardar pending para confirm-mapping ────────────────
                _pending_key = f"pending_mapper/{_slug}/{_load_id}_{safe_filename}"
                try:
                    _sc   = _get_storage_client()
                    _blob = _sc.bucket(GCS_OUTPUT_BUCKET).blob(_pending_key)
                    _blob.upload_from_filename(temp_path, content_type="application/octet-stream")
                    _blob.metadata = {"company_slug": _slug, "sector": _sector,
                                      "founder_email": founder_email or ""}
                    _blob.patch()
                except Exception as _gcs_err:
                    print(f"[API][mapper] GCS pending save failed (non-fatal): {_gcs_err}")

                # ── STATUS: "committed" si gate pasa, "pending_kpis" si no ───
                _status = "committed" if _gate.get("gate_passed") else "pending_kpis"
                print(f"[API][gate] status={_status!r}  msg={_gate.get('message','')[:80]}")

                _bq_summary = None
                if _bq_dispatch is not None:
                    _bq_summary = {
                        "rows_inserted": _bq_dispatch.rows_inserted,
                        "rows_updated":  _bq_dispatch.rows_skipped_dup,
                        "rows_error":    _bq_dispatch.rows_error,
                        "warnings":      _bq_dispatch.warnings,
                    }

                _jero_summary = None
                if _jero_records is not None:
                    _null_count    = sum(1 for r in _jero_records if r["value"] is None)
                    _filled_count  = len(_jero_records) - _null_count
                    _crit_missing  = [
                        r["kpi_id"] for r in _jero_records
                        if r["value"] is None and r["innegociable"]
                    ]
                    _jero_summary = {
                        "generated":       True,
                        "total_kpis":      len(_jero_records),
                        "filled":          _filled_count,
                        "nulls":           _null_count,
                        "critical_missing": _crit_missing,
                        "output_file":     str(_jero_path.name) if _jero_path else None,
                    }

                return JSONResponse(content={
                    "status":          _status,
                    "flow":            "mapper",
                    "file_ext":        file_ext,
                    "commitment_gate": _gate,
                    "preview":         preview,
                    "audit":           _audit_result,
                    "bq_commit":       _bq_summary,
                    "jero_contract":   _jero_summary,
                }, status_code=200)

            except SubmissionBlockedError as sbe:
                raise HTTPException(status_code=422, detail=[
                    {"rule_id": f.rule_id, "severity": f.severity, "msg": f.message}
                    for f in sbe.blocking_flags
                ])
            except Exception as mapper_err:
                print(f"[API][mapper] Error en motor de mapeo: {mapper_err}")
                raise HTTPException(status_code=500,
                    detail=f"Error en motor de mapeo: {mapper_err}")

        # ── 7B. RAMA PDF — Vertex AI / Gemini (flujo original) ───────────────
        # 7. Iniciar procesamiento con Vertex AI
        try:
            # Inicializar adaptadores
            doc_ai = DocumentAIAdapter(PROJECT_ID, LOCATION_DOC_AI, PROCESSOR_ID)
            print(f"DEBUG: El valor de VERTEX_LOCATION es: '{VERTEX_LOCATION}'")
            gemini = GeminiAuditor(PROJECT_ID, VERTEX_LOCATION)
            
            # ── Contexto de vertical (inyectado antes del prompt) ─────────────
            # El bucket_id se conoce en cuanto detectamos la empresa; pasarlo
            # al prompt evita que Gemini busque métricas de otras verticales.
            _bucket_id = COMPANY_BUCKET.get(company_domain, "UNKNOWN")
            _sector_hints: dict[str, str] = {
                "SAAS":  "Prioriza MRR, Churn Rate y CAC. Si el documento menciona 'Monthly Recurring Revenue', 'Churn' o 'Customer Acquisition Cost', extráelos en sector_metrics.",
                "LEND":  "Prioriza Portfolio Size y NPL Ratio. Si el documento menciona 'cartera de crédito', 'morosidad', 'NPL' o 'Non-Performing Loans', extráelos en sector_metrics.",
                "ECOM":  "Prioriza GMV. Si el documento menciona 'Gross Merchandise Value', 'Total Sales Volume', 'Total Transaction Value' o 'GMV', extráelo en sector_metrics.gmv.",
                "INSUR": "Prioriza Loss Ratio. Si el documento menciona 'siniestralidad', 'claims ratio' o 'loss ratio', extráelo en sector_metrics.",
                "OTH":   "Extrae las métricas financieras estándar. No hay métricas sectoriales obligatorias.",
            }
            _sector_instruction = _sector_hints.get(
                _bucket_id,
                "Extrae todas las métricas del esquema. No hay contexto sectorial específico disponible.",
            )

            # ── Prompt CoT + esquema dinámico 109 KPIs ───────────────────────
            # El prefijo dinámico (f-string con bucket_id y sector_instruction)
            # se concatena con el cuerpo que contiene el esquema JSON de los 89
            # KPIs GIVEN generado desde loading_brain_v1.json en tiempo real.
            _prompt_prefix = (
                f"Eres un auditor financiero senior especializado en due diligence de startups.\n"
                f"Tu misión es extraer métricas financieras de un PDF con precisión institucional.\n"
                f"\n"
                f"╔══════════════════════════════════════════════════════════════╗\n"
                f"║  CONTEXTO DE INDUSTRIA                                      ║\n"
                f"╚══════════════════════════════════════════════════════════════╝\n"
                f"\n"
                f"Estás analizando una empresa de la vertical {_bucket_id}.\n"
                f"{_sector_instruction}\n"
                f"No intentes inventar o extraer métricas de otras verticales a menos que el\n"
                f"documento las mencione explícitamente como KPIs de negocio de la empresa.\n"
                f"\n"
                f"╔══════════════════════════════════════════════════════════════╗\n"
                f"║  SINÓNIMOS FINANCIEROS ACEPTADOS                            ║\n"
                f"╚══════════════════════════════════════════════════════════════╝\n"
                f"\n"
                f"Reconoce estas equivalencias aunque el documento use el término alternativo:\n"
                f"  • revenue / ingresos / ventas netas / ingresos de actividades ordinarias\n"
                f"    / total net revenue / net sales / turnover\n"
                f"  • ebitda / utilidad operativa ajustada / resultado operacional\n"
                f"  • gross profit / utilidad bruta / margen bruto\n"
                f"  • cash in bank / efectivo y equivalentes / cash & equivalents\n"
                f"  • mrr / monthly recurring revenue / ingresos recurrentes mensuales\n"
                f"  • arr / annual recurring revenue / ingresos recurrentes anuales\n"
                f"  • churn / tasa de cancelación / abandono de clientes\n"
                f"  • cac / costo de adquisición / customer acquisition cost\n"
                f"  • gmv / volumen bruto de mercancía / gross merchandise value\n"
                f"  • npl / cartera vencida / non-performing loans / morosidad\n"
                f"Si ves cualquiera de estos términos, extrae el valor aunque el nombre exacto\n"
                f"no coincida con el campo del esquema. Prioriza extraer sobre dejar en null.\n"
                f"\n"
            )
            # Generar el esquema dinámico de los 89 KPIs GIVEN desde el brain.
            # _build_gemini_kpi_schema() retorna el bloque JSON de extracción.
            _kpi_schema = _build_gemini_kpi_schema()

            _prompt_body = (
                "\n"
                "╔══════════════════════════════════════════════════════════════╗\n"
                "║  FASE 1 — ANÁLISIS PREVIO  (escribe en _document_context)  ║\n"
                "╚══════════════════════════════════════════════════════════════╝\n"
                "\n"
                "Ejecuta este análisis y materializa las conclusiones en el campo\n"
                "`_document_context` del JSON de salida (ver esquema más abajo).\n"
                "\n"
                "  A. MONEDA PRINCIPAL\n"
                "     ¿Cuál es la moneda del documento? (USD, MXN, EUR, COP, etc.)\n"
                "     Si hay más de una moneda, ¿cuál domina los estados financieros?\n"
                "\n"
                "  B. PERÍODO DE REPORTE\n"
                "     ¿Qué año fiscal o período cubre el documento?\n"
                "     Busca encabezados como 'FY2025', 'Año terminado el 31/12/2025',\n"
                "     'H1 2025', 'Q4 2025'. Anota si el documento cubre períodos parciales.\n"
                "\n"
                "  C. ESCALA NUMÉRICA\n"
                "     ¿Los montos están en unidades base, miles ($K) o millones ($M)?\n"
                "     Busca notas como 'en miles de pesos' o 'amounts in USD thousands'.\n"
                "     Normaliza TODOS los valores usando K, M o B según corresponda.\n"
                "\n"
                "  D. ZONAS DE AMBIGÜEDAD\n"
                "     Identifica métricas donde el dato no aparece explícito, hay dos\n"
                "     cifras posibles, o el documento es un deck sin estados auditados.\n"
                "     Para cada zona asigna confidence < 0.70 y explica en description.\n"
                "\n"
                "Las conclusiones de A, B, C van en `_document_context`.\n"
                "Las zonas de D informan los campos description y confidence de cada métrica.\n"
                "\n"
                "╔══════════════════════════════════════════════════════════════╗\n"
                "║  FASE 2 — EXTRACCIÓN JSON  (109 KPIs del Contrato Cometa)  ║\n"
                "╚══════════════════════════════════════════════════════════════╝\n"
                "\n"
                "REGLAS OBLIGATORIAS:\n"
                "1. Responde ÚNICAMENTE con el objeto JSON. Cero caracteres fuera del JSON.\n"
                "2. Usa EXACTAMENTE las claves del esquema. Sin sinónimos, sin claves extra.\n"
                "3. Cada métrica tiene tres campos obligatorios:\n"
                "\n"
                "   \"value\"       — número con unidad normalizada. Ej: '36%', '$9.7M', '-$320K'.\n"
                "                    NUNCA omitas la unidad. NUNCA escribas un número puro.\n"
                "\n"
                "   \"confidence\"  — Float 0.0–1.0:\n"
                "                    >= 0.90  dato explícito, fuente directa.\n"
                "                    0.70–0.89 cálculo menor o inferencia razonable.\n"
                "                    < 0.70  ambiguo, estimado, parcial o fuente indirecta.\n"
                "                    Sé honesto: subestimar es mejor que inflar confianza.\n"
                "\n"
                "   \"description\" — Cita exacta: tabla, línea y página del documento.\n"
                "                    Si confidence < 0.70: explica por qué el dato es incierto.\n"
                "\n"
                "4. Si una métrica no aparece en el documento: escribe null.\n"
                "5. La clave raíz de las métricas es SIEMPRE \"financial_metrics_2025\".\n"
                "6. Normaliza la escala (FASE 1-C) antes de escribir cada value.\n"
                "7. El campo `_document_context` es OBLIGATORIO con sus 4 sub-campos.\n"
                "   currency usa código ISO 4217 (USD, MXN, EUR, BRL, COP, ARS...).\n"
                "8. Las métricas marcadas ★ INNEGOCIABLE son críticas — prioriza su extracción.\n"
                "\n"
                "ESQUEMA REQUERIDO (89 KPIs GIVEN del Contrato Cometa VC):\n"
                + _kpi_schema + "\n"
                "\n"
                "Analiza el documento adjunto y responde con el JSON completo. Nada más.\n"
            )
            prompt_config = _prompt_prefix + _prompt_body

            # ── Enrutar por tipo de archivo ────────────────────────────────
            if file_ext in ('.csv', '.xlsx', '.xls', '.parquet'):
                resultado_raw = _process_tabular(temp_path, file_ext, gemini, prompt_config)
            elif file_ext in ('.docx', '.doc'):
                resultado_raw = _process_docx(temp_path, gemini, prompt_config)
            else:
                # PDF — Smart Chunking (bloques de 50 páginas)
                resultado_raw = _chunk_and_process_pdf(temp_path, gemini, prompt_config)

            # Normalizar: Gemini devuelve un string JSON; lo parseamos para evitar
            # doble serialización al guardarlo en GCS.
            if isinstance(resultado_raw, str):
                try:
                    resultado = json.loads(resultado_raw)
                except json.JSONDecodeError:
                    # Si el modelo devolvió texto envuelto en ```json ... ```, limpiarlo
                    import re
                    clean = re.sub(r'^```json\s*|\s*```$', '', resultado_raw.strip())
                    resultado = json.loads(clean)
            else:
                resultado = resultado_raw
            
            # 8a. Resolución de identidad — prioridad: filename > contenido PDF > email domain
            #
            # El nombre del archivo es la señal más confiable: el Founder sabe qué empresa
            # está subiendo. El contenido del PDF puede mezclar referencias a otras empresas.
            # El email domain es el último recurso (puede pertenecer al portafolio equivocado).
            _filename_key, _filename_portfolio = detect_company_from_text(file.filename or "")
            _content_key,  _content_portfolio  = detect_company_from_text(json.dumps(resultado))

            if _filename_key != "unknown":
                detected_key       = _filename_key
                detected_portfolio = _filename_portfolio
                print(
                    f"[API] Empresa detectada del NOMBRE DE ARCHIVO: "
                    f"'{detected_key}' -> Fondo {detected_portfolio}"
                )
            elif _content_key != "unknown":
                detected_key       = _content_key
                detected_portfolio = _content_portfolio
                print(
                    f"[API] Empresa detectada del CONTENIDO PDF: "
                    f"'{detected_key}' -> Fondo {detected_portfolio}"
                )
            else:
                detected_key       = "unknown"
                detected_portfolio = _lookup_portfolio(company_domain)
                print(
                    f"[API] Empresa no detectada del PDF — usando header/email: "
                    f"'{company_domain}' -> Fondo {detected_portfolio}"
                )

            if detected_key != "unknown":
                company_domain = detected_key
                portfolio_id   = detected_portfolio
            else:
                portfolio_id = detected_portfolio

            # 8b. Build the canonical data contract (Rule 4 + Rule 8)
            contract = build_contract(
                gemini_json=resultado,
                file_hash=file_hash,
                company_id=company_domain,
                founder_email=founder_email or "",
                original_filename=file.filename,
                portfolio_id=portfolio_id,
            )

            # ── R1 + R2: Normalización de IDs canónicos ──────────────────────────
            # Muta el contrato in-place ANTES de leer ningún campo:
            #   • period_id  →  PYYYYQxMyy  (ej. "P2025Q1M03")
            #   • company_id →  COMP_XXX    (ej. "COMP_SOLVENTO")
            # También propaga company_id a cada kpi_row (requerido por el contrato JSON).
            _raw_period = resultado.get("_document_context", {}).get("period", "") or ""
            _norm_result = _apply_contract_normalization(
                contract=contract,
                raw_company=company_domain,
                raw_period=_raw_period,
            )
            # Sincronizar company_domain local con la clave canónica lowercase para GCS.
            # IMPORTANTE: el contrato ya tiene comp_id ("COMP_SOLVENTO") mutado para BQ.
            # Las rutas GCS deben usar la clave lowercase ("solvento") para coincidir
            # con lo que /api/portfolio-companies devuelve al sidebar del analista.
            if _norm_result["company_ok"]:
                # Empresa conocida: derivar clave lowercase desde comp_id (strip "COMP_")
                company_domain = _norm_result["comp_id"].replace("COMP_", "").lower()
            else:
                # Empresa desconocida: normalizar dominio (strip TLD, lowercase)
                company_domain = company_domain.lower().split(".")[0].replace("-", "").replace("_", "")
            portfolio_id = _norm_result["fund_id"] or portfolio_id

            print(
                f"🔖 [R1/R2] Normalization — "
                f"period: '{_norm_result['period_id']}' (ok={_norm_result['period_ok']}) | "
                f"company: '{_norm_result['comp_id']}' (known={_norm_result['company_ok']}) | "
                f"fund: '{_norm_result['fund_id']}' | bucket: '{_norm_result['bucket_id']}'"
            )
            if _norm_result["errors"]:
                for _err in _norm_result["errors"]:
                    print(f"⚠️  [R1/R2] {_err}")
            # ─────────────────────────────────────────────────────────────────────

            integrity = contract["integrity"]
            _sub        = contract["submission"]
            _kpi_valid  = _sub["kpi_count_valid"]
            _kpi_total  = _sub["kpi_count_total"]
            _period_id  = _sub["period_id"]   # ya normalizado a PYYYYQxMyy
            _period_ok  = integrity["period_consistent"]
            print(
                f"📋 [API] Contract built — "
                f"valid KPIs: {_kpi_valid}/{_kpi_total}, "
                f"period: {_period_id}, "
                f"period_consistent: {_period_ok}"
            )

            if integrity["warnings"]:
                for w in integrity["warnings"]:
                    print(f"⚠️  [API] Integrity warning: {w}")

            # 8c. Sector checklist — usa bucket_id resuelto por R2
            company_bucket = _norm_result["bucket_id"] or COMPANY_BUCKET.get(company_domain, "UNKNOWN")
            checklist_status = build_checklist_status(contract["kpi_rows"], company_bucket)
            # Enrich checklist with per-KPI confidence scores so the frontend can
            # highlight low-confidence fields before the founder manually corrects them.
            _conf_scores = _extract_kpi_confidence_scores(resultado)
            if _conf_scores:
                checklist_status["confidence_scores"] = _conf_scores
            if not checklist_status["is_complete"]:
                _missing_kpis = checklist_status["missing_critical_kpis"]
                print(f"[API] Checklist incompleto ({company_bucket}): faltan {_missing_kpis}")

            # ── SQLite: guardar KPIs de Gemini SIEMPRE (acumulación multi-archivo)
            # Debe ocurrir ANTES de build_accumulated_kpi_grid para que la grilla
            # incluya los hallazgos del archivo actual junto con los previos.
            _pdf_period     = datetime.now(timezone.utc).date().replace(day=1)
            _pdf_period_iso = _pdf_period.isoformat()
            _pdf_load_id    = file_hash[:16]
            _pdf_slug       = company_domain.lower().replace(".", "_")
            _pdf_rows_saved = 0
            try:
                _pdf_db_rows = build_registry_rows_from_contract(
                    kpi_rows     = contract["kpi_rows"],
                    company_slug = _pdf_slug,
                    company_name = company_domain,
                    period_iso   = _pdf_period_iso,
                    load_id      = _pdf_load_id,
                    loaded_by    = token.get("email", "api"),
                )
                _pdf_rows_saved = save_registry_rows(_pdf_db_rows)
                print(
                    f"[API][local_db][pdf] {_pdf_rows_saved} filas guardadas en SQLite "
                    f"(company={_pdf_slug!r} period={_pdf_period_iso!r})"
                )
            except Exception as _pdf_db_err:
                print(f"[API][local_db][pdf] SQLite write non-fatal: {_pdf_db_err}")

            # ── Grilla acumulada 109 KPIs: SQLite (previo) + Gemini (actual) ────
            # Se construye DESPUÉS del save para que la grilla refleje el estado
            # completo del período incluyendo los hallazgos del archivo actual.
            # El Founder ve su progreso real: todos los archivos juntos.
            from src.core.local_db import _CATALOG_BY_METRIC_ID as _CAT_IDX
            _pdf_current_kpis: dict = {}
            for _cr in contract["kpi_rows"]:
                if _cr.get("numeric_value") is not None and _cr.get("is_valid", False):
                    _cat = _CAT_IDX.get(_cr["kpi_key"])
                    if _cat:
                        _pdf_current_kpis[_cat["kpi_id"]] = {
                            "value":         _cr["numeric_value"],
                            "unit":          _cat.get("unit") or _cr.get("unit"),
                            "match_type":    "EXACT",
                            "audit_status":  "GOLD",
                            "quality_score": 1.0,
                            "source":        "gemini",
                        }

            _pdf_kpi_grid    = build_accumulated_kpi_grid(
                company_slug = _pdf_slug,
                period_iso   = _pdf_period_iso,
                current_kpis = _pdf_current_kpis,
            )
            _pdf_found_count = sum(1 for r in _pdf_kpi_grid if r["status"] == "FOUND")
            _pdf_miss_count  = len(_pdf_kpi_grid) - _pdf_found_count
            _pdf_audit_block = {
                "kpi_grid": _pdf_kpi_grid,
                "kpi_grid_summary": {
                    "total":        len(_pdf_kpi_grid),
                    "found":        _pdf_found_count,
                    "missing":      _pdf_miss_count,
                    "source":       "gemini+sqlite",
                    "rows_saved":   _pdf_rows_saved,
                },
            }
            print(
                f"DEBUG: Enviando grilla acumulada al UI con "
                f"{_pdf_found_count}/109 KPIs encontrados "
                f"(periodo={_pdf_period_iso!r} empresa={_pdf_slug!r})."
            )
            # ────────────────────────────────────────────────────────────────────

            # ── Validación de contenido financiero (permisiva) ──────────────────
            # Si Gemini no extrajo KPIs reconocibles, NO se bloquea la carga.
            # El archivo ya está en GCS raw; el Analista puede hacer el mapeo manual
            # desde el KpiReviewPanel. Se marca needs_manual_mapping=True para que
            # el frontend muestre el aviso correspondiente.
            needs_manual_mapping = False
            if not _is_financial_document(resultado):
                print(
                    f"⚠️  [API] Gemini no detectó KPIs financieros reconocibles — "
                    f"se activa mapeo manual. Hash: {file_hash}, archivo: {file.filename}"
                )
                needs_manual_mapping = True
            # ────────────────────────────────────────────────────────────────────

            # 8d. BQ write — BLOQUEADO para FOUNDER (Zero-Trust Gatekeeper).
            # El Analista es el unico que puede certificar datos hacia produccion.
            # FOUNDER solo escribe en GCS vault/ + SQLite.
            # La certificacion ocurre via POST /api/analyst/review-pdf → finalize-analysis.
            db_result = {"inserted": False, "duplicate": False, "submission_id": None,
                         "deferred": False}
            _uploader_role = token.get("role", "FOUNDER")
            if _uploader_role == "ANALISTA":
                # Analistas pueden subir PDFs y certificarlos directamente (caso especial)
                try:
                    db_result = insert_contract(contract)
                    if db_result["duplicate"]:
                        print(f"[API] BQ dedup — submission already exists for hash {file_hash}")
                except Exception as db_err:
                    print(f"[API] BigQuery write failed (non-fatal, GCS copy kept): {db_err}")
            else:
                # FOUNDER: data queda en GCS vault/ + SQLite, pendiente de analista
                db_result["deferred"] = True
                print(
                    f"[API][PDF][FOUNDER] BQ write diferido — hash={file_hash[:16]}… "
                    f"pendiente revision de analista via /api/analyst/review-pdf"
                )

            # 8d. Save to GCS vault bajo el company_domain ya detectado
            storage_client = _get_storage_client()
            bucket = storage_client.bucket(GCS_OUTPUT_BUCKET)
            vault_prefix = f"vault/{company_domain}/"
            result_filename = f"{vault_prefix}{file_hash}_result.json"
            blob = bucket.blob(result_filename)
            blob.metadata = {
                'file_hash': file_hash,
                'original_filename': file.filename,
                'founder_email': founder_email,
                'company_domain': company_domain,
                'portfolio_id': portfolio_id,
                'vault_path': vault_prefix,
                'processed_at': datetime.now(timezone.utc).isoformat()
            }
            # Vínculo permanente: inyectar source_file con la URI raw ANTES de guardar.
            # Esto permite que el analista siempre pueda recuperar el original.
            _raw_uri = f"gs://{GCS_INPUT_BUCKET}/{company_domain}/{file_hash}_{safe_filename}"
            resultado.setdefault("source_file", _raw_uri)

            blob.upload_from_string(
                json.dumps(resultado, indent=2, ensure_ascii=False),
                content_type='application/json'
            )

            print(f"✅ [API] Resultado guardado en GCS: {result_filename}")

            # GCS Medallion Stage bypassed in v2.0 — BD_Cometa_Dev es la única fuente de verdad.
            log.info("[API][stage] GCS Stage upload bypassed in v2.0 (Star Schema activo)")

            # Guardar archivo original en GCS (vault/{company}/raw/)
            raw_filename = f"vault/{company_domain}/raw/{file_hash}_{safe_filename}"
            raw_blob = bucket.blob(raw_filename)
            raw_blob.metadata = {
                'file_hash': file_hash,
                'original_filename': file.filename,
                'founder_email': founder_email,
                'company_domain': company_domain,
                'file_type': file_ext,
                'uploaded_at': pd.Timestamp.now().isoformat(),
            }
            raw_blob.upload_from_string(file_content, content_type=file.content_type or 'application/octet-stream')
            print(f"📦 [API] Archivo original guardado en GCS: {raw_filename}")

            # 8e. Receipt email — Vault Seal SHA-256 (non-fatal)
            try:
                from src.services.hash_service import generate_vault_seal
                from src.services.email_service import send_receipt_email
                _processed_at = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M")
                _vault_seal = generate_vault_seal(
                    company_id   = company_domain,
                    file_hash    = file_hash,
                    kpi_rows     = contract["kpi_rows"],
                    processed_at = _processed_at,
                )
                send_receipt_email(
                    to_email       = founder_email,
                    company_domain = company_domain,
                    period         = _period_id,
                    vault_seal     = _vault_seal,
                    file_hash      = file_hash,
                    kpi_count      = _kpi_valid,
                    processed_at   = _processed_at,
                )
            except Exception as _receipt_err:
                print(f"[API] Receipt email failed (non-fatal): {_receipt_err}")

            # 9. Limpiar archivo temporal
            os.remove(temp_path)

            # 10. Retornar contrato completo al frontend
            kpi_confidence_scores = _extract_kpi_confidence_scores(resultado)

            # ── Sector-gate validator (mandatory fields + sanity ranges) ─────────
            # Runs AFTER GCS/BQ writes — the document is never lost regardless of
            # the outcome.  Returns 400 so the frontend can render the rescue form.
            _blocked = validate_founder_submission(contract["kpi_rows"], company_bucket)
            if _blocked:
                _missing  = _blocked["missing_mandatory_fields"]
                _sanity   = _blocked["sanity_violations"]
                _miss_names = ", ".join(f['kpi_key'] for f in _missing)
                print(
                    f"⚠️  [GATE] Datos parciales detectados. "
                    f"Cambiando a modo de carga manual para completar los 109 KPIs. "
                    f"Faltan: {_miss_names or 'ver sanity_violations'}"
                )
                # ── 200 OK — el frontend necesita los datos para pintar el
                # formulario de inputs manuales y el contador de KPIs.
                # Un 400 rompe la comunicación con la interfaz.
                return JSONResponse(
                    content={
                        "status":                   "pending_kpis",
                        "commitment_gate": {
                            "counter":        f"{_pdf_found_count}/109",
                            "gate_passed":    _pdf_found_count >= 109,
                            "coverage_pct":   round(_pdf_found_count / 109 * 100, 1),
                            "missing_required": [
                                {"kpi_id": f["kpi_key"], "display_name": f["kpi_key"],
                                 "severity": "CRITICAL"}
                                for f in _missing
                            ],
                            "ui_hint": (
                                f"GATE BLOQUEADO — {_pdf_found_count}/109 KPIs acumulados "
                                f"({_pdf_miss_count} faltantes). "
                                f"KPIs criticos pendientes: {_miss_names}. "
                                "Sube otro archivo o completa manualmente para desbloquear."
                            ),
                        },
                        "missing_mandatory_fields": _missing,
                        "sanity_violations":        _sanity,
                        "file_hash":                file_hash,
                        "company_domain":           company_domain,
                        "checklist_status":         checklist_status,
                        "kpi_confidence_scores":    kpi_confidence_scores,
                        "audit":                    _pdf_audit_block,
                    },
                    status_code=200,
                )
            # ─────────────────────────────────────────────────────────────────────

            return JSONResponse(
                content={
                    "status": "success",
                    "message": (
                        "Archivo recibido. Requiere mapeo manual del Analista."
                        if needs_manual_mapping
                        else "Archivo procesado exitosamente"
                    ),
                    "duplicate": False,
                    "needs_manual_mapping": needs_manual_mapping,
                    # Legacy field: raw Gemini dict for existing frontend consumers
                    "result": resultado,
                    # Structured contract for The Vault
                    "submission":        contract["submission"],
                    "kpi_rows":          contract["kpi_rows"],
                    "integrity":         integrity,
                    "db":                db_result,
                    "file_hash":         file_hash,
                    "company_domain":    company_domain,
                    # Sector checklist — feedback inmediato al founder
                    "checklist_status":  checklist_status,
                    # Per-KPI confidence scores extracted from Gemini (0–100 integer)
                    "kpi_confidence_scores": kpi_confidence_scores,
                    # 109-KPI progress grid — siempre presente para el frontend
                    "audit":             _pdf_audit_block,
                },
                status_code=200
            )
            
        except HTTPException:
            # HTTPException ya tiene status_code correcto — propagar sin envolver
            if os.path.exists(temp_path):
                os.remove(temp_path)
            raise
        except Exception as e:
            # Limpiar archivo temporal
            if os.path.exists(temp_path):
                os.remove(temp_path)

            print(f"⚠️  [API] Error en procesamiento KPI — archivo aceptado para revisión manual: {e}")
            traceback.print_exc()

            # 202 Accepted: el PDF ya está en GCS raw; el Analista puede mapearlo.
            return JSONResponse(
                content={
                    "status": "accepted",
                    "message": (
                        "El archivo fue recibido pero el procesamiento automático falló. "
                        "Queda disponible para revisión manual del Analista."
                    ),
                    "needs_manual_mapping": True,
                    "file_hash": file_hash if "file_hash" in dir() else None,
                    "company_domain": company_domain if "company_domain" in dir() else None,
                    "error_detail": str(e),
                },
                status_code=202,
            )
    
    except Exception as e:
        print(f"❌ [API] Error general en upload: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error en el servidor: {str(e)}")


# ═══════════════════════════════════════════════════════════════════════════════
# MAPPER CONFIRM ENDPOINT
# Segundo paso del flujo Excel/CSV: escritura real en fact_portfolio_history.
# El preview (paso 1) fue generado por POST /upload y guardó el archivo
# en GCS bajo pending_mapper/{load_id}_{filename}.
# ═══════════════════════════════════════════════════════════════════════════════

class ConfirmMappingRequest(BaseModel):
    """Body para POST /api/founder/confirm-mapping."""
    load_id:        str           # UUID del batch del preview (del paso 1)
    filename:       str           # nombre del archivo original (para localizar en GCS)
    company_slug:   str           # ej. "simetrik"
    company_name:   str           # ej. "SIMETRIK"
    period_str:     str           # "YYYY-MM" — primer dia del mes reportado
    sector:         str           # "ALL" | "SAAS_SUBSCRIPTION" | etc.
    source_type:    str = "verified"


# ── LEGACY RETIRADO ──────────────────────────────────────────────────────────
# POST /api/founder/confirm-mapping implementaba el flujo Excel de 2 pasos
# (preview → confirm). Reemplazado por el flujo single-step de
# POST /api/founder/process-document.
@app.post("/api/founder/confirm-mapping", deprecated=True, include_in_schema=False)
@limiter.limit("10/minute")
async def founder_confirm_mapping(
    request: Request,
    body: ConfirmMappingRequest,
    token: dict = Depends(_require_auth),
):
    """[RETIRADO] Reemplazado por POST /api/founder/process-document."""
    raise HTTPException(
        status_code=410,
        detail=(
            "Este endpoint ha sido retirado. "
            "Usa POST /api/founder/process-document para subir archivos Excel/CSV."
        ),
    )
    import re as _re
    from datetime import date as _date

    # Validar period_str formato YYYY-MM
    if not _re.match(r"^\d{4}-\d{2}$", body.period_str):
        raise HTTPException(status_code=400,
            detail="period_str debe tener formato YYYY-MM (ej. 2025-03)")

    year_s, month_s = body.period_str.split("-")
    period = _date(int(year_s), int(month_s), 1)

    # Si el caller es un Founder, forzar su propio company_slug (Zero Trust)
    _role = token.get("role", "")
    if _role == "FOUNDER":
        _email  = token.get("email", "")
        _domain = _email.split("@")[-1] if "@" in _email else ""
        _forced_slug = detect_company_from_text(_domain)[0]
        if _forced_slug != "unknown" and _forced_slug != body.company_slug:
            raise HTTPException(status_code=403,
                detail="Founders solo pueden confirmar cargas de su propia empresa.")

    # Recuperar archivo temporal de GCS
    _safe   = _sanitize_filename(body.filename)
    _gcs_key = f"pending_mapper/{body.company_slug}/{body.load_id}_{_safe}"
    _local_tmp = os.path.join('/tmp', f"confirm_{body.load_id}_{_safe}")
    try:
        _sc   = _get_storage_client()
        _blob = _sc.bucket(GCS_OUTPUT_BUCKET).blob(_gcs_key)
        _blob.download_to_filename(_local_tmp)
    except Exception as dl_err:
        raise HTTPException(status_code=404,
            detail=f"Archivo de mapeo no encontrado en GCS ({_gcs_key}): {dl_err}")

    try:
        _prev_cash = get_prev_cash_from_bq(body.company_slug, period)

        mapping_result = map_uploaded_file(
            file_path=_local_tmp,
            sector=body.sector,
            prev_cash=_prev_cash,
        )

        meta = IngestionMetadata(
            company_name     = body.company_name.upper(),
            company_slug     = body.company_slug.lower(),
            period           = period,
            founder_email    = token.get("email", ""),
            sector           = body.sector,
            loaded_by        = token.get("email", "pipeline@cometa.vc"),
        )

        dispatch = dispatch_to_storage(
            result      = mapping_result,
            metadata    = meta,
            source_type = body.source_type,
            load_id     = body.load_id,   # mismo load_id del preview
        )

        # Limpiar GCS pending file (best-effort)
        try:
            _blob.delete()
        except Exception:
            pass

        return JSONResponse(content={
            "status":          "committed",
            "load_id":         dispatch.load_id,
            "rows_inserted":   dispatch.rows_inserted,
            "rows_updated":    dispatch.rows_skipped_dup,
            "rows_error":      dispatch.rows_error,
            "quality_summary": dispatch.quality_summary,
            "warnings":        dispatch.warnings,
        }, status_code=200)

    except SubmissionBlockedError as sbe:
        raise HTTPException(status_code=422, detail=[
            {"rule_id": f.rule_id, "severity": f.severity, "msg": f.message}
            for f in sbe.blocking_flags
        ])
    except Exception as confirm_err:
        print(f"[API][confirm-mapping] Error: {confirm_err}")
        raise HTTPException(status_code=500,
            detail=f"Error al confirmar la carga: {confirm_err}")
    finally:
        try:
            os.unlink(_local_tmp)
        except Exception:
            pass


class ManualUpdateRequest(BaseModel):
    """Body for POST /api/founder/manual-update."""
    file_hash: str
    updates:   dict[str, str]


# ── LEGACY RETIRADO ──────────────────────────────────────────────────────────
# POST /api/founder/manual-update actualizaba blobs JSON en el vault de GCS.
# El nuevo flujo unificado escribe en BD_Cometa_Dev; las correcciones del
# founder se incluyen en el campo manual_kpis de POST /api/founder/finalize.
@app.post("/api/founder/manual-update", deprecated=True, include_in_schema=False)
@limiter.limit("30/minute")
async def founder_manual_update(
    request: Request,
    body: ManualUpdateRequest,
    token: dict = Depends(_require_auth),
) -> JSONResponse:
    """[RETIRADO] Las correcciones van en manual_kpis de POST /api/founder/finalize."""
    raise HTTPException(
        status_code=410,
        detail=(
            "Este endpoint ha sido retirado. "
            "Incluye las correcciones en el campo manual_kpis de "
            "POST /api/founder/finalize."
        ),
    )
    company_domain: str = token.get("company_id") or token.get("sub", "")
    # Normalise to domain-only (strip full email if necessary)
    if "@" in company_domain:
        company_domain = company_domain.split("@")[-1]

    if not company_domain:
        raise HTTPException(status_code=403, detail="company_id no disponible en el token")

    vault_prefix = f"vault/{company_domain}/"

    try:
        storage_client = _get_storage_client()
        bucket_obj     = storage_client.bucket(GCS_OUTPUT_BUCKET)
        blobs          = list(bucket_obj.list_blobs(prefix=vault_prefix))
    except Exception as gcs_err:
        raise HTTPException(
            status_code=500,
            detail=f"Error al conectar con GCS: {gcs_err}",
        )

    # Find the blob for this file_hash
    target_blob = None
    for blob in blobs:
        if blob.name.endswith(".json") and body.file_hash in blob.name:
            target_blob = blob
            break

    if target_blob is None:
        raise HTTPException(
            status_code=404,
            detail=f"No se encontró resultado para el hash '{body.file_hash}' en la bóveda de {company_domain}",
        )

    try:
        existing_raw  = target_blob.download_as_text()
        existing_data = _ensure_dict(json.loads(existing_raw))
    except Exception as read_err:
        raise HTTPException(
            status_code=500,
            detail=f"Error al leer el resultado existente: {read_err}",
        )

    # Apply corrections — inject into the manual_corrections sub-object
    manual_corrections: dict = existing_data.get("manual_corrections") or {}
    for k, v in body.updates.items():
        manual_corrections[k] = v
    existing_data["manual_corrections"] = manual_corrections

    try:
        target_blob.upload_from_string(
            json.dumps(existing_data, indent=2, ensure_ascii=False),
            content_type="application/json",
        )
    except Exception as write_err:
        raise HTTPException(
            status_code=500,
            detail=f"Error al guardar correcciones en GCS: {write_err}",
        )

    updated_fields = list(body.updates.keys())
    print(
        f"[manual-update] {len(updated_fields)} campo(s) corregidos para "
        f"hash={body.file_hash} company={company_domain}: {updated_fields}"
    )
    return JSONResponse(
        content={"status": "ok", "updated_fields": updated_fields},
        status_code=200,
    )


# ── LEGACY RETIRADO ──────────────────────────────────────────────────────────
# GET /api/result/{file_hash} leía resultados del vault GCS por hash.
# Los resultados ahora viven en BD_Cometa_Dev (BigQuery), no en GCS.
@app.get("/api/result/{file_hash}", deprecated=True, include_in_schema=False)
async def get_analysis_result(file_hash: str):
    """[RETIRADO] Los resultados ya viven en BD_Cometa_Dev, no en GCS."""
    raise HTTPException(
        status_code=410,
        detail=(
            "Este endpoint ha sido retirado. "
            "Los resultados de análisis ya no se almacenan en GCS vault."
        ),
    )

@app.get("/api/results")
async def get_all_results(company_id: str = None, token: dict = Depends(_require_auth)):
    """
    Obtiene todos los resultados de análisis guardados en GCS/vault/{company_id}/

    Multi-tenancy rules
    -------------------
    ANALISTA (ANA-*)  — can query any company_id supplied in the URL.
    FOUNDER  (FND-*)  — acceso denegado: los datos del founder se consultan
                        a través de /api/founder/config y BD_Cometa_Dev.
    """
    # ── Rol gate: Founders no acceden al vault GCS legacy ─────────────────────
    _role    = (token.get("role") or "").upper()
    _user_id = (token.get("user_id") or "")
    if _role == "FOUNDER" or _user_id.startswith("FND-"):
        raise HTTPException(
            status_code=403,
            detail="Founders acceden a sus datos via /api/founder/config",
        )

    # ── Multi-tenancy gate ─────────────────────────────────────────────────────
    role    = (token.get("role") or "").upper()
    user_id = (token.get("user_id") or "")

    is_founder = (role == "FOUNDER") or user_id.startswith("FND-")
    if is_founder:
        # Derive company from JWT — never trust the URL param for Founders.
        # token.get("company_id") covers tokens minted with an explicit claim;
        # falling back to "sub" (email) and stripping the domain works for
        # standard Founder accounts whose email matches their company domain.
        raw_company: str = (token.get("company_id") or token.get("sub") or "").strip()
        if "@" in raw_company:
            raw_company = raw_company.split("@")[-1]   # john@solvento.com → solvento.com
        jwt_company = raw_company.lower()
        if not jwt_company:
            raise HTTPException(
                status_code=403,
                detail="Founder sin company_id en el token. Contacta a tu analista.",
            )
        # Override whatever the caller sent — scope is enforced server-side
        company_id = jwt_company

    try:
        storage_client = _get_storage_client()
        bucket = storage_client.bucket(GCS_OUTPUT_BUCKET)

        # Si no se proporciona company_id, devolver error
        if not company_id:
            raise HTTPException(status_code=400, detail="company_id es obligatorio")

        # Buscar en la ruta canónica y en la ruta legada COMP_XXX para compatibilidad
        # con documentos subidos antes del fix de normalización de company_domain.
        cid_clean = company_id.lower().strip()
        vault_prefixes = [
            f"vault/{cid_clean}/",                      # ruta canónica: vault/solvento/
            f"vault/COMP_{cid_clean.upper()}/",         # ruta legada:   vault/COMP_SOLVENTO/
        ]

        results = []
        seen_ids: set[str] = set()  # evitar duplicados si por alguna razón coinciden

        for vault_prefix in vault_prefixes:
            for blob in bucket.list_blobs(prefix=vault_prefix):
                if not blob.name.endswith('.json'):
                    continue
                # Deduplicate by the computed id (hash-based), not the full blob path.
                # Two blobs from canonical and legacy prefixes can share the same derived
                # id if the same file exists in both; track that to avoid React key errors.
                blob_id = blob.name.replace('.json', '').replace(vault_prefix, '')
                if blob_id in seen_ids:
                    continue
                seen_ids.add(blob_id)
                try:
                    # Descargar contenido del JSON
                    content = blob.download_as_text()
                    
                    # Manejar posible doble serialización
                    try:
                        result_data = json.loads(content)
                        # Si el resultado sigue siendo un string, aplicar segundo parse
                        if isinstance(result_data, str):
                            print(f" [API] Detectada doble serialización en {blob.name}")
                            result_data = json.loads(result_data)
                    except json.JSONDecodeError as e:
                        print(f" [API] Error parseando JSON de {blob.name}: {e}")
                        continue
                    
                    # Asegurarse que result_data sea un diccionario
                    if not isinstance(result_data, dict):
                        _blob_name      = blob.name
                        _result_type    = type(result_data).__name__
                        print(f" [API] Resultado no es diccionario en {_blob_name}: {_result_type}")
                        continue
                    
                    # Extraer metadata
                    metadata = blob.metadata or {}
                    
                    # Construir objeto de resultado con estructura clara
                    # Derive portfolio — prefer stored metadata, fall back to lookup
                    blob_portfolio = metadata.get('portfolio_id') or _lookup_portfolio(
                        metadata.get('company_domain', company_id)
                    )
                    _company_domain = metadata.get('company_domain', company_id)
                    result_item = {
                        "id":          blob.name.replace('.json', '').replace(vault_prefix, ''),
                        "company_id":  cid_clean,
                        "slug":        _company_domain.lower() if _company_domain else cid_clean,
                        "data":        result_data,
                        "date":        metadata.get('processed_at', 'unknown'),
                        "metadata": {
                            "original_filename": metadata.get('original_filename', 'unknown'),
                            "founder_email":     metadata.get('founder_email', 'unknown'),
                            "file_hash":         metadata.get('file_hash', ''),
                            "processed_at":      metadata.get('processed_at', 'unknown'),
                            "gcs_path":          blob.name,
                            "company_domain":    _company_domain,
                            "portfolio_id":      blob_portfolio,
                        }
                    }

                    results.append(result_item)
                    _blob_name   = blob.name
                    _data_type   = type(result_data).__name__
                    _data_keys   = list(result_data.keys()) if isinstance(result_data, dict) else "N/A"
                    print(f" [API] Resultado cargado: {_blob_name}")
                    print(f" [API] Tipo de data: {_data_type}")
                    print(f" [API] Keys en data: {_data_keys}")
                    
                except Exception as e:
                    print(f" [API] Error procesando {blob.name}: {e}")
                    continue
        
        # Ordenar por fecha de procesamiento (más reciente primero)
        results.sort(key=lambda x: x['date'], reverse=True)

        # ── BigQuery fallback — carga datos históricos cuando GCS está vacío ──
        # Los registros legacy/missing_legacy no tienen archivos en GCS vault/;
        # solo existen en fact_kpi_values. Los sintetizamos en el mismo formato
        # financial_metrics_2025 que el frontend ya sabe consumir.
        if len(results) == 0:
            results = _build_results_from_bq(cid_clean)

        print(f" [API] Resultados encontrados para {company_id}: {len(results)}")

        return JSONResponse(
            content={
                "status": "success",
                "results": results,
                "company_id": company_id,
                "total": len(results)
            }
        )
    except (DefaultCredentialsError, Forbidden, Unauthorized) as e:
        sa_path = _resolve_service_account_path()
        print(f" [API] Error de credenciales/permisos obteniendo resultados: {e}")
        print(f"   GOOGLE_APPLICATION_CREDENTIALS={os.getenv('GOOGLE_APPLICATION_CREDENTIALS')}")
        print(f"   Service Account resuelto={sa_path}")
        raise HTTPException(
            status_code=500,
            detail="Error de autenticación/permisos con GCS. Verifica credenciales."
        )
    except Exception as e:
        print(f"[API] Error general en get_all_results: {e}")
        raise HTTPException(status_code=500, detail=f"Error obteniendo resultados: {str(e)}")

@app.get("/health")
async def health_check():
    """Endpoint de verificación de salud"""
    return {"status": "healthy", "service": "cometa-pipeline-api"}


# ── Analyst KPI correction ────────────────────────────────────────────────────

class KpiUpdateRequest(BaseModel):
    """
    Body for PUT /api/kpi-update.

    Fields
    ------
    submission_id : UUID of the submission in BigQuery (fact_kpi_values.submission_id).
    metric_id     : KPI key as stored in fact_kpi_values.kpi_key
                    (e.g. "revenue_growth", "ebitda_margin").
    value         : New raw string value typed by the Analyst
                    (e.g. "42%", "$8.5M"). Passed through parse_numeric.
    """
    submission_id: str
    metric_id:     str
    value:         str


@app.put("/api/kpi-update")
async def kpi_update(payload: KpiUpdateRequest, token: dict = Depends(_require_auth)):
    """
    Persist an Analyst correction to a KPI row in BigQuery.

    - Validates the new value via parse_numeric (Rule 4).
    - Sets is_manually_edited=TRUE and edited_at=CURRENT_TIMESTAMP().
    - Preserves the original value in edited_raw_value for audit.

    Returns the updated KPI data including parse result.
    """
    try:
        new_val = float(payload.value) if payload.value else None
        _bq_svc.update_single_kpi(
            submission_id=payload.submission_id,
            metric_id=payload.metric_id,
            new_value=new_val,
        )
        return JSONResponse(
            content={
                "status":  "success",
                "message": f"KPI '{payload.metric_id}' actualizado correctamente",
            }
        )

    except ValueError as e:
        raise HTTPException(status_code=404, detail=str(e))

    except Exception as e:
        log.error("[API] Error en /api/kpi-update: %s", e)
        raise HTTPException(
            status_code=500,
            detail=f"Error actualizando KPI en BigQuery: {str(e)}"
        )


# ── PATCH /api/kpi/update — Analyst correction with fidelity upgrade ─────────
#
# Diferencias clave respecto al PUT /api/kpi-update:
#   1. Solo accesible por ANALISTA (ANA-* / role=ANALISTA) — 403 para FOUNDER.
#   2. Al ser editado manualmente, fact_kpi_values marca is_manually_edited=TRUE,
#      lo que eleva value_status a 'verified' en las queries. Esto hace que
#      el Banner de Fidelidad desaparezca conforme el analista limpia la data
#      periodo a periodo.

@app.patch("/api/kpi/update")
async def kpi_patch_update(
    payload: KpiUpdateRequest,
    token: dict = Depends(_require_auth),
):
    """
    Corrige el valor de un KPI en BigQuery y lo marca como 'verified'.

    - Requiere rol ANALISTA (403 si es FOUNDER o SOCIO).
    - Llama a _bq_svc.update_single_kpi() que hace DML UPDATE en
      BD_Cometa_Dev.fact_kpi_values (Star Schema v2.0).

    Body
    ----
    submission_id : str  — UUID de la submision en BigQuery.
    metric_id     : str  — kpi_key (e.g. "revenue_growth").
    value         : str  — nuevo valor corregido (e.g. "42%", "$8.5M").
    """
    role    = (token.get("role") or "").upper()
    user_id = (token.get("user_id") or "")

    is_analista = (role == "ANALISTA") or user_id.startswith("ANA-")
    if not is_analista:
        raise HTTPException(
            status_code=403,
            detail="Solo los analistas pueden corregir KPIs.",
        )

    try:
        new_val = float(payload.value) if payload.value else None
        _bq_svc.update_single_kpi(
            submission_id=payload.submission_id,
            metric_id=payload.metric_id,
            new_value=new_val,
        )
        return JSONResponse(
            content={
                "status":       "success",
                "message":      f"KPI '{payload.metric_id}' corregido — fidelidad actualizada",
                "value_status": "verified",
                **result,
            }
        )

    except ValueError as e:
        raise HTTPException(status_code=404, detail=str(e))

    except Exception as e:
        print(f"❌ [API] Error en PATCH /api/kpi/update: {e}")
        traceback.print_exc()
        raise HTTPException(
            status_code=500,
            detail=f"Error corrigiendo KPI en BigQuery: {str(e)}",
        )


# ── Analyst batch-edit with audit hash ───────────────────────────────────────

class AnalystEditRequest(BaseModel):
    """Body for POST /api/analyst/audit-edit."""
    submission_id: str                   # file_hash / submission_id in BigQuery
    updates:       dict[str, str]        # { kpi_key: new_raw_value }
    note:          str = ""              # Optional edit justification note


@app.post("/api/analyst/audit-edit")
@limiter.limit("30/minute")
async def analyst_audit_edit(
    request: Request,
    body: AnalystEditRequest,
    token: dict = Depends(_require_auth),
) -> JSONResponse:
    """
    Batch-edit KPI values for a submission and return an audit hash.

    Only accessible to ANALISTA role.  For each entry in ``body.updates``,
    calls ``_bq_svc.update_single_kpi()`` (DML UPDATE en BD_Cometa_Dev.fact_kpi_values).  After all updates a SHA-256 vault seal
    is generated covering: submission_id + sorted kpi_keys + analyst email +
    timestamp.  This hash is the "recibo de edición" the analyst can attach
    to the case.

    Returns
    -------
    JSON ``{ status, audit_hash, updated_kpis, failed_kpis, submission_id }``
    """
    from src.services.hash_service import generate_vault_seal

    # ── A1: ANALISTA-only gate ────────────────────────────────────────────────
    role    = token.get("role", "")
    user_id = token.get("user_id", "")
    if role != "ANALISTA" and not user_id.startswith("ANA-"):
        raise HTTPException(status_code=403, detail="Solo analistas pueden editar en batch.")

    analyst_email: str = (token.get("email") or token.get("sub", "")).strip()

    if not body.updates:
        raise HTTPException(status_code=422, detail="updates no puede estar vacío.")

    processed_at = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M")
    updated_kpis: list[str] = []
    failed_kpis:  list[dict] = []

    for kpi_key, raw_value in body.updates.items():
        raw_value = (raw_value or "").strip()
        if not raw_value:
            continue
        try:
            new_val = float(raw_value) if raw_value else None
            _bq_svc.update_single_kpi(
                submission_id=body.submission_id,
                metric_id=kpi_key,
                new_value=new_val,
            )
            updated_kpis.append(kpi_key)
        except Exception as _err:
            failed_kpis.append({"kpi_key": kpi_key, "error": str(_err)})
            log.warning("[analyst/audit-edit] %s: %s", kpi_key, _err)

    if not updated_kpis:
        raise HTTPException(
            status_code=422,
            detail=f"No se pudo actualizar ningún KPI. Detalles: {failed_kpis}",
        )

    # ── Generate audit hash ───────────────────────────────────────────────────
    audit_hash = generate_vault_seal(
        company_id   = analyst_email,
        file_hash    = body.submission_id,
        kpi_rows     = [
            {"kpi_key": k, "raw_value": body.updates[k], "unit": "", "is_valid": True}
            for k in sorted(updated_kpis)
        ],
        processed_at = processed_at,
    )

    print(
        f"[analyst/audit-edit] analyst={analyst_email!r}  "
        f"submission={body.submission_id[:12]}…  "
        f"updated={len(updated_kpis)}  failed={len(failed_kpis)}  "
        f"hash={audit_hash[:16]}…"
    )

    return JSONResponse(content={
        "status":       "ok",
        "audit_hash":   audit_hash,
        "updated_kpis": updated_kpis,
        "failed_kpis":  failed_kpis,
        "submission_id": body.submission_id,
        "processed_at": processed_at,
    })


# ── Coverage heatmap — ANALISTA only ─────────────────────────────────────────

@app.get("/api/analyst/coverage")
async def analyst_coverage(
    token:        dict = Depends(_require_auth),
    portfolio_id: str  = Query(default=""),
):
    """
    GET /api/analyst/coverage — Portfolio KPI coverage matrix.

    Returns per-company × per-period KPI verification status for the heatmap
    component.  Restricted to ANALISTA role.

    Response shape
    --------------
    {
        "status":    "ok",
        "companies": [{"key": str, "display": str, "portfolio_id": str}],
        "periods":   [str],      # canonical PYYYYQxMyy, chronological
        "cells":     [
            {
                "company":        str,
                "period":         str,
                "status":         "verified" | "legacy" | "missing",
                "kpi_count":      int,
                "verified_count": int,
                "legacy_count":   int
            }
        ]
    }
    """
    print("\n🚀 [API] Petición de Cobertura Entrando...", flush=True)

    role    = (token.get("role") or "").upper()
    user_id = (token.get("user_id") or "")

    is_analista = (role == "ANALISTA") or user_id.startswith("ANA-")
    if not is_analista:
        raise HTTPException(
            status_code=403,
            detail="Acceso restringido a analistas Cometa.",
        )

    try:
        result = _bq_svc.get_portfolio_coverage(
            fund_id=portfolio_id if portfolio_id else None,
        )
        print(f"✅ [API] Coverage OK — {len(result.get('companies', []))} empresas, {len(result.get('cells', []))} celdas", flush=True)
        return result
    except Exception as exc:
        print(f"❌ [API] Coverage ERROR — {exc}", flush=True)
        log.error("analyst_coverage: error al construir cobertura — %s", exc)
        raise HTTPException(
            status_code=500,
            detail=f"Error al calcular la cobertura del portafolio: {exc}",
        ) from exc


# ── Analyst Gold Approval — DISPARADOR del MERGE a BigQuery ──────────────────
#
# Este endpoint es el UNICO punto de entrada al MERGE de fact_portfolio_kpis.
# Flujo:
#   1. Lee el archivo pendiente de GCS stage (pending_mapper/{load_id}_{filename}).
#   2. Re-ejecuta el motor de mapeo (idempotente, < 300 ms).
#   3. Sube el contrato a GCS gold/ (certificacion permanente).
#   4. Ejecuta dispatch_to_storage() con confirmed_by=analyst_email.
#   5. Elimina el archivo pendiente de stage.

class ConfirmGoldRequest(BaseModel):
    """Body para POST /api/analyst/confirm-gold (v2.0 — Star Schema).

    Recibe el submission_id canónico de BD_Cometa_Dev y opcionalmente
    un comentario del analista. Actualiza status → 'VALIDATED' via DML.
    """
    submission_id: str              # ID de la submission a validar (ej. "S1A2B3")
    review_notes:  str | None = None  # Comentario opcional del analista


@app.post("/api/analyst/confirm-gold")
@limiter.limit("20/minute")
async def analyst_confirm_gold(
    request: Request,
    body: ConfirmGoldRequest,
    token: dict = Depends(_require_analyst_auth),
) -> JSONResponse:
    """
    Valida una submission en BD_Cometa_Dev cambiando su status a 'VALIDATED'.

    Solo accesible para el rol ANALISTA con dominio @cometa.vc.
    En el Star Schema v2.0, la validación es un UPDATE sobre submissions —
    no hay movimiento de archivos GCS ni MERGE a fact_portfolio_kpis.

    Returns
    -------
    JSON {status, submission_id, validated_by}
    """
    analyst_email: str = (token.get("email") or token.get("sub", "")).strip()

    if not body.submission_id or len(body.submission_id) < 3:
        raise HTTPException(status_code=400, detail="submission_id inválido.")

    try:
        _bq_svc.update_submission_status(
            submission_id=body.submission_id,
            status="VALIDATED",
            review_notes=body.review_notes or f"Validated by {analyst_email}",
        )
        print(
            f"[analyst/confirm-gold] analyst={analyst_email!r} "
            f"submission={body.submission_id!r} → VALIDATED"
        )
        return JSONResponse(content={
            "status":        "committed",
            "submission_id": body.submission_id,
            "validated_by":  analyst_email,
        }, status_code=200)

    except Exception as err:
        print(f"[analyst/confirm-gold] Error: {err}")
        raise HTTPException(status_code=500,
            detail=f"Error al validar la submission: {err}")


# ── Manual KPI entry (Analista Auditoría tab) ────────────────────────────────

_KPI_ENTRY_FIELDS = frozenset({
    "revenue_growth", "gross_profit_margin", "ebitda_margin",
    "cash_in_bank_end_of_year", "annual_cash_flow", "working_capital_debt",
    "revenue", "ebitda", "cogs",
    "mrr", "churn_rate", "cac", "portfolio_size", "npl_ratio", "gmv", "loss_ratio",
})


class ManualEntryRequest(BaseModel):
    """Body for POST /api/manual-entry — Analyst enters KPIs without a PDF.

    Supports all 16 KPIs of the current data contract:
      - 6 core financial KPIs
      - 3 base metrics (inputs for the derivation engine)
      - 7 sector metrics (sector-specific KPIs)

    Empty strings from the frontend are coerced to None before validation.
    Numeric values sent as float/int are coerced to str for parse_numeric.
    """
    # ── Identifiers ────────────────────────────────────────────────────────
    company_id:               str
    portfolio_id:             str
    period_id:                str  = "FY2025"
    founder_email:            str  = ""
    submission_id:            str | None = None   # links manual entry to original upload
    # ── Core financial KPIs ────────────────────────────────────────────────
    revenue_growth:           str | None = None
    gross_profit_margin:      str | None = None
    ebitda_margin:            str | None = None
    cash_in_bank_end_of_year: str | None = None
    annual_cash_flow:         str | None = None
    working_capital_debt:     str | None = None
    # ── Base metrics (derivation engine inputs) ────────────────────────────
    revenue:                  str | None = None
    ebitda:                   str | None = None
    cogs:                     str | None = None
    # ── Sector metrics ─────────────────────────────────────────────────────
    mrr:                      str | None = None   # SAAS
    churn_rate:               str | None = None   # SAAS
    cac:                      str | None = None   # SAAS / ECOM / INSUR
    portfolio_size:           str | None = None   # LEND
    npl_ratio:                str | None = None   # LEND
    gmv:                      str | None = None   # ECOM
    loss_ratio:               str | None = None   # INSUR

    @model_validator(mode="before")
    @classmethod
    def _coerce_kpi_fields(cls, data: dict) -> dict:
        """Convert empty strings → None and numbers → str for all KPI fields."""
        if not isinstance(data, dict):
            return data
        for field in _KPI_ENTRY_FIELDS:
            val = data.get(field)
            if val is None:
                continue
            if isinstance(val, (int, float)):
                # Frontend sent a bare number — convert to string for parse_numeric
                data[field] = str(val)
            elif isinstance(val, str) and val.strip() == "":
                # Empty string — treat as not provided
                data[field] = None
        return data


def _manual_node(value: str | None) -> dict:
    return {
        "value":       value,
        "confidence":  1.0,
        "description": "Entrada manual del Analista",
    }


@app.post("/api/manual-entry")
async def manual_entry(payload: ManualEntryRequest, token: dict = Depends(_require_auth)):
    """
    Persist analyst-entered KPIs directly to BigQuery without a PDF.
    Converts string KPI values with parse_numeric(), maps to canonical metric_ids,
    and writes via BQDataService.insert_submission_and_facts().
    """
    from datetime import date as _date
    import re as _re

    # ── Resolve canonical company_id from BQ ─────────────────────────────
    bq_company_id, bq_fund_id, _, is_known = get_company_id(payload.company_id)
    if not is_known:
        # Fall back to payload values if company not in BQ (analyst override)
        bq_company_id = payload.company_id
        bq_fund_id    = payload.portfolio_id

    # ── Derive period_start (partition key) from period_id string ─────────
    year_match = _re.search(r"(\d{4})", payload.period_id)
    period_year = int(year_match.group(1)) if year_match else datetime.now(timezone.utc).year
    period_start = _date(period_year, 1, 1)

    # ── Field → metric_id mapping (cash_in_bank_end_of_year → "cash") ────
    _FIELD_TO_METRIC: dict[str, str] = {
        "cash_in_bank_end_of_year": "cash",
        "revenue_growth":           "revenue_growth",
        "gross_profit_margin":      "gross_profit_margin",
        "ebitda_margin":            "ebitda_margin",
        "annual_cash_flow":         "annual_cash_flow",
        "working_capital_debt":     "working_capital_debt",
        "revenue":                  "revenue",
        "ebitda":                   "ebitda",
        "cogs":                     "cogs",
        "mrr":                      "mrr",
        "churn_rate":               "churn_rate",
        "cac":                      "cac",
        "portfolio_size":           "portfolio_size",
        "npl_ratio":                "npl_ratio",
        "gmv":                      "gmv",
        "loss_ratio":               "loss_ratio",
    }

    # ── Build kpi_rows — skip fields where parse_numeric returns None ─────
    kpi_rows: list[dict] = []
    for field, metric_id in _FIELD_TO_METRIC.items():
        raw_val = getattr(payload, field, None)
        if raw_val is None:
            continue
        numeric, unit = parse_numeric(raw_val)
        if numeric is None:
            continue
        kpi_rows.append({
            "metric_id":   metric_id,
            "value":       numeric,
            "value_notes": f"Entrada manual del Analista{f' ({unit})' if unit else ''}",
        })

    if not kpi_rows:
        raise HTTPException(
            status_code=422,
            detail="Ningún campo KPI contiene un valor numérico válido.",
        )

    try:
        result = _bq_svc.insert_submission_and_facts(
            company_id   = bq_company_id,
            fund_id      = bq_fund_id,
            period_id    = payload.period_id,
            period_start = period_start,
            submitted_by = payload.founder_email or token.get("email", "analista"),
            source_file  = f"[manual] {payload.company_id} {payload.period_id}",
            kpi_rows     = kpi_rows,
        )
        return JSONResponse(content={
            "status":       "success",
            "message":      f"Datos de {bq_company_id} guardados correctamente",
            "submission_id": result["submission_id"],
            "rows_inserted": result["rows_inserted"],
            "timestamp":     result["timestamp"],
        })
    except Exception as e:
        print(f"[API] Error en /api/manual-entry: {e}")
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))


# ── Duplicate audit — delete a vault submission ───────────────────────────────

@app.delete("/api/submission")
async def delete_submission(file_hash: str, company_id: str, token: dict = Depends(_require_auth)):
    """
    Delete a specific submission from the GCS vault.
    Identifies the blob by matching file_hash in metadata.
    BigQuery row is NOT deleted (preserves audit trail).
    """
    try:
        storage_client = _get_storage_client()
        bucket         = storage_client.bucket(GCS_OUTPUT_BUCKET)
        vault_prefix   = f"vault/{company_id}/"
        blobs          = list(bucket.list_blobs(prefix=vault_prefix))
        deleted        = 0
        for blob in blobs:
            if blob.metadata and blob.metadata.get("file_hash") == file_hash:
                blob.delete()
                deleted += 1
                print(f"[API] Deleted vault blob: {blob.name}")
        if deleted == 0:
            raise HTTPException(status_code=404, detail=f"No blob found for hash {file_hash} in {company_id}")
        return JSONResponse(content={"status": "success", "deleted": deleted, "file_hash": file_hash})
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# ── Portfolio registry endpoints ──────────────────────────────────────────────

@app.get("/api/kpi-metadata")
async def get_kpi_metadata(vertical: str | None = None):
    """
    Returns the KPI dictionary from dim_kpi_metadata.

    Query params
    ------------
    vertical : optional — 'SAAS' | 'FINTECH' | 'MARKETPLACE' | 'INSURTECH'
               When provided, returns GENERAL KPIs plus vertical-specific ones.
               When omitted, returns the full catalogue.

    No auth required: founders need this before they are logged in to
    the analyst cockpit (UploadFlow step 0 vertical selector).
    """
    try:
        items = _bq_svc.get_kpi_metadata(vertical=vertical)
    except Exception as e:
        print(f"[kpi-metadata] BQ error (non-fatal, returning empty): {e}")
        items = []
    return JSONResponse(content={"status": "ok", "kpis": items, "vertical": vertical})


@app.get("/api/portfolio-companies")
async def get_portfolio_companies(portfolio_id: str = None):
    """
    Returns the canonical company list grouped by portfolio (fund_id).

    Source: dim_company JOIN dim_bucket desde BD_Cometa_Dev (caché 5 min).
    Fallback a PORTFOLIO_MAP estático si BQ no está disponible.

    Each company includes has_data: bool — True when fact_kpi_values has at
    least one row with value_status in ('legacy', 'missing_legacy', 'verified').

    Query params
    ------------
    portfolio_id : optional — fund_id ("F001", "CIII", etc.).
                   If omitted, returns all funds.
    """
    # ── 1. Catálogo desde BQ (con caché) — fallback a PORTFOLIO_MAP ───────────
    try:
        catalog = _bq_svc.get_portfolio_catalog()
    except Exception as _cat_err:
        print(f"[portfolio-companies] BQ catalog failed, using legacy PORTFOLIO_MAP: {_cat_err}")
        catalog = None  # type: ignore[assignment]

    # ── 2. Empresas con datos en BQ (non-fatal) ───────────────────────────────
    companies_with_data: set[str] | None = set()
    try:
        _bq  = _get_bq_client_for_api()
        _sql = f"""
            SELECT DISTINCT LOWER(company_id) AS company_id
            FROM `{_bq.project}.{BQ_DATASET}.submissions`
            WHERE company_id IS NOT NULL
        """
        for row in _bq.query(_sql).result():
            if row.company_id:
                companies_with_data.add(row.company_id)
    except Exception as _bq_err:
        print(f"[portfolio-companies] has_data check failed (non-fatal): {_bq_err}")
        companies_with_data = None  # default has_data=True so UI stays usable

    def _sort_entries(entries: list[dict]) -> list[dict]:
        overviews = [e for e in entries if e.get("is_overview")]
        rest      = sorted((e for e in entries if not e.get("is_overview")), key=lambda e: e["label"])
        return overviews + rest

    # ── 3a. BQ path ───────────────────────────────────────────────────────────
    if catalog is not None:
        grouped: dict[str, list[dict]] = {}
        for company in catalog:
            fid = company["fund_id"]
            if portfolio_id and fid != portfolio_id:
                continue
            cid      = company["company_id"]
            has_data = True if companies_with_data is None else (cid.lower() in companies_with_data)
            grouped.setdefault(fid, []).append({
                "key":         cid,
                "id":          cid,
                "company_id":  cid,
                "slug":        cid.lower(),
                "label":       company["company_name"],
                "bucket":      company["bucket_name"],
                "is_overview": False,
                "has_data":    has_data,
            })

        return JSONResponse(content={
            "status":     "success",
            "source":     "bigquery",
            "portfolios": [
                {
                    "portfolio_id":   fid,
                    "portfolio_name": f"Fondo {fid}",
                    "companies":      _sort_entries(entries),
                }
                for fid, entries in sorted(grouped.items())
            ],
        })

    # ── 3b. Legacy fallback — PORTFOLIO_MAP ───────────────────────────────────
    grouped_legacy: dict[str, list[dict]] = {}
    for key, info in PORTFOLIO_MAP.items():
        pid = info["portfolio_id"]
        if portfolio_id and pid != portfolio_id:
            continue
        label       = info.get("display_name") or key.capitalize()
        is_overview = info.get("entity_type") == "FUND_OVERVIEW"
        has_data    = True if companies_with_data is None else (key.lower() in companies_with_data)
        grouped_legacy.setdefault(pid, []).append({
            "key":         key,
            "id":          key,
            "company_id":  key,
            "slug":        key.lower(),
            "label":       label,
            "is_overview": is_overview,
            "has_data":    has_data,
        })

    return JSONResponse(content={
        "status":     "success",
        "source":     "legacy_static",
        "portfolios": [
            {
                "portfolio_id":   pid,
                "portfolio_name": f"Fondo {pid}",
                "companies":      _sort_entries(entries),
            }
            for pid, entries in sorted(grouped_legacy.items())
        ],
    })


@app.get("/api/results/all", deprecated=True, include_in_schema=False)
async def get_all_results_global(token: dict = Depends(_require_auth)):
    """[RETIRADO] Este endpoint ya no está en uso activo.
    Los dashboards del analista deben usar GET /api/results (por portfolio) o
    la capa de BigQuery directamente vía BQDataService.
    """
    return JSONResponse(content={
        "status":      "success",
        "results":     [],
        "total":       0,
        "_deprecated": "Este endpoint ha sido retirado. Usa GET /api/results.",
    })


@app.get("/api/analytics/portfolio")
async def get_portfolio_analytics(portfolio_id: str = "CIII", token: dict = Depends(_require_auth)):
    """
    Aggregated KPI analytics from BD_Cometa_Dev for the requested portfolio.
    Groups by (month, company_id) and returns per-KPI averages via BQDataService.
    """
    try:
        result = _bq_svc.get_portfolio_analytics(fund_id=portfolio_id)
        return JSONResponse(content={"status": "success", **result})
    except Exception as e:
        print(f"[API/analytics] Error: {e}")
        raise HTTPException(status_code=500, detail=f"Error obteniendo analytics: {str(e)}")


@app.get("/api/audit")
async def get_audit_report(portfolio_id: str = None):
    """
    Runs the BigQuery post-insert audit across fact_kpi_values.
    Returns all rows flagged as ERROR or WARNING with their audit_status.

    Query params
    ------------
    portfolio_id : optional — "VII" or "CIII". If omitted, audits all funds.

    audit_status values
    -------------------
    PASS                          — row is clean
    ERROR: Duplicado              — more than one row for same (company, metric, period)
    ERROR: Valor no numérico      — is_valid = FALSE
    ERROR: Confianza crítica (<0.70) — Gemini was highly uncertain
    ADVERTENCIA: Confianza baja (<0.85) — flagged for human review
    """
    # Migrado a Star Schema v2.0 — este endpoint apuntaba a fact_kpi_values (legacy).
    # Para auditoría de calidad, usar los datos directamente desde BD_Cometa_Dev.
    raise HTTPException(
        status_code=410,
        detail="GET /api/audit retirado en v2.0. Usa BD_Cometa_Dev.fact_kpi_values para auditoría.",
    )


@app.get("/api/audit/fidelity/{submission_id}")
async def get_fidelity_audit(submission_id: str):
    """
    Reporte de Fidelidad de Datos — Auditor Senior Cometa.

    Ejecuta tres auditorías encadenadas sobre una submission específica:

    1. identity_check
       Verifica que company_id esté en los 30 registros oficiales de dim_company
       y que el bucket_id asignado coincida con el registro canónico COMPANY_BUCKET.

    2. calculator_audit
       Clasifica cada KPI como 'gemini' (extraído del PDF) o 'calculated' (derivado
       por Python).  Para gross_profit_margin y ebitda_margin re-ejecuta la fórmula
       matemática y compara contra el valor almacenado.  Discrepancias > 0.5pp
       generan WARN; > 2pp generan ERROR (posible manipulación de reporte founder).

    3. checklist_diagnosis
       Cruza los KPIs válidos del reporte con SECTOR_REQUIREMENTS del vertical
       (SAAS / LEND / ECOM / INSUR / OTH) y devuelve missing_kpis con mensaje
       listo para mostrar al founder.

    overall_status
    --------------
    PASS  — sin hallazgos de error ni advertencia
    WARN  — advertencias presentes; el reporte es usable pero debe revisarse
    FAIL  — errores bloqueantes detectados (identidad incorrecta, discrepancia alta)

    Path param
    ----------
    submission_id : UUID de la submission (devuelto por /upload o /api/submissions)
    """
    # Migrado a Star Schema v2.0 — auditoría de fidelidad disponible en BD_Cometa_Dev.
    raise HTTPException(
        status_code=410,
        detail="GET /api/audit/fidelity retirado en v2.0. Consulta BD_Cometa_Dev.submissions.",
    )


# ── CSV Export endpoint ───────────────────────────────────────────────────────

import csv as _csv_mod
import io  as _io_mod

from fastapi.responses import Response as _PlainResponse

@app.get("/api/export/csv")
@limiter.limit("10/minute")
async def export_kpi_csv(
    request:      Request,
    portfolio_id: str | None = None,
    company_id:   str | None = None,
    token:        dict = Depends(_require_auth),
    _origin:      None = Depends(_verify_origin),
):
    """
    Export KPI data as UTF-8 CSV from BigQuery.

    Query params
    ------------
    portfolio_id : filter by fund  (e.g. "CIII" or "VII")
    company_id   : filter by company domain  (e.g. "solvento.com")

    Multi-tenant rule (A1):
    - FND- founders are hard-locked to their own company_id from JWT.
    - ANA- analysts may export any scope requested.

    Columns: Empresa, Fondo, Período, KPI, Valor, Unidad, Confianza, Procesado
    """
    from google.cloud import bigquery as _bq

    # A1 — tenant isolation
    jwt_company      = _derive_tenant_from_token(token)
    effective_company = jwt_company if jwt_company is not None else company_id

    # Build parameterised BigQuery query
    ds      = f"{PROJECT_ID}.{BQ_DATASET}"
    filters = ["f.is_valid = TRUE", "f.raw_value IS NOT NULL"]
    params: list[_bq.ScalarQueryParameter] = []

    if portfolio_id:
        filters.append("s.portfolio_id = @portfolio_id")
        params.append(_bq.ScalarQueryParameter("portfolio_id", "STRING", portfolio_id))
    if effective_company:
        filters.append("LOWER(s.company_id) LIKE @company_id")
        params.append(_bq.ScalarQueryParameter("company_id", "STRING",
                                               f"%{effective_company.lower()}%"))

    sql = f"""
        SELECT
            s.company_id,
            s.portfolio_id,
            s.period_id,
            f.kpi_label,
            f.raw_value,
            f.unit,
            f.confidence_score,
            s.submitted_at
        FROM `{ds}.fact_kpi_values` f
        JOIN `{ds}.submissions`     s ON f.submission_id = s.submission_id
        WHERE {' AND '.join(filters)}
        ORDER BY s.company_id, s.submitted_at DESC, f.kpi_label
        LIMIT 10000
    """

    try:
        client  = _get_bq_client_for_api()
        job     = client.query(sql, job_config=_bq.QueryJobConfig(query_parameters=params))
        rows    = list(job.result())
    except Exception as exc:
        print(f"[export/csv] BigQuery error: {exc}")
        raise HTTPException(status_code=500,
                            detail=f"Error consultando BigQuery: {str(exc)}")

    # Build CSV in memory
    buf    = _io_mod.StringIO()
    writer = _csv_mod.writer(buf)
    writer.writerow(["Empresa", "Fondo", "Período", "KPI", "Valor",
                     "Unidad", "Confianza", "Procesado"])
    for r in rows:
        conf_str = f"{float(r.confidence_score):.2f}" if r.confidence_score is not None else ""
        date_str = str(r.submitted_at)[:10] if r.submitted_at else ""
        writer.writerow([
            r.company_id   or "",
            r.portfolio_id or "",
            r.period_id    or "",
            r.kpi_label    or "",
            r.raw_value    or "",
            r.unit         or "",
            conf_str,
            date_str,
        ])

    scope_tag = (effective_company or portfolio_id or "all").replace("/", "-")
    filename  = f"cometa_kpis_{scope_tag}_{datetime.now().strftime('%Y%m%d')}.csv"

    return _PlainResponse(
        content=buf.getvalue().encode("utf-8-sig"),   # BOM for Excel compatibility
        media_type="text/csv; charset=utf-8",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


# ── RAG Chat endpoint ─────────────────────────────────────────────────────────

from src.ai_engine import (
    build_rag_prompt,
    call_gemini        as _call_gemini_engine,
    call_gemini_stream as _call_gemini_stream,
)

_CHAT_MAX_QUESTION_CHARS   = 500
_CHAT_MAX_SUMMARY_CHARS    = 300  # cap the frontend-supplied executive summary

# ── UI Action marker extraction ────────────────────────────────────────────────
# Gemini appends <!--ACTION:{...}--> at the end of answers when is_analyst=True.
# The blocking /api/chat endpoint strips it server-side and surfaces it as
# ui_action in the JSON response.  The SSE /api/chat/stream endpoint lets tokens
# pass through — the frontend scans for the marker at [DONE].

_ACTION_MARKER_RE = re.compile(r"<!--ACTION:(.*?)-->", re.DOTALL)


def _extract_ui_action(text: str) -> tuple[str, dict | None]:
    """Strip the <!--ACTION:{...}--> marker from *text*.

    Returns ``(clean_text, action_dict)`` where ``action_dict`` is ``None``
    when no marker is found or the JSON inside it is malformed.
    """
    import json as _j

    match = _ACTION_MARKER_RE.search(text)
    if not match:
        return text, None
    try:
        action = _j.loads(match.group(1))
    except Exception:
        return text, None
    clean = _ACTION_MARKER_RE.sub("", text).strip()
    return clean, action

class ChatRequest(BaseModel):
    """
    Body for POST /api/chat.

    Fields
    ------
    question:           The analyst's question (max 500 chars — C5).
    portfolio_id:       Optional portfolio filter.
    company_id:         Company in focus — only respected for ANA- users (A1).
    executive_summary:  Pre-computed KPI one-liner from ExecutiveSummaryText;
                        injected into the Gemini prompt when caller is ANA-.
    """
    question:           str
    portfolio_id:       str | None = None
    company_id:         str | None = None       # ignored for FND- founders (A1)
    executive_summary:  str | None = None       # ANA- analyst context only

@app.post("/api/chat")
@limiter.limit("20/minute")
async def portfolio_chat(
    request: Request,
    req: ChatRequest,
    token: dict = Depends(_require_auth),
    _origin: None = Depends(_verify_origin),
):
    """
    RAG chat: consulta BigQuery → arma contexto → llama Gemini → devuelve respuesta.
    Controles activos: C5 (prompt injection), A1 (multi-tenant), A2 (rate limit), C4 (origin).
    """
    # C5 — Límite de longitud de la pregunta
    question_raw = (req.question or "").strip()
    if not question_raw:
        raise HTTPException(status_code=400, detail="La pregunta no puede estar vacía.")
    if len(question_raw) > _CHAT_MAX_QUESTION_CHARS:
        raise HTTPException(
            status_code=400,
            detail=f"La pregunta excede el límite de {_CHAT_MAX_QUESTION_CHARS} caracteres "
                   f"({len(question_raw)} recibidos).",
        )

    # A1 — Aislamiento multi-tenant: company_id derivado del JWT, nunca del body
    jwt_company_id       = _derive_tenant_from_token(token)
    effective_company_id = jwt_company_id if jwt_company_id is not None else req.company_id

    # I1 — Extracción de identidad del JWT (Capa Humana)
    user_id_claim: str = token.get("user_id", "")
    user_name:     str = (token.get("name")  or token.get("sub") or "").strip()
    user_role:     str = (token.get("role")  or "").strip()
    is_analyst         = user_id_claim.startswith("ANA-")

    # C5 — Cap the executive summary length to avoid prompt stuffing
    executive_summary: str | None = None
    if is_analyst and req.executive_summary:
        executive_summary = req.executive_summary.strip()[:_CHAT_MAX_SUMMARY_CHARS]

    # 1. Recuperar contexto de BigQuery
    raw_rows = _query_rag_context(req.portfolio_id, effective_company_id)

    # A3 — RAG Leak Protection: verifica que todos los rows pertenezcan al tenant
    rows = _verify_rag_integrity(raw_rows, effective_company_id or "")

    # Detectar si algún KPI en el contexto aún no ha sido verificado manualmente
    has_legacy_data = any(not row.get("is_manually_edited", False) for row in rows)

    # KPI Dictionary — enriquece el prompt con definiciones y año de alta (non-fatal)
    kpi_dict = _fetch_kpi_dict_for_rag()

    # Resolver nombre legible de la empresa desde el catálogo BQ (non-fatal)
    company_name = _resolve_company_name(effective_company_id)

    # 2. Build structured prompt via ai_engine — con identidad, advertencia legacy y diccionario
    prompt = build_rag_prompt(
        question=question_raw,
        context_rows=rows,
        company_id=effective_company_id,
        company_name=company_name,
        portfolio_id=req.portfolio_id,
        executive_summary=executive_summary,
        is_analyst=is_analyst,
        user_name=user_name,
        user_role=user_role,
        has_legacy_data=has_legacy_data,
        kpi_dict=kpi_dict,
    )

    # 3. Llamar Gemini via ai_engine
    try:
        answer = _call_gemini_engine(prompt, PROJECT_ID, VERTEX_LOCATION)
    except Exception as e:
        print(f"[RAG/chat] Gemini error: {e}")
        raise HTTPException(status_code=500, detail=f"Error generando respuesta: {str(e)}")

    # 4. AI Audit Log — Cloud Logging nativo (gratis en Cloud Run, sin BQ write)
    log.info(
        "[chat] user=%r role=%r company=%r portfolio=%r context_rows=%d has_legacy=%s q=%r",
        user_id_claim, user_role, effective_company_id, req.portfolio_id,
        len(rows), has_legacy_data, question_raw[:120],
    )

    # Extract optional ui_action marker that Gemini may have appended
    clean_answer, ui_action = _extract_ui_action(answer)

    response_data: dict = {
        "status":          "success",
        "answer":          clean_answer,
        "sources_count":   len(rows),
        "has_legacy_data": has_legacy_data,
        "portfolio_id":    req.portfolio_id,
        "company_id":      effective_company_id,
    }
    if ui_action:
        response_data["ui_action"] = ui_action

    return JSONResponse(content=response_data)


@app.post("/api/chat/stream")
@limiter.limit("20/minute")
async def portfolio_chat_stream(
    request: Request,
    req: ChatRequest,
    token: dict = Depends(_require_auth),
    _origin: None = Depends(_verify_origin),
):
    """
    SSE streaming chat — identical security controls as /api/chat.

    Response format (text/event-stream):
      data: {"token": "<chunk>"}\\n\\n   — incremental Gemini output
      data: {"error":  "<msg>"}\\n\\n   — on Gemini error
      data: [DONE]\\n\\n               — stream completed
    """
    import json as _json_mod

    # ── Input validation (mirrors /api/chat) ─────────────────────────────────
    question_raw = (req.question or "").strip()
    if not question_raw:
        raise HTTPException(status_code=400, detail="La pregunta no puede estar vacía.")
    if len(question_raw) > _CHAT_MAX_QUESTION_CHARS:
        raise HTTPException(
            status_code=400,
            detail=f"La pregunta excede {_CHAT_MAX_QUESTION_CHARS} caracteres.",
        )

    # ── A1 — tenant isolation ──────────────────────────────────────────────────
    jwt_company_id       = _derive_tenant_from_token(token)
    effective_company_id = jwt_company_id if jwt_company_id is not None else req.company_id

    # I1 — Extracción de identidad del JWT (Capa Humana)
    user_id_claim: str = token.get("user_id", "")
    user_name:     str = (token.get("name")  or token.get("sub") or "").strip()
    user_role:     str = (token.get("role")  or "").strip()
    is_analyst         = user_id_claim.startswith("ANA-")

    executive_summary: str | None = None
    if is_analyst and req.executive_summary:
        executive_summary = req.executive_summary.strip()[:_CHAT_MAX_SUMMARY_CHARS]

    # ── Build prompt with A3 leak protection ──────────────────────────────────
    raw_rows = _query_rag_context(req.portfolio_id, effective_company_id)
    rows     = _verify_rag_integrity(raw_rows, effective_company_id or "")

    has_legacy_data = any(not row.get("is_manually_edited", False) for row in rows)

    kpi_dict = _fetch_kpi_dict_for_rag()

    # Resolver nombre legible de la empresa desde el catálogo BQ (non-fatal)
    company_name = _resolve_company_name(effective_company_id)

    prompt = build_rag_prompt(
        question=question_raw,
        context_rows=rows,
        company_id=effective_company_id,
        company_name=company_name,
        portfolio_id=req.portfolio_id,
        executive_summary=executive_summary,
        is_analyst=is_analyst,
        user_name=user_name,
        user_role=user_role,
        has_legacy_data=has_legacy_data,
        kpi_dict=kpi_dict,
    )

    # ── AI Audit Log — Cloud Logging nativo ────────────────────────────────────
    log.info(
        "[chat/stream] user=%r role=%r company=%r portfolio=%r context_rows=%d has_legacy=%s q=%r",
        user_id_claim, user_role, effective_company_id, req.portfolio_id,
        len(rows), has_legacy_data, question_raw[:120],
    )

    # ── SSE generator ──────────────────────────────────────────────────────────
    async def _sse_generator():
        try:
            for chunk in _call_gemini_stream(prompt, PROJECT_ID, VERTEX_LOCATION):
                payload = _json_mod.dumps({"token": chunk}, ensure_ascii=False)
                yield f"data: {payload}\n\n"
        except Exception as exc:
            err = _json_mod.dumps({"error": str(exc)}, ensure_ascii=False)
            yield f"data: {err}\n\n"
        finally:
            yield "data: [DONE]\n\n"

    return StreamingResponse(
        _sse_generator(),
        media_type="text/event-stream",
        headers={
            "Cache-Control":     "no-cache",
            "X-Accel-Buffering": "no",   # disables nginx proxy buffering
        },
    )


# ═══════════════════════════════════════════════════════════════════════════════
# AUTENTICACIÓN — /api/login  &  /api/me
# ═══════════════════════════════════════════════════════════════════════════════

# ── CompanyMapper — resolución determinista email → identidad empresa ──────────
# Clave: dominio completo en minúsculas. Valor: (slug, nombre de display).
# El slug coincide con las claves de PORTFOLIO_MAP / COMPANY_BUCKET.
_DOMAIN_SLUG_MAP: dict[str, tuple[str, str]] = {
    "solvento.com":     ("solvento",     "Solvento"),
    "hunty.com":        ("hunty",         "Hunty"),
    "hunty.io":         ("hunty",         "Hunty"),
    "kueski.com":       ("kueski",        "Kueski"),
    "conekta.com":      ("conekta",       "Conekta"),
    "simetrik.com":     ("simetrik",      "Simetrik"),
    "skydropx.com":     ("skydropx",      "Skydropx"),
    "yotepresto.com":   ("yotepresto",    "Yo Te Presto"),
    "mpower.mx":        ("mpower",        "mPower"),
    "ivoy.mx":          ("ivoy",          "iVoy"),
    "bewe.io":          ("bewe",          "Bewe"),
    "guros.com":        ("guros",         "Guros"),
    "quinio.com":       ("quinio",        "Quinio"),
    "hackmetrix.com":   ("hackmetrix",    "Hackmetrix"),
    "atani.com":        ("atani",         "Atani"),
    "cluvi.com":        ("cluvi",         "Cluvi"),
    "kuona.io":         ("kuona",         "Kuona"),
    "prometeo.io":      ("prometeo",      "Prometeo"),
    "territorium.com":  ("territorium",   "Territorium"),
    "morgana.mx":       ("morgana",       "Morgana"),
    "duppla.mx":        ("duppla",        "Duppla"),
    "kala.mx":          ("kala",          "Kala"),
    "pulsar.mx":        ("pulsar",        "Pulsar"),
    "numia.mx":         ("numia",         "Numia"),
    "bnext.com":        ("bnext",         "Bnext"),
    "clip.mx":          ("clip",          "Clip"),
    "stori.mx":         ("stori",         "Stori"),
    "treinta.co":       ("treinta",       "Treinta"),
    "kushki.com":       ("kushki",        "Kushki"),
    "pomelo.la":        ("pomelo",        "Pomelo"),
    "mentis.mx":        ("mentis",        "Mentis"),
    # Dominios internos Cometa — analistas, no founders
    "cometa.vc":        ("cometa",        "Cometa"),
    "cometa.com":       ("cometa",        "Cometa"),
    "cometavc.com":     ("cometa",        "Cometa"),
    "cometa.fund":      ("cometa",        "Cometa"),
}

_TEST_LOCAL_RE = re.compile(r"^(founder_test|test|demo|prueba|sandbox)\b", re.IGNORECASE)
_TEST_DOMAIN_RE = re.compile(r"^(test|demo|sandbox|example|localhost)\b", re.IGNORECASE)
_DEMO_IDENTITY: dict[str, str] = {
    "company_slug": "demo-startup",
    "company_name": "Startup Demo",
}


def _resolve_company_identity(email: str) -> dict[str, str]:
    """
    Mapea el email del Founder a ``{company_slug, company_name}`` de forma
    determinista. El resultado se inyecta en el JWT — el usuario no puede
    modificarlo mediante peticiones externas.

    Prioridad (para en el primer match):
      1. Patrones de prueba/demo  → "demo-startup" / "Startup Demo"
      2. Dominio exacto en _DOMAIN_SLUG_MAP
      3. Base del dominio (sin TLD) en PORTFOLIO_MAP → humanizar nombre
      4. Fallback seguro         → "demo-startup" (nunca lanza excepción)
    """
    if not email or "@" not in email:
        return _DEMO_IDENTITY.copy()

    local, domain = email.lower().split("@", 1)
    domain_base = domain.split(".")[0]

    # 1. Patrones de prueba
    if _TEST_LOCAL_RE.match(local) or _TEST_DOMAIN_RE.match(domain_base):
        return _DEMO_IDENTITY.copy()

    # 2. Dominio exacto en mapa estático
    if domain in _DOMAIN_SLUG_MAP:
        slug, name = _DOMAIN_SLUG_MAP[domain]
        return {"company_slug": slug, "company_name": name}

    # 3. Lookup en dim_company (BQ) por company_name normalizado
    # Usa el catálogo en caché de _bq_svc (BQ v2.0).
    try:
        _catalog  = _bq_svc.get_portfolio_catalog()
        _stripped = domain_base.replace("-", "").replace("_", "").lower()
        for _c in _catalog:
            _norm = _c["company_name"].lower().replace(" ", "").replace("-", "").replace("_", "")
            if _norm == _stripped or _c["company_name"].lower() == domain_base:
                _name = _c["company_name"]
                return {"company_slug": domain_base, "company_name": _name}
    except Exception:
        # BQ no disponible o empresa no encontrada → seguir al fallback
        pass

    # 4. Fallback seguro
    return _DEMO_IDENTITY.copy()


class LoginRequest(BaseModel):
    email:    str
    password: str


# ── Helpers de contraseña (bcrypt) ────────────────────────────────────────────

def _is_bcrypt_hash(value: str) -> bool:
    """Detecta si `value` es un hash bcrypt válido ($2b$, $2a$ o $2y$)."""
    return isinstance(value, str) and value.startswith(("$2b$", "$2a$", "$2y$"))


def _hash_password(plaintext: str) -> str:
    """Genera un hash bcrypt con salt aleatorio (12 rondas por defecto)."""
    return bcrypt.hashpw(plaintext.encode("utf-8"), bcrypt.gensalt()).decode("utf-8")


def _verify_password(plaintext: str, stored: str) -> bool:
    """
    Verifica la contraseña contra el valor almacenado.
    - Hash bcrypt → bcrypt.checkpw (timing-safe)
    - Texto plano legacy → comparación directa (solo durante migración)
    """
    if _is_bcrypt_hash(stored):
        return bcrypt.checkpw(plaintext.encode("utf-8"), stored.encode("utf-8"))
    return stored == plaintext


def _load_users() -> list[dict]:
    """Lee users.json desde disco. No cachea para reflejar cambios en caliente."""
    try:
        with open(_USERS_FILE, "r", encoding="utf-8") as fh:
            return json.load(fh).get("users", [])
    except FileNotFoundError:
        return []


def _save_users(users: list[UserSchema]) -> None:
    """
    Persiste usuarios validados en users.json (escritura atómica vía .tmp).

    La firma `list[UserSchema]` es la barrera de seguridad principal:
    es imposible llamar esta función con datos sin validar, ya que
    UserSchema aplica todas sus validaciones en el momento de construcción.

    Flujo garantizado:
      1. Caller construye list[UserSchema]  ← validación ocurre AQUÍ
      2. Si falla → ValidationError antes de abrir ningún archivo
      3. Si pasa → serialización + write atómico (.tmp → replace)
    """
    tmp = _USERS_FILE.with_suffix(".json.tmp")
    payload = [u.model_dump() for u in users]
    tmp.write_text(json.dumps({"users": payload}, indent=2, ensure_ascii=False), encoding="utf-8")
    tmp.replace(_USERS_FILE)


@app.post("/api/login")
@limiter.limit("10/minute")
async def login(request: Request, body: LoginRequest):
    """
    Valida credenciales contra users.json y emite un JWT de 24 h.

    Flujo de seguridad:
      1. Buscar usuario por email.
      2. Verificar contraseña con bcrypt.checkpw (o comparación directa si legacy).
      3. Migración perezosa si cualquier campo está desactualizado:
         - Contraseña en texto plano → hashear con bcrypt antes de guardar.
         - ID en formato legacy      → generar ID Híbrido.
         - Todo el archivo validado con UserSchema ANTES de la escritura atómica.
      4. Emitir JWT con user_id para auditoría.
    """
    email_lc = (body.email or "").strip().lower()
    users     = _load_users()

    idx  = next((i for i, u in enumerate(users) if u.get("email", "").lower() == email_lc), None)
    user = users[idx] if idx is not None else None

    # ── Verificación de credenciales ──────────────────────────────────────────
    if not user or not _verify_password(body.password, user.get("password", "")):
        raise HTTPException(status_code=401, detail="Credenciales inválidas")

    # ── Bloquear cuentas pendientes de activación ──────────────────────────────
    if user.get("status") == "PENDING_INVITE":
        raise HTTPException(
            status_code=401,
            detail="Cuenta pendiente de activación. Revisa tu correo para configurar tu acceso.",
        )

    # ── Migración perezosa: ID Híbrido + hash bcrypt ──────────────────────────
    needs_id_migration = not is_hybrid_id(user.get("id", ""))
    needs_pw_migration = not _is_bcrypt_hash(user.get("password", ""))

    if needs_id_migration or needs_pw_migration:
        if needs_id_migration:
            users[idx]["id"] = generate_hybrid_id(email_lc)
        if needs_pw_migration:
            # Hashear antes de validar con UserSchema (garantía: campo no vacío)
            users[idx]["password"] = _hash_password(body.password)

        # Construir list[UserSchema] — validación completa ANTES de abrir el disco.
        # Migración perezosa de cualquier otro usuario legacy en el mismo archivo.
        # Si falla → ValidationError → handler global 422, disco intacto.
        validated: list[UserSchema] = [
            UserSchema.model_validate(
                u if (is_hybrid_id(u.get("id", "")) and _is_bcrypt_hash(u.get("password", "")))
                else {
                    **u,
                    "id":       u["id"] if is_hybrid_id(u.get("id", ""))
                                else generate_hybrid_id(u.get("email", "")),
                    "password": u["password"] if _is_bcrypt_hash(u.get("password", ""))
                                else _hash_password(u.get("password", "")),
                }
            )
            for u in users
        ]

        _save_users(validated)    # escritura atómica: .tmp → replace
        user = users[idx]         # refrescar referencia local

    user_id = user["id"]
    role    = enforce_internal_role(email_lc, user.get("role", "FOUNDER"))
    name    = user.get("name", "")

    # ── Resolución de identidad de empresa — inyectada en JWT ─────────────────
    # Para FOUNDERs, la identidad viene del dominio del email (no del body).
    # Analistas internos reciben sus propios campos para compatibilidad.
    # Si users.json tiene company_slug / company_name explícitos, tienen prioridad
    # sobre el resolver client-side (útil cuando el email de prueba no coincide
    # con el dominio real de la empresa, ej. founder@demo.com → Quinio).
    company_identity = _resolve_company_identity(email_lc)
    company_slug = user.get("company_slug") or company_identity["company_slug"]
    company_name = user.get("company_name") or company_identity["company_name"]

    token = create_access_token(
        email=email_lc,
        role=role,
        name=name,
        user_id=user_id,
        extra_claims={
            "company_slug": company_slug,
            "company_name": company_name,
            "company_id":   user.get("company_id", company_slug),
        },
    )

    return {
        "access_token": token,
        "token_type":   "bearer",
        "user": {
            "user_id":      user_id,
            "email":        email_lc,
            "name":         name,
            "role":         role,
            "company_id":   user.get("company_id", company_slug),
            "company_slug": company_slug,
            "company_name": company_name,
        },
    }


@app.get("/api/me")
async def get_me(token: dict = Depends(_require_auth)):
    """
    Endpoint protegido. Decodifica el JWT y devuelve la identidad completa.
    company_slug y company_name están en el JWT desde el momento del login —
    sin llamadas adicionales a la base de datos.
    """
    # Tokens de ANALISTA deben provenir del dominio @cometa.vc.
    # Founders usan el mismo endpoint desde su portal — no se restringe por dominio.
    if token.get("role") == "ANALISTA":
        _check_cometa_domain(token)

    email = token.get("email") or token.get("sub", "")
    # Tokens legacy sin company_slug: resolver en caliente (backwards-compat)
    if not token.get("company_slug"):
        identity = _resolve_company_identity(email)
    else:
        identity = {
            "company_slug": token["company_slug"],
            "company_name": token.get("company_name", ""),
        }
    return {
        "user_id":      token.get("user_id", ""),
        "email":        email,
        "name":         token.get("name", ""),
        "role":         token.get("role", ""),
        "company_id":   token.get("company_id", identity["company_slug"]),
        "company_slug": identity["company_slug"],
        "company_name": identity["company_name"],
    }


# ── Invite / Setup-password flow ─────────────────────────────────────────────

_INVITE_TOKEN_TYPE = "invite"
_INVITE_EXPIRE_HOURS = 48

# Regex: min 8 chars, at least one digit, at least one non-alphanumeric char
_PASSWORD_RE = re.compile(r"^(?=.*\d)(?=.*[\W_]).{8,}$")


class SetupPasswordRequest(BaseModel):
    """Body for POST /api/auth/setup-password."""
    token:            str
    password:         str
    password_confirm: str


@app.post("/api/auth/setup-password")
@limiter.limit("10/minute")
async def setup_password(
    request: Request,
    body: SetupPasswordRequest,
) -> JSONResponse:
    """
    Activate a PENDING_INVITE account by setting the initial password.

    Flow
    ----
    1. Decode & verify the invite JWT (type="invite", not expired).
    2. Find the user in users.json — must have status=PENDING_INVITE.
    3. Validate password strength (≥8 chars, ≥1 digit, ≥1 symbol).
    4. Hash password with bcrypt, set status="ACTIVE".
    5. Atomic write via _save_users().
    6. Issue a 24-h access JWT so the founder is logged in immediately.

    Returns
    -------
    Same shape as POST /api/login: { access_token, token_type, user }.
    """
    # ── 1. Validate invite token ───────────────────────────────────────────────
    try:
        claims = jwt.decode(
            body.token,
            _JWT_SECRET,
            algorithms=[_JWT_ALGORITHM],
            options={"verify_aud": False},
        )
    except JWTError as exc:
        raise HTTPException(status_code=400, detail=f"Token inválido o expirado: {exc}")

    if claims.get("type") != _INVITE_TOKEN_TYPE:
        raise HTTPException(status_code=400, detail="Token no es de tipo invitación")

    invite_email: str = (claims.get("sub") or "").strip().lower()
    if not invite_email:
        raise HTTPException(status_code=400, detail="Token no contiene email")

    # ── 2. Find pending user ───────────────────────────────────────────────────
    users = _load_users()
    idx   = next(
        (i for i, u in enumerate(users) if u.get("email", "").lower() == invite_email),
        None,
    )
    if idx is None:
        raise HTTPException(status_code=404, detail="Usuario no encontrado")

    user = users[idx]
    if user.get("status") != "PENDING_INVITE":
        raise HTTPException(
            status_code=409,
            detail="Esta cuenta ya está activa. Inicia sesión normalmente.",
        )

    # ── 3. Validate password strength ─────────────────────────────────────────
    if body.password != body.password_confirm:
        raise HTTPException(status_code=422, detail="Las contraseñas no coinciden")

    if not _PASSWORD_RE.match(body.password):
        raise HTTPException(
            status_code=422,
            detail="La contraseña debe tener al menos 8 caracteres, un número y un símbolo.",
        )

    # ── 4 & 5. Hash + atomic save ─────────────────────────────────────────────
    users[idx]["password"] = _hash_password(body.password)
    users[idx]["status"]   = "ACTIVE"

    validated: list[UserSchema] = [UserSchema.model_validate(u) for u in users]
    _save_users(validated)

    # ── 6. Issue access token ──────────────────────────────────────────────────
    activated = users[idx]
    role      = enforce_internal_role(invite_email, activated.get("role", "FOUNDER"))
    token     = create_access_token(
        email=invite_email,
        role=role,
        name=activated.get("name", ""),
        user_id=activated.get("id", ""),
    )
    print(f"[setup-password] Account activated: {invite_email}")

    return JSONResponse(
        content={
            "access_token": token,
            "token_type":   "bearer",
            "user": {
                "user_id":    activated.get("id", ""),
                "email":      invite_email,
                "name":       activated.get("name", ""),
                "role":       role,
                "company_id": activated.get("company_id", ""),
            },
        },
        status_code=200,
    )


# ── Google OAuth — Analista Cometa ───────────────────────────────────────────
# Solo permite acceso a cuentas @cometa.vc verificadas por Google.
# No requiere contraseña — el ID token de Google es la credencial.
# El analista se registra automáticamente en users.json + dim_analyst (BQ).

_GOOGLE_CLIENT_ID   = os.getenv("GOOGLE_CLIENT_ID", "")
_COMETA_AUTH_DOMAIN = "cometa.vc"


def _upsert_analyst_in_bq(*, user_id: str, email: str, name: str) -> None:
    """
    MERGE del analista en dim_analyst de BigQuery.
    Non-fatal: si la tabla no existe o BQ no está disponible, solo loguea.
    """
    try:
        from google.cloud import bigquery as _bq_mod
        from src.core.bq_data_service import BQ_DATASET as _BQ_DS, _client as _bq_cli
        bq  = _bq_cli()
        sql = f"""
            MERGE `{_BQ_DS}.dim_analyst` AS T
            USING (
                SELECT @user_id     AS user_id,
                       @email       AS email,
                       @name        AS display_name,
                       CURRENT_TIMESTAMP() AS now_ts
            ) AS S
            ON T.email = S.email
            WHEN MATCHED THEN
                UPDATE SET display_name  = S.display_name,
                           last_login_at = S.now_ts
            WHEN NOT MATCHED THEN
                INSERT (user_id, email, display_name, created_at, last_login_at)
                VALUES (S.user_id, S.email, S.display_name, S.now_ts, S.now_ts)
        """
        job = bq.query(sql, job_config=_bq_mod.QueryJobConfig(query_parameters=[
            _bq_mod.ScalarQueryParameter("user_id", "STRING", user_id),
            _bq_mod.ScalarQueryParameter("email",   "STRING", email),
            _bq_mod.ScalarQueryParameter("name",    "STRING", name),
        ]))
        job.result()
        log.info("[google_auth] dim_analyst upserted: %s", email)
    except Exception as _exc:
        log.warning("[google_auth] dim_analyst upsert non-fatal: %s", _exc)


class GoogleAuthRequest(BaseModel):
    """Body para POST /api/auth/google."""
    id_token: str


@app.post("/api/auth/google")
@limiter.limit("20/minute")
async def google_auth(request: Request, body: GoogleAuthRequest):
    """
    Autenticación OAuth con Google — exclusiva para analistas @cometa.vc.

    Flujo
    -----
    1. Verificar el ID token de Google con la clave pública de Google.
    2. Extraer email y verificar que termine en @cometa.vc.
    3. Buscar o crear el usuario en users.json con role=ANALISTA.
    4. MERGE (upsert) en BigQuery dim_analyst — non-fatal.
    5. Emitir un JWT HS256 igual que /api/login.

    La contraseña NO se usa ni se requiere.
    """
    from google.oauth2 import id_token as _g_id_token
    from google.auth.transport import requests as _g_requests

    if not _GOOGLE_CLIENT_ID:
        raise HTTPException(
            status_code=500,
            detail="GOOGLE_CLIENT_ID no está configurado en el servidor.",
        )

    # ── 1. Verificar token con Google ──────────────────────────────────────
    try:
        idinfo = _g_id_token.verify_oauth2_token(
            body.id_token,
            _g_requests.Request(),
            _GOOGLE_CLIENT_ID,
        )
    except ValueError as exc:
        raise HTTPException(status_code=401, detail=f"Token de Google inválido: {exc}")

    email = (idinfo.get("email") or "").strip().lower()

    if not idinfo.get("email_verified"):
        raise HTTPException(status_code=401, detail="El email de Google no está verificado.")

    # ── 2. Restringir a @cometa.vc ─────────────────────────────────────────
    if not email.endswith(f"@{_COMETA_AUTH_DOMAIN}"):
        raise HTTPException(
            status_code=403,
            detail=(
                f"Acceso denegado. Solo cuentas @{_COMETA_AUTH_DOMAIN} "
                "pueden acceder como analistas."
            ),
        )

    name    = idinfo.get("name") or email.split("@")[0].replace(".", " ").title()
    picture = idinfo.get("picture", "")

    # ── 3. Buscar o crear usuario en users.json ────────────────────────────
    users = _load_users()
    idx   = next(
        (i for i, u in enumerate(users) if u.get("email", "").lower() == email),
        None,
    )

    if idx is None:
        # Primer login: registrar automáticamente como ANALISTA
        new_user: dict = {
            "id":            generate_hybrid_id(email),
            "email":         email,
            "password":      _hash_password(secrets.token_urlsafe(32)),  # pw aleatorio, nunca usado
            "name":          name,
            "role":          "ANALISTA",
            "company_id":    "COMETA_INTERNAL",
            "auth_provider": "google",
            "status":        "ACTIVE",
        }
        UserOut.model_validate(new_user)   # validar antes de escribir
        users.append(new_user)
        _save_users([UserSchema.model_validate(u) for u in users])
        user = new_user
        log.info("[google_auth] Nuevo analista registrado: %s", email)
    else:
        user = users[idx]
        # Actualizar nombre si cambió en Google
        if user.get("name") != name:
            users[idx]["name"] = name
            _save_users([UserSchema.model_validate(u) for u in users])
            user = users[idx]

    user_id: str = user["id"]

    # ── 4. Upsert en BigQuery dim_analyst (non-fatal) ──────────────────────
    _upsert_analyst_in_bq(user_id=user_id, email=email, name=name)

    # ── 5. Emitir JWT ─────────────────────────────────────────────────────
    token = create_access_token(
        email   = email,
        role    = "ANALISTA",
        name    = name,
        user_id = user_id,
        extra_claims={
            "company_slug":  "cometa",
            "company_name":  "Cometa",
            "company_id":    "COMETA_INTERNAL",
            "auth_provider": "google",
            "picture":       picture,
        },
    )

    return {
        "access_token": token,
        "token_type":   "bearer",
        "user": {
            "user_id":      user_id,
            "email":        email,
            "name":         name,
            "role":         "ANALISTA",
            "company_id":   "COMETA_INTERNAL",
            "company_slug": "cometa",
            "company_name": "Cometa",
        },
    }


# ── Founder notification endpoints ───────────────────────────────────────────

class NotifyUploadRequest(BaseModel):
    """Body for POST /api/notify/upload."""
    founder_email:  str
    file_hash:      str
    company_domain: str = ""


@app.post("/api/notify/upload")
@limiter.limit("20/minute")
async def notify_upload(
    request: Request,
    body: NotifyUploadRequest,
    token: dict = Depends(_require_auth),
) -> JSONResponse:
    """
    Best-effort upload notification hook.

    Called fire-and-forget by the frontend after each successful document
    upload.  Logs the event; in production this is where a real-time
    Slack/Teams notification could be triggered.

    Always returns 200 so transient failures never block the UI.
    """
    email = (token.get("email") or token.get("sub", "")).strip()
    print(
        f"[notify/upload] hash={body.file_hash!r}  "
        f"company={body.company_domain!r}  founder={email!r}"
    )
    return JSONResponse(content={"status": "ok"}, status_code=200)


_BUCKET_TO_VERTICAL: dict[str, str] = {
    "SAAS":  "SAAS",
    "LEND":  "FINTECH",
    "ECOM":  "MARKETPLACE",
    "INSUR": "INSURTECH",
    "OTH":   "GENERAL",
}


@app.get("/api/founder/config")
@limiter.limit("30/minute")
async def founder_config(
    request: Request,
    token: dict = Depends(_require_auth),
) -> JSONResponse:
    """
    Auto-detects company_id and vertical for the authenticated founder.

    Derives company context from the JWT email domain so the Founder Portal
    never needs to ask the user to choose their company manually.

    Returns
    -------
    JSON ``{ "company_id", "vertical", "is_known", "domain" }``
    """
    email: str = (token.get("email") or token.get("sub", "")).strip()
    if not email or "@" not in email:
        raise HTTPException(status_code=422, detail="email no disponible en el token")
    domain = email.split("@", 1)[1].lower()
    comp_id, _, bucket_id, is_known = get_company_id(domain)
    vertical = _BUCKET_TO_VERTICAL.get(bucket_id, "GENERAL")

    # ── Display name: test accounts, known portfolio, or humanized domain ─────
    _test_local_re = re.compile(r"^(founder_test|test|demo|prueba|sandbox)\b", re.IGNORECASE)
    _test_domain_re = re.compile(r"^(test|demo|sandbox|example|localhost)\b", re.IGNORECASE)
    local_part = email.split("@")[0]
    domain_base = domain.split(".")[0]

    # Demo-startup portfolio entry has an explicit display_name we can look up
    _demo_slugs = {"demo-startup", "demostartup"}
    if _test_local_re.match(local_part) or _test_domain_re.match(domain_base) or comp_id.replace("COMP_", "").lower().replace("_", "-") in _demo_slugs:
        company_display_name = "Startup Demo"
    elif is_known:
        # Resolve display name from BQ catalog (comp_id is the canonical BQ ID)
        try:
            _catalog = _bq_svc.get_portfolio_catalog()
            _bq_entry = next((e for e in _catalog if e["company_id"] == comp_id), None)
            company_display_name = (
                _bq_entry["company_name"] if _bq_entry
                else domain_base.replace("-", " ").replace("_", " ").title()
            )
        except Exception:
            company_display_name = domain_base.replace("-", " ").replace("_", " ").title()
    else:
        company_display_name = domain_base.replace("-", " ").replace("_", " ").title() or "tu empresa"

    return JSONResponse(content={
        "company_id":           comp_id,
        "vertical":             vertical,
        "is_known":             is_known,
        "domain":               domain,
        "company_display_name": company_display_name,
    })


class FinalizeRequest(BaseModel):
    """Body for POST /api/founder/finalize."""
    file_hashes:    list[str]
    company_domain: str = ""   # ignored — slug is always derived from JWT for security
    file_names:     list[str] = []
    manual_kpis:    dict[str, str] | None = None


@app.post("/api/founder/finalize")
@limiter.limit("10/minute")
async def founder_finalize(
    request: Request,
    body: FinalizeRequest,
    token: dict = Depends(_require_auth),
) -> JSONResponse:
    """
    Finalize a founder's expediente.

    Marks the submission set as complete and dispatches an HTML confirmation
    email to the founder.  Safe to call even when no email transport is
    configured — the dev fallback prints to stdout and the endpoint still
    returns 200.

    Parameters
    ----------
    body.file_hashes    : SHA-256 prefixes of every processed document.
    body.company_domain : Company slug, e.g. ``"solvento.com"``.
    body.file_names     : Display names of the processed files (for the email).
    body.manual_kpis    : Any KPI key/value pairs supplied manually by the founder.

    Returns
    -------
    JSON ``{ "status": "ok", "message": "...", "sent_to": email }``
    """
    from src.services.email_service import send_receipt_email
    from src.services.hash_service  import generate_vault_seal

    founder_email: str = (token.get("email") or token.get("sub", "")).strip()

    if not founder_email:
        raise HTTPException(status_code=403, detail="email no disponible en el token")

    # Company slug — JWT is the authoritative source; body.company_domain is ignored.
    # Priority: JWT company_slug → JWT company_id → email domain fallback.
    company_domain = (token.get("company_slug") or "").strip()
    if not company_domain:
        company_id_claim = (token.get("company_id") or "").strip()
        if "@" in company_id_claim:
            company_domain = company_id_claim.split("@")[-1]
        elif company_id_claim:
            company_domain = company_id_claim
        else:
            company_domain = founder_email.split("@")[-1] if "@" in founder_email else "cometa"

    file_names  = body.file_names or [f[:16] + "…" for f in body.file_hashes]
    raw_manual  = body.manual_kpis or {}
    processed_at = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M")

    # ══════════════════════════════════════════════════════════════════════════
    # GATE DE SEGURIDAD FINAL — Regla de hierro: 109/109 o nada.
    # Se evalúa antes de generar el vault_seal, el upload_log y el email.
    # El cliente NO puede bypassear este gate aunque manipule el payload.
    # ══════════════════════════════════════════════════════════════════════════

    # Los KPIs ingresados manualmente en el frontend llegan en body.manual_kpis
    # con keys en formato "KPI-001". Los separamos de las notas (_note suffix).
    _raw_manual_keys = {k for k in raw_manual if not k.endswith("_note")}
    _gate = check_gate_for_finalize(
        company_slug    = company_domain.replace(".", "_").lower(),
        manual_kpi_refs = _raw_manual_keys,
    )

    if not _gate["gate_passed"]:
        _present  = _gate["present"]
        _total    = _gate["total"]
        _missing  = _gate["missing"]
        print(
            f"[founder/finalize] BLOQUEADO — gate {_present}/{_total} "
            f"({_missing} faltantes) company={company_domain!r}"
        )
        raise HTTPException(
            status_code=403,
            detail=(
                "Sincronización abortada: El ADN financiero está incompleto. "
                "No se puede generar un recibo para un reporte parcial. "
                f"({_present}/{_total} KPIs presentes — faltan {_missing})"
            ),
        )

    print(
        f"[founder/finalize] GATE OK — {_gate['present']}/{_gate['total']} KPIs "
        f"period={_gate['period']!r} company={company_domain!r}"
    )
    # ══════════════════════════════════════════════════════════════════════════

    # ── Separate KPI values from justification notes ──────────────────────────
    # MissingDataPanel sends notes with a `_note` suffix: churn_rate_note → note.
    # Values (without suffix) go to manual_kpis for the Vault Seal + upload_log.
    # Notes go to log.info() — Cloud Logging persiste en Cloud Run sin coste de BQ write.
    manual_kpis: dict[str, str] = {}
    rescue_notes: dict[str, str] = {}
    for k, v in raw_manual.items():
        if k.endswith("_note"):
            kpi_key = k[:-5]          # strip "_note" suffix → "churn_rate"
            rescue_notes[kpi_key] = v
        else:
            manual_kpis[k] = v

    # ── Persist rescue notes for each uploaded file (Truth Shield) ─────────
    if rescue_notes:
        log.info(
            "[founder/finalize] rescue_notes: hashes=%d kpis=%s",
            len(body.file_hashes), list(rescue_notes.keys()),
        )
    period_id    = datetime.now(timezone.utc).strftime("%Y")

    # ── Vault Seal — SHA-256 integridad del expediente ────────────────────────
    # Cubre: company_id + file_hashes (ordenados) + timestamp de finalización.
    # Genera un fingerprint determinista y auditable que va al correo y a BQ.
    vault_seal = generate_vault_seal(
        company_id   = company_domain,
        file_hash    = body.file_hashes[0] if body.file_hashes else "",
        kpi_rows     = [
            {"kpi_key": k, "raw_value": v, "unit": "", "is_valid": True}
            for k, v in manual_kpis.items()
        ],
        processed_at = processed_at,
    )

    print(
        f"[founder/finalize] company={company_domain!r}  "
        f"files={len(body.file_hashes)}  founder={founder_email!r}  "
        f"seal={vault_seal[:16]}…"
    )

    # ── Upload log — Cloud Logging nativo (v2.0: no más escritura a BQ por evento) ──
    log.info(
        "[founder/finalize] company=%r founder=%r files=%d manual_kpis=%d seal=%s…",
        company_domain, founder_email, len(body.file_hashes),
        len(manual_kpis), vault_seal[:16],
    )

    # ── Correo de confirmación con Sello de Bóveda ───────────────────────────
    send_receipt_email(
        to_email       = founder_email,
        company_domain = company_domain,
        period         = period_id,
        vault_seal     = vault_seal,
        file_hash      = body.file_hashes[0] if body.file_hashes else "",
        kpi_count      = len(manual_kpis),
        processed_at   = processed_at,
    )

    return JSONResponse(
        content={
            "status":     "ok",
            "message":    "Expediente registrado. Se ha enviado tu Recibo Digital al correo.",
            "sent_to":    founder_email,
            "vault_seal": vault_seal,
        },
        status_code=200,
    )


# ── Admin: invite founder ──────────────────────────────────────────────────────

_INVITE_EXPIRE_HOURS = 48  # noqa: F811 — already defined near setup-password; safe duplicate

_ADMIN_INVITE_FRONTEND_URL = os.getenv(
    "NEXTAUTH_URL",
    "https://cometa-vault-frontend-92572839783.us-central1.run.app",
)

_EMAIL_RE_ADMIN = re.compile(r"^[^\s@]+@[^\s@]+\.[^\s@]+$")


class AdminInviteRequest(BaseModel):
    """Body for POST /api/admin/invite — restricted to ANALISTA role."""
    email:        str
    company_name: str
    name:         str = ""


# ── System Settings ───────────────────────────────────────────────────────────

from src.schemas import SystemSettings, SystemSettingsResponse  # noqa: E402 (local re-import for clarity)

_SETTINGS_FILE = Path(__file__).parent / "settings.json"


def _load_settings() -> SystemSettings:
    """Lee settings.json; devuelve defaults si no existe o está corrupto."""
    try:
        if _SETTINGS_FILE.exists():
            raw = json.loads(_SETTINGS_FILE.read_text(encoding="utf-8"))
            return SystemSettings.model_validate(raw)
    except Exception:
        pass
    return SystemSettings()


def _save_settings(s: SystemSettings) -> None:
    """Escritura atómica de settings.json (mismo patrón que _save_users)."""
    tmp = _SETTINGS_FILE.with_suffix(".json.tmp")
    tmp.write_text(s.model_dump_json(indent=2), encoding="utf-8")
    tmp.replace(_SETTINGS_FILE)


def _settings_to_response(s: SystemSettings) -> SystemSettingsResponse:
    return SystemSettingsResponse(
        llm_model=s.llm_model,
        gcp_region=s.gcp_region,
        gcp_project_id=s.gcp_project_id,
        bq_dataset=s.bq_dataset,
        service_account_key_set=bool(s.service_account_key),
        looker_url=s.looker_url,
        gcs_bucket=s.gcs_bucket,
        alert_webhook=s.alert_webhook,
        min_confidence=s.min_confidence,
        irr_alert_below=s.irr_alert_below,
        notification_email=s.notification_email,
    )


@app.get("/api/admin/settings", response_model=SystemSettingsResponse)
@limiter.limit("60/minute")
async def get_settings(
    request: Request,
    token: dict = Depends(_require_auth),
) -> JSONResponse:
    """
    Devuelve la configuración del sistema.
    Las claves secretas (ai_api_key, db_password) se reemplazan por flags booleanos.
    Requiere rol ANALISTA.
    """
    role = token.get("role", "")
    if role != "ANALISTA":
        raise HTTPException(status_code=403, detail="Acceso restringido a ANALISTA")

    s = _load_settings()
    return JSONResponse(content=_settings_to_response(s).model_dump())


class SettingsUpdateRequest(BaseModel):
    # Vertex AI / LLM
    llm_model:           str = "gemini-1.5-pro"
    gcp_region:          str = "us-central1"
    # GCP Infrastructure
    gcp_project_id:      str = ""
    bq_dataset:          str = ""
    service_account_key: str = ""   # cadena vacía = conservar valor anterior
    # Looker & Export
    looker_url:          str = ""
    gcs_bucket:          str = ""
    alert_webhook:       str = ""
    # Thresholds
    min_confidence:      int = 70
    irr_alert_below:     int = 12
    notification_email:  str = ""


@app.post("/api/admin/settings", response_model=SystemSettingsResponse)
@limiter.limit("20/minute")
async def update_settings(
    request: Request,
    body: SettingsUpdateRequest,
    token: dict = Depends(_require_auth),
) -> JSONResponse:
    """
    Persiste la configuración del sistema GCP (escritura atómica).
    service_account_key solo se actualiza si se envía con contenido.
    Requiere rol ANALISTA.
    """
    role = token.get("role", "")
    if role != "ANALISTA":
        raise HTTPException(status_code=403, detail="Acceso restringido a ANALISTA")

    current = _load_settings()

    updated = SystemSettings(
        llm_model=body.llm_model,
        gcp_region=body.gcp_region,
        gcp_project_id=body.gcp_project_id,
        bq_dataset=body.bq_dataset,
        service_account_key=body.service_account_key if body.service_account_key else current.service_account_key,
        looker_url=body.looker_url,
        gcs_bucket=body.gcs_bucket,
        alert_webhook=body.alert_webhook,
        min_confidence=body.min_confidence,
        irr_alert_below=body.irr_alert_below,
        notification_email=body.notification_email,
    )

    _save_settings(updated)
    return JSONResponse(content=_settings_to_response(updated).model_dump())


@app.get("/api/admin/invitations")
@limiter.limit("30/minute")
async def admin_invitations(
    request: Request,
    token: dict = Depends(_require_auth),
) -> JSONResponse:
    """
    Return all FOUNDER users with their activation status.
    Restricted to ANALISTA role.

    Returns
    -------
    JSON ``{ "invitations": [{ email, name, company_id, status }] }``
    """
    if token.get("role") not in ("ANALISTA",):
        raise HTTPException(status_code=403, detail="Solo analistas pueden ver invitaciones.")

    users = _load_users()
    founders = [
        {
            "email":      u.get("email", ""),
            "name":       u.get("name", ""),
            "company_id": u.get("company_id", ""),
            "status":     u.get("status", "ACTIVE"),
        }
        for u in users
        if u.get("role") == "FOUNDER"
    ]
    return JSONResponse(content={"invitations": founders})


@app.post("/api/admin/invite")
@limiter.limit("20/minute")
async def admin_invite(
    request: Request,
    body: AdminInviteRequest,
    token: dict = Depends(_require_auth),
) -> JSONResponse:
    """
    Create and dispatch a secure founder invitation (ANALISTA-only).

    Flow
    ----
    1. Enforce ANALISTA role.
    2. Validate email format.  Any existing record (ACTIVE or PENDING_INVITE)
       is dropped and recreated — this allows re-inviting founders who lost
       their setup link or whose access needs to be reset.
    3. Auto-derive role: @cometa.vc / @cometa.fund / @cometavc.com → ANALISTA,
       everything else → FOUNDER.
    4. Generate a signed JWT invite token (48 h, type="invite").
    5. Register PENDING_INVITE record in users.json atomically.
    6. Send invite email via email_service.send_invite_email().
    7. Return { status, email, company_name, setup_url }.
    """
    from src.services.email_service import send_invite_email  # lazy import

    # ── 1. Role guard ─────────────────────────────────────────────────────────
    caller_role = token.get("role", "")
    if caller_role not in ("ANALISTA",):
        raise HTTPException(status_code=403, detail="Solo analistas pueden enviar invitaciones.")

    email_lc = (body.email or "").strip().lower()
    if not _EMAIL_RE_ADMIN.match(email_lc):
        raise HTTPException(status_code=422, detail=f"Email inválido: {email_lc!r}")

    company_name = body.company_name.strip()
    if not company_name:
        raise HTTPException(status_code=422, detail="El nombre de la empresa es obligatorio.")

    # ── 2. Duplicate check — ACTIVE users are reset to PENDING_INVITE (re-invite) ─
    users = _load_users()
    existing = next((u for u in users if u.get("email", "").lower() == email_lc), None)
    # Both ACTIVE and PENDING_INVITE: drop stale record and recreate fresh
    if existing:
        users = [u for u in users if u.get("email", "").lower() != email_lc]

    # ── 3. Derive role from email domain ──────────────────────────────────────
    invite_role = "ANALISTA" if any(
        email_lc.endswith(d) for d in _INTERNAL_DOMAINS
    ) else "FOUNDER"

    # ── 4. Generate invite token ──────────────────────────────────────────────
    now = datetime.now(timezone.utc)
    invite_payload = {
        "type":         _INVITE_TOKEN_TYPE,
        "sub":          email_lc,
        "email":        email_lc,
        "company_name": company_name,
        "iat":          now,
        "exp":          now + timedelta(hours=_INVITE_EXPIRE_HOURS),
    }
    invite_token = jwt.encode(invite_payload, _JWT_SECRET, algorithm=_JWT_ALGORITHM)
    setup_url    = f"{_ADMIN_INVITE_FRONTEND_URL}/auth/setup-password?token={invite_token}"

    # ── 5. Register PENDING_INVITE user ──────────────────────────────────────
    company_domain = email_lc.split("@")[1] if "@" in email_lc else ""
    company_id     = company_domain or company_name.lower().replace(" ", "_")
    placeholder_pw = f"LOCKED:{secrets.token_hex(24)}"
    new_user_dict  = {
        "id":         generate_hybrid_id(email_lc),
        "email":      email_lc,
        "password":   placeholder_pw,
        "name":       body.name.strip() or company_name,
        "role":       invite_role,
        "company_id": company_id,
        "status":     "PENDING_INVITE",
    }

    all_users = users + [new_user_dict]
    validated: list[UserSchema] = [UserSchema.model_validate(u) for u in all_users]
    _save_users(validated)
    print(f"[admin/invite] Registered {email_lc!r} as PENDING_INVITE role={invite_role} (company={company_name!r})")

    # ── 5. Send invite email ──────────────────────────────────────────────────
    sent, email_error = send_invite_email(
        to_email=email_lc,
        company_name=company_name,
        setup_url=setup_url,
    )
    if sent:
        print(f"[admin/invite] Email sent to {email_lc!r}")
    else:
        print(f"[admin/invite] WARN: Email not sent — {email_error}")

    return JSONResponse(
        content={
            "status":       "ok",
            "email":        email_lc,
            "company_name": company_name,
            "setup_url":    setup_url,
            "email_sent":   sent,
            "email_error":  email_error,
        },
        status_code=200,
    )


# ── Admin: Import MasterDatabase → BigQuery ───────────────────────────────────

@app.post("/api/admin/import-master-db")
@limiter.limit("5/minute")
async def import_master_db(
    request: Request,
    token: dict = Depends(_require_auth),
) -> JSONResponse:
    """
    Preprocess all MasterDatabase CSVs and load them into fact_portfolio_kpis
    in BigQuery using MERGE (upsert). ANALISTA-only.

    Datos certificados por analistas (confidence_score >= 80) nunca son
    sobreescritos por esta operacion. Solo datos legacy (confidence < 60)
    pueden ser actualizados.

    Flow
    ----
    1. Enforce ANALISTA role.
    2. Call process_all() → returns cleaned DataFrame (wide→tall).
    3. [RETIRADO v2.0] Datos migrados a BD_Cometa_Dev.
    4. Return { status, rows_loaded, table, loaded_at, warnings }.
    """
    # Migrado a Star Schema v2.0 — la MasterDB ya fue importada a BD_Cometa_Dev.
    # fact_portfolio_kpis (destino legacy) ya no recibe cargas nuevas.
    raise HTTPException(
        status_code=410,
        detail="POST /api/admin/import-master-db retirado en v2.0. Los datos históricos viven en BD_Cometa_Dev.",
    )


# ══════════════════════════════════════════════════════════════════════════════
# ANALYST REVIEW — Cerebro + Finalize
# ══════════════════════════════════════════════════════════════════════════════

from src.schemas import KpiReviewRow, FinalizeAnalysisRequest, FinalizeAnalysisResponse
from src.core.vc_validator import run_cerebro


class CerebroRequest(BaseModel):
    kpi_rows: list[dict]


@app.post("/api/analyst/cerebro")
@limiter.limit("30/minute")
async def analyst_cerebro(
    request: Request,
    body: CerebroRequest,
    token: dict = Depends(_require_analyst_auth),
) -> JSONResponse:
    """
    Corre el Cerebro de validacion sobre una lista de kpi_rows extraidos por Gemini.

    Aplica reglas de fisica financiera (VIO-001..004) y calcula cross-checks
    derivados (Net Burn, Runway). Devuelve las filas enriquecidas listas para
    la Vista de Analista.

    Input
    -----
    { "kpi_rows": [...] }   — lista de kpi_rows del contrato de datos

    Output
    ------
    {
      "enriched_rows"         : [...],  -- filas con physics_violation + cerebro_alert
      "derived_rows"          : [...],  -- Net Burn y Runway calculados
      "violations"            : [...],  -- mensajes de errores de fisica
      "missing_required"      : [...],  -- KPIs VC obligatorios ausentes
      "has_physics_violations": bool,
      "cross_checks"          : {...},  -- metricas derivadas
      "approval_blocked"      : bool
    }
    """
    if not body.kpi_rows:
        raise HTTPException(status_code=422, detail="kpi_rows no puede estar vacio.")

    try:
        result = run_cerebro(body.kpi_rows)
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"Cerebro error: {exc}")

    return JSONResponse(content=result)


@app.get("/api/analyst/buckets")
@limiter.limit("60/minute")
async def list_analyst_buckets(
    request:    Request,
    layer:      str = "stage",
    company:    str = "",
    limit:      int = 50,
    page_token: str = "",
    token:      dict = Depends(_require_analyst_auth),
) -> JSONResponse:
    """
    Lista archivos reales de GCS para la Consola de Supervisión del Analista.

    Capas disponibles
    -----------------
    - raw          : GCS_OUTPUT_BUCKET/vault/{slug}/
    - stage        : MEDALLION_BUCKET/stage/
    - gold         : BUCKET_GOLD/ (cometa-vc-gold-prod)
    - historicofund: HISTORICOFUND_BUCKET/

    Por cada blob extrae el company_slug del path, lo cruza contra el CSV de
    historicofund para indicar si la empresa tiene datos certificados y cuál
    es su nombre oficial.

    Query params
    ------------
    layer       — raw | stage | gold | historicofund  (default: stage)
    company     — filtrar por slug de empresa (opcional)
    limit       — máximo de resultados (default 50, max 200)
    page_token  — token de paginación de GCS (devuelto en next_page_token)
    """
    import re as _re

    valid_layers = {"raw", "stage", "vault", "gold", "historicofund", "pending"}
    if layer not in valid_layers:
        raise HTTPException(status_code=422, detail=f"layer debe ser uno de {sorted(valid_layers)}")

    limit = min(max(1, limit), 200)

    # ── Bucket y prefix por capa ──────────────────────────────────────────────
    # ── Bucket config desde la fuente única de verdad ────────────────────────
    # RAW_BUCKET   — archivos originales subidos por founders (NUNCA se borran)
    # STAGE_BUCKET — JSONs de Gemini + vault/ + pending_mapper/
    # GOLD_BUCKET  — KPIs certificados por el analista
    # HIST_BUCKET  — CSV maestro del fondo (solo lectura)

    layer_cfg: dict[str, tuple[str, str]] = {
        "raw":          (RAW_BUCKET,   ""),            # Origen: PDFs/XLSX originales
        "stage":        (STAGE_BUCKET, "stage/"),      # Extracción: JSONs de Gemini
        "vault":        (STAGE_BUCKET, "vault/"),      # Vault: resultados procesados
        "gold":         (GOLD_BUCKET,  ""),            # Certificado: KPIs aprobados
        "historicofund":(HIST_BUCKET,  ""),            # Histórico: CSV maestro fondo
        "pending":      (STAGE_BUCKET, "pending_mapper/"),  # Excel/CSV en cola
    }
    bucket_name, prefix = layer_cfg[layer]

    # ── Cargar mapa histórico para validación de empresas ────────────────────
    hist_map = _load_historicofund_map()   # {slug -> official_name}

    def _slug_from_path(blob_name: str, pfx: str) -> str:
        """Extrae el slug (primer segmento después del prefix)."""
        tail = blob_name[len(pfx):] if blob_name.startswith(pfx) else blob_name
        parts = tail.strip("/").split("/")
        raw   = parts[0] if parts else ""
        return _re.sub(r"[^a-z0-9]+", "_", raw.strip().lower()).strip("_")

    # ── Listar blobs — resiliente: nunca lanza 502 ───────────────────────────
    # Si GCS no está disponible (credenciales, red, bucket inexistente) devolvemos
    # una respuesta vacía con un aviso amigable en lugar de colapsar la consola.
    # Esto es especialmente importante para la capa historicofund, que puede no
    # estar accesible en entornos de desarrollo o staging.
    blobs: list      = []
    next_token: str  = ""
    gcs_warning: str = ""

    try:
        gcs = _get_storage_client()
        bkt = gcs.bucket(bucket_name)

        list_prefix = prefix
        if company:
            safe_co  = _re.sub(r"[^a-z0-9_\-]", "", company.lower())
            list_prefix = f"{prefix}{safe_co}/"

        kwargs: dict = {"prefix": list_prefix, "max_results": limit}
        if page_token:
            kwargs["page_token"] = page_token

        page_iter  = bkt.list_blobs(**kwargs)
        blobs      = list(page_iter)
        next_token = getattr(page_iter, "next_page_token", "") or ""

    except Exception as exc:
        _short = str(exc)[:200]
        gcs_warning = (
            f"No se pudo acceder al bucket '{bucket_name}' (capa: {layer}). "
            f"Verifica credenciales GCS o que el bucket exista. Detalle: {_short}"
        )
        print(f"[buckets] WARN (non-fatal): {gcs_warning}")

    # ── Construir lista de BucketFile ─────────────────────────────────────────
    files: list[dict] = []
    for blob in blobs:
        # Skip directory placeholders
        if blob.name.endswith("/"):
            continue

        slug = _slug_from_path(blob.name, prefix)
        official = hist_map.get(slug, "")
        found    = bool(official)

        # Extract load_id from filename pattern  {load_id}_gemini.json etc.
        fname       = blob.name.rsplit("/", 1)[-1]
        load_id_m   = _re.match(r"^([a-f0-9]{8,}(?:_[a-f0-9]+)?)", fname)
        load_id_val = load_id_m.group(1) if load_id_m else ""

        updated_str = blob.updated.isoformat() if blob.updated else ""

        files.append({
            "uri":           f"gs://{bucket_name}/{blob.name}",
            "name":          blob.name,
            "layer":         layer,
            "company_slug":  slug,
            "size_bytes":    blob.size or 0,
            "updated_at":    updated_str,
            "load_id":       load_id_val,
            "company_found": found,
            "official_name": official,
        })

    response: dict = {
        "layer":           layer,
        "files":           files,
        "next_page_token": next_token,
        "total":           len(files),
    }
    if gcs_warning:
        response["warning"] = gcs_warning

    return JSONResponse(content=response)


def _stage_json_to_kpi_rows(stage_data: dict, period_id: str = "") -> list[dict]:
    """
    Convierte cualquier estructura JSON de stage al formato list[dict] que
    run_cerebro() espera (kpi_key, kpi_label, numeric_value, unit, …).

    Orden de intentos:
      1. kpi_rows                — formato nativo del pipeline (pass-through)
      2. financial_metrics_2025  — output crudo de Gemini (anidado)
      3. data / metrics / kpis   — claves alternativas conocidas
      4. raíz del objeto         — walk total como último recurso
    """
    from src.core.data_contract import parse_numeric as _parse_num

    def _leaf_to_row(key: str, obj: object, path: str) -> "dict | None":
        """Convierte un nodo hoja a un kpi_row. Retorna None si no es numérico."""
        if isinstance(obj, dict) and "value" in obj:
            raw = str(obj.get("value") or "")
            unit_hint = obj.get("unit") or None
            numeric, unit_det = _parse_num(raw)
            if numeric is None:
                return None
            return {
                "kpi_key":            key,
                "kpi_label":          key.replace("_", " ").title(),
                "numeric_value":      numeric,
                "unit":               unit_hint or unit_det,
                "is_valid":           True,
                "confidence":         "MEDIUM",
                "source_description": f"Gemini — {path}" if path else "Gemini",
                "raw_value":          raw,
                "period_id":          period_id,
            }
        if isinstance(obj, (int, float)):
            return {
                "kpi_key":            key,
                "kpi_label":          key.replace("_", " ").title(),
                "numeric_value":      float(obj),
                "unit":               None,
                "is_valid":           True,
                "confidence":         "MEDIUM",
                "source_description": f"Gemini — {path}" if path else "Gemini",
                "raw_value":          str(obj),
                "period_id":          period_id,
            }
        if isinstance(obj, str):
            numeric, unit_det = _parse_num(obj)
            if numeric is None:
                return None
            return {
                "kpi_key":            key,
                "kpi_label":          key.replace("_", " ").title(),
                "numeric_value":      numeric,
                "unit":               unit_det,
                "is_valid":           True,
                "confidence":         "LOW",
                "source_description": f"Gemini — {path}" if path else "Gemini",
                "raw_value":          obj,
                "period_id":          period_id,
            }
        return None

    def _flatten(node: object, path: str = "") -> list[dict]:
        """Recorre recursivamente y devuelve todas las hojas numéricas."""
        rows: list[dict] = []
        if not isinstance(node, dict):
            return rows
        for k, v in node.items():
            full_path = f"{path}/{k}" if path else k
            if isinstance(v, dict) and "value" in v:
                row = _leaf_to_row(k, v, path)
                if row:
                    rows.append(row)
            elif isinstance(v, dict):
                rows.extend(_flatten(v, full_path))
            else:
                row = _leaf_to_row(k, v, path)
                if row:
                    rows.append(row)
        return rows

    # 1 — Formato nativo
    if stage_data.get("kpi_rows"):
        return list(stage_data["kpi_rows"])

    # 2 — financial_metrics_2025 (Gemini crudo)
    fm = stage_data.get("financial_metrics_2025")
    if fm and isinstance(fm, dict):
        rows = _flatten(fm, "financial_metrics_2025")
        if rows:
            return rows

    # 3 — Claves alternativas conocidas
    for alt_key in ("data", "metrics", "kpis"):
        alt = stage_data.get(alt_key)
        if isinstance(alt, list) and alt:
            return alt
        if isinstance(alt, dict):
            rows = _flatten(alt, alt_key)
            if rows:
                return rows

    # 4 — Walk completo del objeto raíz (excluyendo metadata de strings)
    return _flatten(stage_data)


def _resolve_raw_uri(stage_data: dict, stage_gcs_uri: str) -> str:
    """
    Devuelve la URI del archivo ORIGINAL en RAW_BUCKET.

    Estrategia (en orden):
    1. Campo ``source_file`` en el JSON de stage — vínculo permanente inyectado
       al guardar (implementado desde la v2.1).
    2. Búsqueda en RAW_BUCKET por prefijo {company}/{hash}: si el stage_uri tiene
       la forma ``vault/{company}/{hash}_result.json``, extrae company + hash y
       lista blobs en RAW_BUCKET/{company}/ que empiecen por ese hash.
    3. Fallback: devuelve el propio stage URI para que el analista no quede bloqueado.
    """
    source = (stage_data.get("source_file") or "").strip()
    if source.startswith("gs://"):
        return source

    # Extraer company y hash desde el URI de stage
    # Formatos conocidos:
    #   gs://bucket/vault/{company}/{hash}_result.json
    #   gs://bucket/stage/{company}/{hash}_{filename}.json
    _m = re.match(r"gs://[^/]+/(?:vault|stage)/([^/]+)/([a-f0-9]{8,})", stage_gcs_uri)
    if _m:
        company   = _m.group(1)
        file_hash = _m.group(2)
        try:
            _gcs = _get_storage_client()
            _blobs = list(
                _gcs.bucket(RAW_BUCKET).list_blobs(
                    prefix=f"{company}/{file_hash}", max_results=5
                )
            )
            if _blobs:
                return f"gs://{RAW_BUCKET}/{_blobs[0].name}"
        except Exception as _lookup_err:
            log.debug("[_resolve_raw_uri] GCS lookup failed: %s", _lookup_err)

    return stage_gcs_uri  # fallback: al menos el analista ve algo


@app.get("/api/analyst/stage-review")
@limiter.limit("30/minute")
async def analyst_stage_review(
    request: Request,
    gcs_uri: str,
    token:   dict = Depends(_require_analyst_auth),
) -> JSONResponse:
    """
    Descarga el JSON de stage desde GCS y ejecuta el Cerebro de validación.

    Acepta múltiples formatos de JSON de stage:
      • kpi_rows nativo        — pipeline original
      • financial_metrics_2025 — output crudo de Gemini (anidado)
      • data / metrics / kpis  — claves alternativas
      • raw fallback            — devuelve el contenido tal cual si no hay KPIs

    Query param
    -----------
    gcs_uri — URI completa de GCS: gs://{bucket}/{blob_path}
    """
    if not gcs_uri.startswith("gs://"):
        raise HTTPException(status_code=422, detail="gcs_uri debe comenzar con gs://")

    import re as _re
    _uri_m = _re.match(r"^gs://([^/]+)/(.+)$", gcs_uri)
    if not _uri_m:
        raise HTTPException(status_code=422, detail="Formato de URI inválido. Esperado: gs://bucket/path")

    bucket_name = _uri_m.group(1)
    blob_path   = _uri_m.group(2)

    # ── Descargar JSON de stage ───────────────────────────────────────────────
    try:
        gcs     = _get_storage_client()
        content = gcs.bucket(bucket_name).blob(blob_path).download_as_text(encoding="utf-8")
    except Exception as exc:
        raise HTTPException(status_code=404, detail=f"Archivo no encontrado en GCS: {exc}")

    try:
        stage_data = json.loads(content)
    except Exception as exc:
        raise HTTPException(status_code=422, detail=f"El archivo no es JSON válido: {exc}")

    period_id = stage_data.get("period") or stage_data.get("period_id") or ""
    kpi_rows  = _stage_json_to_kpi_rows(stage_data, period_id=period_id)

    # ── Fallback seguro: no hay KPIs extraíbles — devolver crudo ─────────────
    if not kpi_rows:
        return JSONResponse(content={
            "enriched_rows":         [],
            "derived_rows":          [],
            "violations":            [],
            "missing_required":      [],
            "has_physics_violations": False,
            "cross_checks":          {},
            "approval_blocked":      False,
            "load_id":               stage_data.get("load_id", ""),
            "slug":                  _resolve_slug(stage_data.get("company_slug", ""))[0],
            "company_name":          _resolve_slug(stage_data.get("company_slug", ""))[1],
            "periodo":               period_id,
            "source_file_uri":       _resolve_raw_uri(stage_data, gcs_uri),
            "stage_uri":             gcs_uri,
            "raw_fallback":          True,
            "raw_data":              stage_data,
            "warning": (
                "El archivo JSON no contiene KPIs en ningún formato conocido. "
                "Se devuelve el contenido crudo para inspección manual."
            ),
        })

    # ── Ejecutar Cerebro ──────────────────────────────────────────────────────
    try:
        cerebro_result = run_cerebro(kpi_rows)
    except Exception as exc:
        # Cerebro falló — devolver las filas extraídas sin enriquecimiento
        cerebro_result = {
            "enriched_rows":         kpi_rows,
            "derived_rows":          [],
            "violations":            [],
            "missing_required":      [],
            "has_physics_violations": False,
            "cross_checks":          {},
            "approval_blocked":      False,
            "cerebro_error":         str(exc),
        }

    # ── Inyectar metadata de stage en la respuesta ────────────────────────────
    cerebro_result["load_id"]        = stage_data.get("load_id", "")
    _resolved_slug, _company_name    = _resolve_slug(stage_data.get("company_slug", ""))
    cerebro_result["slug"]           = _resolved_slug
    cerebro_result["company_name"]   = _company_name
    cerebro_result["periodo"]        = period_id
    cerebro_result["source_file_uri"]= _resolve_raw_uri(stage_data, gcs_uri)
    cerebro_result["stage_uri"]      = gcs_uri
    cerebro_result["raw_fallback"]   = False

    return JSONResponse(content=cerebro_result)


# ── FIX-03: Review PDF — Analista certifica datos de Gemini ──────────────────
# El Founder sube un PDF → Gemini extrae KPIs → queda en vault/ + SQLite.
# El Analista llama a este endpoint para cargar esos KPIs en KpiReviewPanel,
# editar valores, y luego certificar via POST /api/analyst/finalize-analysis.
# Esto cierra el circuito: PDF y Excel/CSV terminan en la misma tabla (fact_portfolio_kpis).

class ReviewPdfRequest(BaseModel):
    """Body para POST /api/analyst/review-pdf."""
    company_slug: str   # ej. "simetrik"
    file_hash:    str   # hash del PDF (primeros 64 hex chars del SHA-256)


@app.post("/api/analyst/review-pdf")
@limiter.limit("20/minute")
async def analyst_review_pdf(
    request: Request,
    body:    ReviewPdfRequest,
    token:   dict = Depends(_require_analyst_auth),
) -> JSONResponse:
    """
    Carga la extraccion de Gemini de un PDF desde vault/{slug}/{hash}_result.json,
    construye kpi_rows y los enriquece con el Cerebro (reglas de fisica financiera).

    Devuelve la misma forma que /api/analyst/stage-review — compatible con
    KpiReviewPanel y con POST /api/analyst/finalize-analysis.

    El Analista es el unico que puede llamar este endpoint (ANALISTA-only via
    _require_analyst_auth). El Founder no tiene acceso.
    """
    if not body.file_hash or len(body.file_hash) < 16:
        raise HTTPException(status_code=422, detail="file_hash debe tener al menos 16 caracteres.")

    _gcs_key = f"vault/{body.company_slug}/{body.file_hash}_result.json"
    try:
        _sc      = _get_storage_client()
        _content = _sc.bucket(GCS_OUTPUT_BUCKET).blob(_gcs_key).download_as_text(encoding="utf-8")
    except Exception as exc:
        raise HTTPException(
            status_code=404,
            detail=f"PDF result no encontrado en vault ({_gcs_key}): {exc}",
        )

    try:
        gemini_json = json.loads(_content)
    except Exception as exc:
        raise HTTPException(status_code=422, detail=f"El archivo de vault no es JSON valido: {exc}")

    # Construir el contrato canonico desde el JSON de Gemini
    try:
        contract = build_contract(
            gemini_json      = gemini_json,
            file_hash        = body.file_hash,
            company_id       = body.company_slug,
            founder_email    = "",
            original_filename= f"{body.file_hash}_result.json",
        )
        kpi_rows = contract.get("kpi_rows", [])
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"Error construyendo contrato desde Gemini JSON: {exc}")

    if not kpi_rows:
        raise HTTPException(status_code=422, detail="El JSON de vault no contiene kpi_rows procesables.")

    # Enriquecer con el Cerebro (fisica financiera + cross-checks)
    try:
        cerebro_result = run_cerebro(kpi_rows)
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"Cerebro error: {exc}")

    # Metadata para que el frontend pueda llamar finalize-analysis directamente
    cerebro_result["file_hash"]      = body.file_hash
    cerebro_result["slug"]           = body.company_slug
    cerebro_result["source_file_uri"]= f"gs://{GCS_OUTPUT_BUCKET}/{_gcs_key}"
    # load_id derivado del hash (mismo convenio que el upload handler PDF)
    cerebro_result["load_id"]        = body.file_hash[:16]
    cerebro_result["pipeline_source"]= "pdf_gemini"

    print(
        f"[analyst/review-pdf] analyst={token.get('email')!r} "
        f"slug={body.company_slug!r} hash={body.file_hash[:16]}… "
        f"kpi_rows={len(kpi_rows)}"
    )

    return JSONResponse(content=cerebro_result)


@app.post("/api/analyst/finalize-analysis", response_model=FinalizeAnalysisResponse)
@limiter.limit("10/minute")
async def finalize_analysis(
    request: Request,
    body: FinalizeAnalysisRequest,
    token: dict = Depends(_require_analyst_auth),
) -> JSONResponse:
    """
    Cierra el ciclo de validacion: certifica el analisis del analista y escribe
    los datos en la capa Gold de GCS y en BigQuery.

    Flujo
    -----
    1. Validar rol ANALISTA.
    2. Construir el JSON final con metadatos de trazabilidad completos.
    3. Subir JSON a gs://cometa-vc-gold-prod/{slug}/{year}/{month}/{load_id}_final.json
    4. Copiar el PDF original de raw/ a gold/certified/ (server-side copy en GCS).
    5. UPSERT en fact_portfolio_history via MERGE con load_id como clave.
    6. Retornar gold_uri, pdf_gold_uri, bq_rows_upserted, timestamp_gold, warnings.

    Input
    -----
    FinalizeAnalysisRequest — contiene kpi_rows con analyst_value + analyst_note.

    Output
    ------
    FinalizeAnalysisResponse
    """
    # ── 1b. Resolver slug y validar contra CSV histórico de historicofund ────────
    resolved_slug, display_name = _resolve_slug(body.slug)
    is_valid_slug, valid_slugs  = _validate_company_slug(resolved_slug)
    if not is_valid_slug and valid_slugs and not body.force_approve:
        raise HTTPException(
            status_code=422,
            detail={
                "error":      "company_not_found",
                "slug":       resolved_slug,
                "company_name": display_name or resolved_slug,
                "message": (
                    f"La empresa '{display_name or resolved_slug}' no existe todavía "
                    "en el registro histórico del fondo. Si es una empresa nueva, "
                    "activa 'Forzar aprobación' para certificar igualmente."
                ),
                "valid_slugs": valid_slugs,
            },
        )

    year, month = body.periodo.split("-")
    timestamp_gold = datetime.now(timezone.utc).isoformat()
    warnings: list[str] = []
    if body.force_approve and not is_valid_slug:
        warnings.append(
            f"Empresa '{display_name or resolved_slug}' no encontrada en historicofund "
            "— aprobación forzada por el analista."
        )

    # ── 2. Construir JSON final de trazabilidad ────────────────────────────────
    final_contract: dict = {
        "startup_id":       body.slug,
        "period":           body.periodo,
        "currency":         body.currency,
        "source_file_uri":  body.source_file_uri,
        "analyst_id":       body.analyst_id,
        "timestamp_gold":   timestamp_gold,
        "load_id":          body.load_id,
        "kpi_rows": [
            {
                "kpi_key":         r.kpi_key,
                "kpi_label":       r.kpi_label,
                "ai_value":        r.ai_value,
                "analyst_value":   r.analyst_value,
                "final_value":     r.analyst_value if r.analyst_value is not None else r.ai_value,
                "unit":            r.unit,
                "confidence":      r.confidence,
                "physics_violation": r.physics_violation,
                "analyst_note":    r.analyst_note,
                "source":          "analyst_approved" if r.analyst_value is not None else r.source,
            }
            for r in body.kpi_rows
        ],
    }

    # ── 3. Subir JSON final a gold/ ────────────────────────────────────────────
    gold_uri    = ""
    bucket_gold = os.getenv("BUCKET_GOLD", "cometa-vc-gold-prod")

    try:
        gcs_client   = _get_storage_client()
        gold_bucket  = gcs_client.bucket(bucket_gold)
        json_blob_path = f"{body.slug}/{year}/{month}/{body.load_id}_final.json"
        json_blob    = gold_bucket.blob(json_blob_path)
        json_blob.upload_from_string(
            json.dumps(final_contract, ensure_ascii=False, indent=2).encode("utf-8"),
            content_type="application/json",
        )
        gold_uri = f"gs://{bucket_gold}/{json_blob_path}"
        print(f"[finalize] Gold JSON written: {gold_uri}")
    except Exception as exc:
        warnings.append(f"GCS gold JSON upload failed: {exc}")
        print(f"[finalize] WARN: {exc}")

    # ── 4. Copiar PDF de raw/ a gold/certified/ ───────────────────────────────
    pdf_gold_uri = ""
    if body.source_file_uri.startswith("gs://"):
        try:
            uri_parts      = body.source_file_uri[5:].split("/", 1)
            src_bucket_name = uri_parts[0]
            src_blob_path   = uri_parts[1] if len(uri_parts) > 1 else ""
            gcs_c          = _get_storage_client()
            src_bucket     = gcs_c.bucket(src_bucket_name)
            src_blob       = src_bucket.blob(src_blob_path)
            dest_bucket    = gcs_c.bucket(bucket_gold)
            dest_path      = f"{body.slug}/{year}/{month}/{body.load_id}_certified.pdf"
            src_bucket.copy_blob(src_blob, dest_bucket, dest_path)
            pdf_gold_uri   = f"gs://{bucket_gold}/{dest_path}"
            print(f"[finalize] PDF certified: {pdf_gold_uri}")
        except Exception as exc:
            warnings.append(f"PDF copy to gold failed: {exc}")
            print(f"[finalize] WARN PDF copy: {exc}")
    else:
        warnings.append("source_file_uri no es un URI gs:// — PDF no copiado a gold.")

    # ── 5. Insertar en BD_Cometa_Dev (Star Schema v2.0) ──────────────────────────
    # Reemplaza el MERGE a fact_portfolio_kpis (tabla vieja) por
    # insert_submission_and_facts() + update_submission_status("VALIDATED").
    bq_rows_upserted = 0
    try:
        from datetime import date as _date

        period_date = _date(int(year), int(month), 1)
        quarter     = f"Q{((int(month) - 1) // 3) + 1}"
        period_id   = f"P{year}{quarter}M{month.zfill(2)}"

        # Resolver company_id y fund_id desde el catálogo en caché (TTL 5 min).
        # dim_company usa IDs como "C001"; intentamos matchear por company_id o name.
        _catalog    = _bq_svc.get_portfolio_catalog()
        _slug_lower = body.slug.lower()
        _company_row = next(
            (c for c in _catalog
             if c["company_id"].lower() == _slug_lower
             or c.get("company_name", "").lower() == _slug_lower),
            None,
        )
        if _company_row:
            company_id = _company_row["company_id"]
            fund_id    = _company_row["fund_id"]
        else:
            # Slug no encontrado en dim_company — usamos el slug en mayúsculas
            # como company_id de fallback y logueamos para que el data team lo corrija.
            company_id = body.slug.upper()
            fund_id    = "F001"
            warnings.append(
                f"Empresa '{body.slug}' no encontrada en dim_company. "
                "Usando slug como company_id (fallback). Verifica dim_company."
            )
            log.warning("[finalize] slug=%r no resuelto en dim_company — fallback company_id=%r", body.slug, company_id)

        # Mapear KpiReviewRow → formato de kpi_rows de insert_submission_and_facts.
        kpi_rows: list[dict] = [
            {
                "metric_id":   r.kpi_key,
                "value":       float(r.analyst_value) if r.analyst_value is not None
                               else (float(r.ai_value) if r.ai_value is not None else None),
                "value_notes": r.analyst_note,
            }
            for r in body.kpi_rows
            if r.analyst_value is not None or r.ai_value is not None
        ]

        analyst_email: str = token.get("email", body.analyst_id)

        result_sub = _bq_svc.insert_submission_and_facts(
            company_id    = company_id,
            fund_id       = fund_id,
            period_id     = period_id,
            period_start  = period_date,
            submitted_by  = analyst_email,
            source_file   = body.source_file_uri,
            kpi_rows      = kpi_rows,
            review_notes  = f"Finalized by {analyst_email} via finalize-analysis",
        )
        sub_id           = result_sub["submission_id"]
        bq_rows_upserted = result_sub["rows_inserted"]

        # El analista ya aprobó — marcar VALIDATED inmediatamente.
        _bq_svc.update_submission_status(
            submission_id = sub_id,
            status        = "VALIDATED",
            review_notes  = f"Auto-validated on finalize by {analyst_email}",
        )
        log.info(
            "[finalize] Star Schema → submission=%r company=%r period=%r %d KPIs VALIDATED",
            sub_id, company_id, period_id, bq_rows_upserted,
        )

    except _CompanyNotFoundError as exc:
        warnings.append(f"Empresa no encontrada en dim_company: {exc}")
        log.error("[finalize] CompanyNotFoundError: %s", exc)
    except _BQCatalogError as exc:
        warnings.append(f"BQ Star Schema write failed: {exc}")
        log.error("[finalize] BQInsertError: %s", exc)
    except Exception as exc:
        warnings.append(f"BQ write failed: {exc}")
        log.error("[finalize] WARN BQ: %s", exc)

    dashboard_url = os.getenv("DASHBOARD_URL", "")

    return JSONResponse(content={
        "gold_uri":         gold_uri,
        "pdf_gold_uri":     pdf_gold_uri,
        "bq_rows_upserted": bq_rows_upserted,
        "timestamp_gold":   timestamp_gold,
        "dashboard_url":    dashboard_url,
        "warnings":         warnings,
    })


def _QUALITY_MAP(confidence: float) -> float:
    """Mapea confidence de Gemini (0-1) a quality_score de BQ."""
    if confidence >= 0.95:
        return 1.00
    if confidence >= 0.85:
        return 0.95
    if confidence >= 0.70:
        return 0.80
    return 0.60


@app.get("/api/analyst/view-file")
@limiter.limit("60/minute")
async def analyst_view_file(
    request: Request,
    uri: str = Query(..., description="gs://bucket/path URI del archivo a visualizar"),
    token: dict = Depends(_require_auth),
) -> JSONResponse:
    """
    GET /api/analyst/view-file?uri=gs://bucket/path

    Genera una Signed URL de lectura (15 min) para que el analista pueda
    visualizar el PDF original directamente desde GCS sin exponerlo publicamente.

    Solo ANALISTA puede acceder.
    """
    role = token.get("role")
    if role != "ANALISTA":
        raise HTTPException(status_code=403, detail="Solo analistas pueden visualizar archivos del pipeline.")

    if not uri.startswith("gs://"):
        raise HTTPException(status_code=400, detail="uri debe comenzar con gs://")

    parts = uri[5:].split("/", 1)
    if len(parts) != 2 or not parts[1]:
        raise HTTPException(status_code=400, detail="uri gs:// invalido — formato esperado: gs://bucket/path/al/archivo")

    bucket_name, blob_path = parts[0], parts[1]

    try:
        gcs_client = _get_storage_client()
        blob       = gcs_client.bucket(bucket_name).blob(blob_path)
        expiry     = timedelta(minutes=15)
        signed_url = blob.generate_signed_url(
            version="v4",
            expiration=expiry,
            method="GET",
        )
        expires_at = (datetime.now(timezone.utc) + expiry).isoformat()

        return JSONResponse(content={
            "signed_url": signed_url,
            "expires_at": expires_at,
            "uri":        uri,
        })
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"Error generando Signed URL de lectura: {exc}")


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=int(os.environ.get("PORT", 8080)))
